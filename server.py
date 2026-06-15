"""codedoc platform server — with JWT auth, background tasks, per-user data, Vue 3 UI."""
from __future__ import annotations

import hashlib
import json
import os
import re
import shutil
import subprocess
import threading
import time
import uuid
from datetime import datetime, timedelta, timezone
from pathlib import Path
from typing import Any

import bcrypt
import jwt
from fastapi import Depends, FastAPI, File, Form, HTTPException, Request, UploadFile
from fastapi.responses import HTMLResponse, JSONResponse, PlainTextResponse
from pydantic import BaseModel
import tempfile
import zipfile

# ---------------------------------------------------------------------------
# App & constants
# ---------------------------------------------------------------------------
JWT_SECRET = os.environ.get("CODEDOC_JWT_SECRET", "codedoc-secret-key-2026-change-me")
JWT_ALGORITHM = "HS256"
JWT_EXPIRE_HOURS = 72

REPOS_DIR = Path("/home/ubuntu/apps/codedoc/repos")
REPOS_DIR.mkdir(parents=True, exist_ok=True)

app = FastAPI(title="codedoc", version="2.0.0", root_path="/codedoc")


try:
    import resource as _resource
except Exception:
    _resource = None


def _clone_limits():
    """clone 子进程降权(fork 后子进程里调;只用已 import 的 resource、不在子进程内 import,
    避免"多线程 + fork"在子进程里取 import 锁死锁)。限 CPU/文件大小/进程数/打开文件。"""
    if _resource is None:
        return
    try:
        _resource.setrlimit(_resource.RLIMIT_CPU, (180, 200))
        _resource.setrlimit(_resource.RLIMIT_FSIZE, (600 * 1024 * 1024, 600 * 1024 * 1024))
        _resource.setrlimit(_resource.RLIMIT_NPROC, (400, 400))
        _resource.setrlimit(_resource.RLIMIT_NOFILE, (2048, 2048))
    except Exception:
        pass


def _rl_key(request) -> str:
    """限流维度:有合法 JWT 按 user,否则按客户端 IP。"""
    auth = request.headers.get("authorization", "")
    if auth.startswith("Bearer "):
        try:
            return "u:%s" % _decode_token(auth[7:]).get("sub")
        except Exception:
            pass
    c = getattr(request, "client", None)
    return "ip:%s" % (c.host if c else "?")


@app.middleware("http")
async def _observability_mw(request: Request, call_next):
    """每请求:request_id + 指标 + 结构化日志 + PG 固定窗口限流(429,跨 worker 精确)。"""
    from codedoc import obs as _obs
    import time as _t
    rid = uuid.uuid4().hex[:16]
    path = request.url.path
    if path.endswith("/health") or path.endswith("/metrics"):
        return await call_next(request)
    if _obs.is_limited(request.method, path):
        import asyncio as _aio
        ok, retry = await _aio.to_thread(_obs.pg_allow, _rl_key(request),
                                         _obs.RL_LIMIT, _obs.RL_WINDOW, _t.time())
        if not ok:
            _obs.METRICS.inc_reject()
            return JSONResponse({"detail": "请求过于频繁,请稍后再试"}, status_code=429,
                                headers={"Retry-After": str(int(retry) + 1), "X-Request-ID": rid})
    _obs.METRICS.inc_inflight(1)
    start = _t.monotonic()
    status = 500
    try:
        resp = await call_next(request)
        status = resp.status_code
        resp.headers["X-Request-ID"] = rid
        return resp
    finally:
        dur = _t.monotonic() - start
        _obs.METRICS.inc_inflight(-1)
        route = request.scope.get("route")
        tmpl = getattr(route, "path", path) if route else path
        _obs.METRICS.observe(request.method, tmpl, status, dur)
        _obs.log_request(rid, request.method, tmpl, status, dur * 1000.0, _rl_key(request))


@app.get("/metrics")
def metrics():
    from codedoc import obs as _obs
    return PlainTextResponse(_obs.METRICS.render(), media_type="text/plain; version=0.0.4")

# ---------------------------------------------------------------------------
# In-memory state
# ---------------------------------------------------------------------------
indexed_repos: dict[str, dict] = {}          # repo_name -> {store, cfg, ...}
task_state: dict[str, dict[str, Any]] = {}   # task_id  -> {status, progress, ...}

# Doc-generation task queue + worker (background, single-worker, FIFO).
# Mirrors the queue model used by devbot's `_run_review_task` — but here we
# enforce serial execution so queue position is predictable for the UI.
_docgen_tasks: dict[str, dict[str, Any]] = {}    # task_id -> 本 worker 处理中任务的本地缓存
_docgen_lock = threading.Lock()
_docgen_worker_thread: threading.Thread | None = None  # single worker
_repo_worker_thread: threading.Thread | None = None    # repo clone/index 领取循环

# ---------------------------------------------------------------------------
# SQLite helpers
# ---------------------------------------------------------------------------

# 生产部署不用 SQLite(4 worker 共写单文件有写锁竞争、CAS 靠文件锁脆弱)。
# 业务表全部落 Postgres(codedoc 库,与图谱/向量/队列同库)。_get_db() 返回一个薄兼容垫片,
# 保持原 sqlite3 调用面:conn.execute("...?...", params) / .fetchone() / .fetchall() /
# .commit() / .close()、row['col'] 与 row[0] 双访问、cursor.rowcount。占位符 ? 自动转 %s。
import psycopg as _psycopg
try:
    from codedoc.graph.graph_persist import PG_DSN as _APP_PG_DSN
except Exception:
    _APP_PG_DSN = os.environ["CODEDOC_PG_DSN"]


class _Row(dict):
    """像 sqlite3.Row:既支持 row['col'] 也支持 row[0]。"""
    __slots__ = ("_vals",)
    def __init__(self, cols, vals):
        super().__init__(zip(cols, vals))
        self._vals = list(vals)
    def __getitem__(self, k):
        if isinstance(k, int):
            return self._vals[k]
        return dict.__getitem__(self, k)


class _Cur:
    def __init__(self, cur):
        self._cur = cur
    def _cols(self):
        return [d.name for d in self._cur.description] if self._cur.description else []
    def fetchone(self):
        r = self._cur.fetchone()
        return _Row(self._cols(), r) if r is not None else None
    def fetchall(self):
        cols = self._cols()
        return [_Row(cols, r) for r in self._cur.fetchall()]
    def __iter__(self):
        return iter(self.fetchall())
    @property
    def rowcount(self):
        return self._cur.rowcount


class _Conn:
    """sqlite3.Connection 兼容垫片(覆盖本项目用到的子集),底层 psycopg/PG。"""
    def __init__(self, conn):
        self._conn = conn
    def execute(self, sql, params=()):
        cur = self._conn.cursor()
        cur.execute(sql.replace("?", "%s"), tuple(params))
        return _Cur(cur)
    def commit(self):
        self._conn.commit()
    def close(self):
        try:
            self._conn.close()
        except Exception:
            pass


def _get_db():
    return _Conn(_psycopg.connect(_APP_PG_DSN))


def _init_db():
    conn = _get_db()
    try:
        for stmt in (
            "CREATE TABLE IF NOT EXISTS users (id BIGINT GENERATED BY DEFAULT AS IDENTITY PRIMARY KEY, "
            "username TEXT UNIQUE NOT NULL, password_hash TEXT NOT NULL, created_at TIMESTAMPTZ DEFAULT now())",
            "CREATE TABLE IF NOT EXISTS user_repos (id BIGINT GENERATED BY DEFAULT AS IDENTITY PRIMARY KEY, "
            "user_id BIGINT NOT NULL, name TEXT NOT NULL, url TEXT NOT NULL, status TEXT DEFAULT 'pending', "
            "nodes INTEGER DEFAULT 0, edges INTEGER DEFAULT 0, task_id TEXT, created_at TIMESTAMPTZ DEFAULT now())",
            "CREATE INDEX IF NOT EXISTS idx_user_repos_user ON user_repos(user_id)",
            "CREATE INDEX IF NOT EXISTS idx_user_repos_task ON user_repos(task_id)",
            "CREATE TABLE IF NOT EXISTS user_docs (id BIGINT GENERATED BY DEFAULT AS IDENTITY PRIMARY KEY, "
            "user_id BIGINT NOT NULL, repo_name TEXT NOT NULL, filename TEXT NOT NULL, "
            "size_bytes INTEGER NOT NULL DEFAULT 0, text_length INTEGER NOT NULL DEFAULT 0, "
            "content TEXT NOT NULL, uploaded_at TIMESTAMPTZ NOT NULL DEFAULT now())",
            "CREATE INDEX IF NOT EXISTS idx_user_docs_user_repo ON user_docs(user_id, repo_name)",
            "CREATE TABLE IF NOT EXISTS qa_history (id BIGINT GENERATED BY DEFAULT AS IDENTITY PRIMARY KEY, "
            "user_id BIGINT NOT NULL, repo TEXT NOT NULL, question TEXT NOT NULL, answer TEXT NOT NULL, "
            "groundedness REAL, total_refs INTEGER, valid_refs_count INTEGER, invalid_refs TEXT, "
            "context_nodes INTEGER, user_docs TEXT, created_at TIMESTAMPTZ DEFAULT now())",
            "CREATE INDEX IF NOT EXISTS idx_qa_history_user_repo ON qa_history(user_id, repo, created_at DESC)",
            "CREATE TABLE IF NOT EXISTS conversation_memory (user_id TEXT, repo TEXT, mem_json TEXT, "
            "updated_at TEXT, version INTEGER NOT NULL DEFAULT 0, PRIMARY KEY(user_id, repo))",
        ):
            conn.execute(stmt)
        conn.commit()
    finally:
        conn.close()

_init_db()

# ---------------------------------------------------------------------------
# Worker-safe repo loader: handles cross-worker state since uvicorn forks.
# Looks up name in user_repos and parses /home/ubuntu/apps/codedoc/repos/<repo>.
# ---------------------------------------------------------------------------

def _ensure_repo_indexed(repo_name: str) -> dict | None:
    """If repo not in this worker's indexed_repos, try parsing from disk."""
    info = indexed_repos.get(repo_name)
    if info and info.get("status") == "ready":
        return info
    # Try to locate on disk: name is "owner/repo" or just "repo"
    candidate_dirs = []
    if "/" in repo_name:
        owner, repo = repo_name.split("/", 1)
        candidate_dirs.append(REPOS_DIR / repo)
    else:
        candidate_dirs.append(REPOS_DIR / repo_name)
    candidate_dirs.append(REPOS_DIR / repo_name.replace("/", "_"))

    for d in candidate_dirs:
        if d.exists() and d.is_dir():
            try:
                from codedoc.config import load_config
                from codedoc.graph.memory_backend import MemoryGraphStore
                from codedoc.parser.runner import parse_repo
                cfg = load_config(d)
                from codedoc.graph import graph_persist
                _snap = graph_persist.load_graph(repo_name)
                if _snap:
                    nodes, edges = _snap
                    print("[graph.persist] loaded graph repo=%s (%d nodes)" % (repo_name, len(nodes)))
                else:
                    nodes, edges = parse_repo(cfg)
                    try:
                        graph_persist.save_graph(repo_name, nodes, edges)
                    except Exception:
                        pass
                store = MemoryGraphStore(cfg)
                store.upsert_nodes(nodes)
                store.upsert_edges(edges)
                info = {
                    "path": str(d),
                    "status": "ready",
                    "nodes": len(nodes),
                    "edges": len(edges),
                    "store": store,
                    "cfg": cfg,
                }
                indexed_repos[repo_name] = info
                return info
            except Exception as _e:
                import logging as _lg
                _lg.getLogger("codedoc").exception("ensure_repo_indexed failed for %s in %s: %s", repo_name, d, _e)
                continue
    return None


# ---------------------------------------------------------------------------
# Auth helpers
# ---------------------------------------------------------------------------

def _hash_password(password: str) -> str:
    return bcrypt.hashpw(password.encode("utf-8"), bcrypt.gensalt()).decode("utf-8")

def _verify_password(password: str, hashed: str) -> bool:
    return bcrypt.checkpw(password.encode("utf-8"), hashed.encode("utf-8"))

def _create_token(user_id: int, username: str) -> str:
    payload = {
        "sub": str(user_id),
        "username": username,
        "exp": datetime.now(timezone.utc) + timedelta(hours=JWT_EXPIRE_HOURS),
    }
    return jwt.encode(payload, JWT_SECRET, algorithm=JWT_ALGORITHM)

def _decode_token(token: str) -> dict:
    try:
        data = jwt.decode(token, JWT_SECRET, algorithms=[JWT_ALGORITHM])
        data["sub"] = int(data["sub"])
        return data
    except jwt.ExpiredSignatureError:
        raise HTTPException(401, "令牌已过期")
    except jwt.InvalidTokenError as e:
        raise HTTPException(401, f"无效令牌: {e}")

def get_current_user(request: Request) -> dict:
    auth = request.headers.get("Authorization", "")
    if not auth.startswith("Bearer "):
        raise HTTPException(401, "未提供认证令牌")
    return _decode_token(auth[7:])

# ---------------------------------------------------------------------------
# Request models
# ---------------------------------------------------------------------------
class RegisterRequest(BaseModel):
    username: str
    password: str

class LoginRequest(BaseModel):
    username: str
    password: str

class RepoAddRequest(BaseModel):
    url: str

class AskRequest(BaseModel):
    question: str
    repo: str = ""
    repos: list[str] = []
    model: str = "auto"

class DocgenRequest(BaseModel):
    repo: str = ""
    repos: list[str] = []  # 多仓文档:>=2 个仓时走跨仓文档生成
    template: str = "default"  # one of DOC_TEMPLATES keys, or "custom"
    sections: list[str] | None = None  # explicit section list (overrides template)

# ---------------------------------------------------------------------------
# Auth endpoints
# ---------------------------------------------------------------------------

@app.post("/api/v1/auth/register")
def register(req: RegisterRequest):
    if len(req.username) < 2 or len(req.username) > 30:
        raise HTTPException(400, "用户名长度需要2-30个字符")
    if len(req.password) < 4:
        raise HTTPException(400, "密码长度至少4个字符")
    conn = _get_db()
    try:
        existing = conn.execute("SELECT id FROM users WHERE username=?", (req.username,)).fetchone()
        if existing:
            raise HTTPException(409, "用户名已存在")
        cur = conn.execute(
            "INSERT INTO users (username, password_hash) VALUES (?, ?) RETURNING id",
            (req.username, _hash_password(req.password)),
        )
        user_id = cur.fetchone()["id"]
        conn.commit()
        token = _create_token(user_id, req.username)
        return {"token": token, "user_id": user_id, "username": req.username}
    finally:
        conn.close()

@app.post("/api/v1/auth/login")
def login(req: LoginRequest):
    conn = _get_db()
    try:
        row = conn.execute("SELECT id, username, password_hash FROM users WHERE username=?",
                           (req.username,)).fetchone()
        if not row:
            raise HTTPException(401, "用户名或密码错误")
        if not _verify_password(req.password, row["password_hash"]):
            raise HTTPException(401, "用户名或密码错误")
        token = _create_token(row["id"], row["username"])
        return {"token": token, "user_id": row["id"], "username": row["username"]}
    finally:
        conn.close()

# ---------------------------------------------------------------------------
# Background clone + index
# ---------------------------------------------------------------------------

def _background_clone_and_index(task_id: str, user_id: int, url: str, name: str, repo_dir: Path):
    """Run git clone + parse + index in a background thread."""
    conn = _get_db()
    try:
        # -- clone --
        task_state[task_id] = {"status": "cloning", "progress": "正在克隆仓库...", "nodes": 0, "edges": 0}
        conn.execute("UPDATE user_repos SET status='cloning' WHERE task_id=?", (task_id,))
        conn.commit()

        if repo_dir.exists():
            shutil.rmtree(repo_dir, ignore_errors=True)
        # Try multiple methods to fetch the repo (China network resilience)
        import zipfile as _zip, io, tempfile
        _last_err = ""
        _success = False
        _clean_url = url.rstrip("/").removesuffix(".git")
        _gh_path = _clean_url.replace("https://github.com/", "").replace("http://github.com/", "")
        _parts = _gh_path.split("/")
        _owner = _parts[0] if len(_parts) >= 1 else ""
        _repo_name = _parts[1] if len(_parts) >= 2 else ""

        # Strategy 1: download zip via curl (handles network resets with retry)
        if _owner and _repo_name:
            _zip_urls = [
                "https://codeload.github.com/" + _owner + "/" + _repo_name + "/zip/refs/heads/main",
                "https://codeload.github.com/" + _owner + "/" + _repo_name + "/zip/refs/heads/master",
                "https://gh-proxy.com/https://github.com/" + _owner + "/" + _repo_name + "/archive/refs/heads/main.zip",
                "https://gh-proxy.com/https://github.com/" + _owner + "/" + _repo_name + "/archive/refs/heads/master.zip",
            ]
            for _zurl in _zip_urls:
                _tmpzip = tempfile.NamedTemporaryFile(suffix=".zip", delete=False).name
                try:
                    _cr = subprocess.run(
                        ["curl", "-sfL", "--retry", "3", "--retry-delay", "2",
                         "--max-time", "120", "-o", _tmpzip, _zurl],
                        capture_output=True, text=True, timeout=140,
                        preexec_fn=_clone_limits, start_new_session=True,
                    )
                    if _cr.returncode != 0 or not os.path.exists(_tmpzip):
                        _last_err = "curl rc=" + str(_cr.returncode) + " " + (_cr.stderr or "")[:200]
                        continue
                    _size = os.path.getsize(_tmpzip)
                    if _size < 1000:
                        _last_err = "zip too small: " + str(_size)
                        continue
                    _z = _zip.ZipFile(_tmpzip)
                    shutil.rmtree(repo_dir, ignore_errors=True)
                    repo_dir.mkdir(parents=True, exist_ok=True)
                    _members = _z.namelist()
                    _top = _members[0].split("/")[0] if _members else ""
                    for _mname in _members:
                        if _mname.endswith("/"):
                            continue
                        _rel = _mname[len(_top)+1:] if _mname.startswith(_top + "/") else _mname
                        if not _rel:
                            continue
                        _dst = repo_dir / _rel
                        _dst.parent.mkdir(parents=True, exist_ok=True)
                        with open(_dst, "wb") as _fo:
                            _fo.write(_z.read(_mname))
                    _success = True
                    break
                except Exception as _e:
                    _last_err = _zurl[:60] + ": " + str(_e)[:200]
                finally:
                    try:
                        os.unlink(_tmpzip)
                    except OSError:
                        pass

        # Strategy 2: git clone via mirrors
        if not _success:
            _git_urls = [
                "https://gh-proxy.com/" + url,
                "https://hub.gitmirror.com/" + url,
                url,
            ]
            for _gurl in _git_urls:
                try:
                    shutil.rmtree(repo_dir, ignore_errors=True)
                    _r = subprocess.run(
                        ["git", "clone", "--depth=1", "--no-tags", _gurl, str(repo_dir)],
                        capture_output=True, text=True, timeout=180,
                        preexec_fn=_clone_limits, start_new_session=True,
                    )
                    if _r.returncode == 0:
                        _success = True
                        break
                    _last_err = (_r.stderr or "")[:300]
                except Exception as _e:
                    _last_err = str(_e)[:300]

        class _R:
            pass
        result = _R()
        result.returncode = 0 if _success else 1
        result.stderr = "" if _success else _last_err
        if result.returncode != 0:
            err_msg = (result.stderr or "")[:500] if hasattr(result, "stderr") else "no stderr"
            import logging as _lg
            _lg.getLogger("codedoc").error("CLONE FAILED url=%s err=%s", url, err_msg)
            task_state[task_id] = {"status": "error", "progress": f"克隆失败: {err_msg}", "nodes": 0, "edges": 0, "error": err_msg}
            conn.execute("UPDATE user_repos SET status='error' WHERE task_id=?", (task_id,))
            conn.commit()
            return

        # -- index --
        task_state[task_id] = {"status": "indexing", "progress": "正在解析代码结构...", "nodes": 0, "edges": 0}
        conn.execute("UPDATE user_repos SET status='indexing' WHERE task_id=?", (task_id,))
        conn.commit()

        from codedoc.config import load_config
        from codedoc.graph.memory_backend import MemoryGraphStore, MemoryGraphQuery
        from codedoc.parser.runner import parse_repo

        cfg = load_config(repo_dir)
        nodes, edges = parse_repo(cfg)
        store = MemoryGraphStore(cfg)
        store.upsert_nodes(nodes)
        store.upsert_edges(edges)

        # 图谱落盘持久化(冷启动秒加载,不再 re-parse)
        try:
            from codedoc.graph import graph_persist
            graph_persist.save_graph(name, nodes, edges)
        except Exception as _ge:
            print('graph persist failed:', _ge)

        # 向量索引(BGE-M3 + pgvector HNSW)
        try:
            from codedoc.index import pg_vectors
            pg_vectors.upsert_repo(name, nodes)
        except Exception as _ve:
            print('vector index failed:', _ve)

        indexed_repos[name] = {
            "path": str(repo_dir),
            "status": "ready",
            "nodes": len(nodes),
            "edges": len(edges),
            "store": store,
            "cfg": cfg,
        }
        task_state[task_id] = {
            "status": "ready",
            "progress": f"索引完成: {len(nodes)} 节点, {len(edges)} 边",
            "nodes": len(nodes),
            "edges": len(edges),
        }
        conn.execute(
            "UPDATE user_repos SET status='ready', nodes=?, edges=? WHERE task_id=?",
            (len(nodes), len(edges), task_id),
        )
        conn.commit()

    except subprocess.TimeoutExpired:
        task_state[task_id] = {"status": "error", "progress": "克隆超时", "nodes": 0, "edges": 0}
        conn.execute("UPDATE user_repos SET status='error' WHERE task_id=?", (task_id,))
        conn.commit()
    except Exception as e:
        task_state[task_id] = {"status": "error", "progress": f"错误: {str(e)[:300]}", "nodes": 0, "edges": 0}
        conn.execute("UPDATE user_repos SET status='error' WHERE task_id=?", (task_id,))
        conn.commit()
    finally:
        conn.close()

# ---------------------------------------------------------------------------
# Repo endpoints (auth required)
# ---------------------------------------------------------------------------

@app.get("/api/v1/repos")
def list_repos(user: dict = Depends(get_current_user)):
    conn = _get_db()
    try:
        rows = conn.execute(
            "SELECT id, name, url, status, nodes, edges, task_id, created_at "
            "FROM user_repos WHERE user_id=? ORDER BY created_at DESC",
            (user["sub"],),
        ).fetchall()
        result = []
        for r in rows:
            tid = r["task_id"]
            # Merge live task_state if exists
            ts = task_state.get(tid, {})
            status = ts.get("status", r["status"])
            nodes = ts.get("nodes", r["nodes"])
            edges = ts.get("edges", r["edges"])
            progress = ts.get("progress", "")
            result.append({
                "id": r["id"],
                "name": r["name"],
                "url": r["url"],
                "status": status,
                "nodes": nodes,
                "edges": edges,
                "task_id": tid,
                "progress": progress,
                "created_at": r["created_at"],
            })
        return {"repos": result}
    finally:
        conn.close()

@app.post("/api/v1/repos")
def add_repo(req: RepoAddRequest, user: dict = Depends(get_current_user)):
    m = re.search(r"github\.com[/:]([^/]+)/([^/.]+)", req.url)
    if not m:
        raise HTTPException(400, "无效的 GitHub 地址")
    owner, repo = m.group(1), m.group(2)
    name = f"{owner}/{repo}"
    task_id = str(uuid.uuid4())
    repo_dir = REPOS_DIR / repo

    conn = _get_db()
    try:
        # Remove existing entry for this user+name if any
        conn.execute("DELETE FROM user_repos WHERE user_id=? AND name=?", (user["sub"], name))
        conn.execute(
            "INSERT INTO user_repos (user_id, name, url, status, task_id) VALUES (?, ?, ?, 'cloning', ?)",
            (user["sub"], name, req.url.strip(), task_id),
        )
        conn.commit()
    finally:
        conn.close()

    # 投递到 PG 派发队列(任意 worker 经 SKIP LOCKED 领取、崩溃可重跑),不再绑定提交方 worker
    from codedoc import repo_jobs
    repo_jobs.enqueue(task_id, user["sub"], name, req.url.strip())
    _start_repo_worker()

    return {"task_id": task_id, "name": name, "status": "cloning"}

def _background_index_only(task_id: str, user_id: int, name: str, repo_dir: Path):
    """Run parse + index on an already-extracted directory (used by upload)."""
    conn = _get_db()
    try:
        task_state[task_id] = {"status": "indexing", "progress": "正在解析代码结构...", "nodes": 0, "edges": 0}
        conn.execute("UPDATE user_repos SET status='indexing' WHERE task_id=?", (task_id,))
        conn.commit()

        from codedoc.config import load_config
        from codedoc.graph.memory_backend import MemoryGraphStore
        from codedoc.parser.runner import parse_repo

        cfg = load_config(repo_dir)
        nodes, edges = parse_repo(cfg)
        store = MemoryGraphStore(cfg)
        store.upsert_nodes(nodes)
        store.upsert_edges(edges)

        # 图谱落盘持久化(冷启动秒加载,不再 re-parse)
        try:
            from codedoc.graph import graph_persist
            graph_persist.save_graph(name, nodes, edges)
        except Exception as _ge:
            print('graph persist failed:', _ge)

        # 向量索引(BGE-M3 + pgvector HNSW)
        try:
            from codedoc.index import pg_vectors
            pg_vectors.upsert_repo(name, nodes)
        except Exception as _ve:
            print('vector index failed:', _ve)

        indexed_repos[name] = {
            "path": str(repo_dir),
            "status": "ready",
            "nodes": len(nodes),
            "edges": len(edges),
            "store": store,
            "cfg": cfg,
        }
        task_state[task_id] = {
            "status": "ready",
            "progress": f"索引完成: {len(nodes)} 节点, {len(edges)} 边",
            "nodes": len(nodes),
            "edges": len(edges),
        }
        conn.execute(
            "UPDATE user_repos SET status='ready', nodes=?, edges=? WHERE task_id=?",
            (len(nodes), len(edges), task_id),
        )
        conn.commit()
    except Exception as e:
        task_state[task_id] = {"status": "error", "progress": f"错误: {str(e)[:300]}", "nodes": 0, "edges": 0}
        conn.execute("UPDATE user_repos SET status='error' WHERE task_id=?", (task_id,))
        conn.commit()
    finally:
        conn.close()


def _start_repo_worker():
    """启动本进程的 repo 领取循环(若未在跑)。每个 uvicorn worker 一个,经 PG
    FOR UPDATE SKIP LOCKED 抢单 → 真正跨 worker 的克隆/索引队列。env CODEDOC_REPO_WORKER=0 可关。"""
    global _repo_worker_thread
    if os.environ.get("CODEDOC_REPO_WORKER", "1") != "1":
        return
    if _repo_worker_thread and _repo_worker_thread.is_alive():
        return
    _repo_worker_thread = threading.Thread(target=_repo_worker_loop, daemon=True)
    _repo_worker_thread.start()


def _repo_worker_loop():
    """持续从 PG 队列领取克隆/索引作业(多 worker 安全),逐个执行。启动先做崩溃孤儿重建。"""
    import time as _time
    from codedoc import repo_jobs
    try:
        repo_jobs.ensure_schema()
        n = repo_jobs.reclaim_orphans(900)
        if n:
            print("[repo] reclaimed %d orphan job(s)" % n)
    except Exception:
        pass
    while True:
        row = None
        try:
            row = repo_jobs.claim_one()
        except Exception:
            row = None
        if not row:
            _time.sleep(3)
            continue
        task_id, user_id, name, url = row["task_id"], row["user_id"], row["name"], row["url"]
        repo_dir = REPOS_DIR / name.split("/")[-1]
        try:
            if url.startswith("local:"):
                _background_index_only(task_id, user_id, name, repo_dir)
            else:
                _background_clone_and_index(task_id, user_id, url, name, repo_dir)
            repo_jobs.complete(task_id)
            print("[repo] done job=%s name=%s" % (task_id, name))
        except Exception as e:  # noqa: BLE001
            try:
                repo_jobs.fail(task_id, str(e))
            except Exception:
                pass


# Upload limit: 50 MB
UPLOAD_MAX_BYTES = 50 * 1024 * 1024


def _sanitize_repo_name(name: str) -> str:
    """Keep alphanumeric, dash, underscore only."""
    cleaned = re.sub(r"[^A-Za-z0-9_-]+", "-", name.strip())
    cleaned = cleaned.strip("-_")
    return cleaned or "repo"


@app.post("/api/v1/repos/upload")
async def upload_repo(
    file: UploadFile = File(...),
    name: str = Form(...),
    user: dict = Depends(get_current_user),
):
    """Upload a ZIP file as a new repo, extract and index it."""
    if not file.filename or not file.filename.lower().endswith(".zip"):
        raise HTTPException(400, "仅支持 .zip 文件")

    raw_name = name.strip()
    if not raw_name:
        raise HTTPException(400, "请提供仓库名称")
    safe_name = _sanitize_repo_name(raw_name)
    repo_dir = REPOS_DIR / safe_name
    display_name = f"local/{safe_name}"

    if repo_dir.exists():
        raise HTTPException(409, "仓库名已存在")

    # Read upload to a temp file while enforcing the size limit.
    tmp_fd, tmp_path_str = tempfile.mkstemp(prefix="codedoc_upload_", suffix=".zip", dir="/tmp")
    tmp_path = Path(tmp_path_str)
    size_total = 0
    try:
        with os.fdopen(tmp_fd, "wb") as out_fp:
            while True:
                chunk = await file.read(1024 * 1024)
                if not chunk:
                    break
                size_total += len(chunk)
                if size_total > UPLOAD_MAX_BYTES:
                    raise HTTPException(413, "ZIP 文件超过 50MB")
                out_fp.write(chunk)

        # Extract
        repo_dir.mkdir(parents=True, exist_ok=True)
        try:
            with zipfile.ZipFile(tmp_path, "r") as zf:
                # Basic safety: reject paths with absolute or .. components
                for info in zf.infolist():
                    target = (repo_dir / info.filename).resolve()
                    if not str(target).startswith(str(repo_dir.resolve())):
                        raise HTTPException(400, "ZIP 文件包含非法路径")
                zf.extractall(repo_dir)
        except zipfile.BadZipFile:
            shutil.rmtree(repo_dir, ignore_errors=True)
            raise HTTPException(400, "无效的 ZIP 文件")

        # If the zip contains a single top-level dir, flatten it for nicer paths
        entries = list(repo_dir.iterdir())
        if len(entries) == 1 and entries[0].is_dir():
            inner = entries[0]
            for child in inner.iterdir():
                shutil.move(str(child), str(repo_dir / child.name))
            shutil.rmtree(inner, ignore_errors=True)
    finally:
        try:
            tmp_path.unlink(missing_ok=True)
        except Exception:
            pass

    # Record + kick off indexing
    task_id = str(uuid.uuid4())
    conn = _get_db()
    try:
        conn.execute("DELETE FROM user_repos WHERE user_id=? AND name=?", (user["sub"], display_name))
        conn.execute(
            "INSERT INTO user_repos (user_id, name, url, status, task_id) VALUES (?, ?, ?, 'indexing', ?)",
            (user["sub"], display_name, f"local:{raw_name}", task_id),
        )
        conn.commit()
    finally:
        conn.close()

    from codedoc import repo_jobs
    repo_jobs.enqueue(task_id, user["sub"], display_name, f"local:{raw_name}")
    _start_repo_worker()

    return {"task_id": task_id, "name": display_name, "status": "queued"}


@app.get("/api/v1/repos/{task_id}/status")
def repo_status(task_id: str, user: dict = Depends(get_current_user)):
    # Check ownership
    conn = _get_db()
    try:
        row = conn.execute(
            "SELECT * FROM user_repos WHERE task_id=? AND user_id=?",
            (task_id, user["sub"]),
        ).fetchone()
        if not row:
            raise HTTPException(404, "任务不存在")
    finally:
        conn.close()

    ts = task_state.get(task_id, {})
    return {
        "task_id": task_id,
        "name": row["name"],
        "status": ts.get("status", row["status"]),
        "progress": ts.get("progress", ""),
        "nodes": ts.get("nodes", row["nodes"]),
        "edges": ts.get("edges", row["edges"]),
    }

@app.delete("/api/v1/repos/{owner}/{repo}")
def delete_repo(owner: str, repo: str, user: dict = Depends(get_current_user)):
    name = f"{owner}/{repo}"
    conn = _get_db()
    try:
        conn.execute("DELETE FROM user_repos WHERE user_id=? AND name=?", (user["sub"], name))
        conn.commit()
    finally:
        conn.close()
    indexed_repos.pop(name, None)
    repo_path = REPOS_DIR / repo
    if repo_path.exists():
        shutil.rmtree(repo_path, ignore_errors=True)
    return {"status": "deleted"}

# ---------------------------------------------------------------------------
# User-uploaded reference documents (PDF / DOCX / MD / TXT / HTML)
# ---------------------------------------------------------------------------

_USER_DOC_EXTS = (".pdf", ".docx", ".md", ".txt", ".html", ".htm")
_USER_DOC_MAX_BYTES = 20 * 1024 * 1024  # 20MB
_USER_DOC_MAX_TEXT = 500_000            # cap stored text


def _user_owns_repo(user_id: int, repo_name: str) -> bool:
    conn = _get_db()
    try:
        row = conn.execute(
            "SELECT id FROM user_repos WHERE user_id=? AND name=?",
            (user_id, repo_name),
        ).fetchone()
        return row is not None
    finally:
        conn.close()


def _extract_user_doc_text(filename: str, content: bytes) -> str:
    fname = filename.lower()
    if fname.endswith(".pdf"):
        try:
            from pypdf import PdfReader
        except Exception as e:  # pragma: no cover
            raise HTTPException(500, f"PDF 解析依赖未安装: {e}")
        import io
        reader = PdfReader(io.BytesIO(content))
        pages = []
        for p in reader.pages:
            try:
                pages.append(p.extract_text() or "")
            except Exception:
                continue
        return "\n\n".join(pages)
    if fname.endswith(".docx"):
        try:
            from docx import Document as DocxDoc
        except Exception as e:  # pragma: no cover
            raise HTTPException(500, f"DOCX 解析依赖未安装: {e}")
        import io
        d = DocxDoc(io.BytesIO(content))
        return "\n".join(p.text for p in d.paragraphs)
    if fname.endswith(".md") or fname.endswith(".txt"):
        return content.decode("utf-8", errors="replace")
    if fname.endswith(".html") or fname.endswith(".htm"):
        raw = content.decode("utf-8", errors="replace")
        no_tags = re.sub(r'<script[\s\S]*?</script>', ' ', raw, flags=re.IGNORECASE)
        no_tags = re.sub(r'<style[\s\S]*?</style>', ' ', no_tags, flags=re.IGNORECASE)
        no_tags = re.sub(r'<[^>]+>', ' ', no_tags)
        return re.sub(r'\s+', ' ', no_tags).strip()
    raise HTTPException(400, "不支持的文件类型")


@app.post("/api/v1/repos/{owner}/{repo}/docs/upload")
async def upload_user_doc(
    owner: str,
    repo: str,
    file: UploadFile = File(...),
    user: dict = Depends(get_current_user),
):
    repo_name = f"{owner}/{repo}"
    if not _user_owns_repo(user["sub"], repo_name):
        raise HTTPException(404, "仓库不存在或无权访问")

    fname = (file.filename or "upload").strip()
    flow = fname.lower()
    if not any(flow.endswith(ext) for ext in _USER_DOC_EXTS):
        raise HTTPException(400, "支持的格式：PDF / DOCX / MD / TXT / HTML")

    content = await file.read()
    if len(content) > _USER_DOC_MAX_BYTES:
        raise HTTPException(413, "文件超过 20MB 上限")
    if not content:
        raise HTTPException(400, "文件为空")

    text = _extract_user_doc_text(fname, content)
    text = (text or "").strip()
    if not text:
        raise HTTPException(400, "未能从文件中提取出文本内容")
    stored_text = text[:_USER_DOC_MAX_TEXT]

    conn = _get_db()
    try:
        cur = conn.execute(
            "INSERT INTO user_docs (user_id, repo_name, filename, size_bytes, text_length, content) "
            "VALUES (?, ?, ?, ?, ?, ?) RETURNING id",
            (user["sub"], repo_name, fname, len(content), len(text), stored_text),
        )
        doc_id = cur.fetchone()["id"]
        conn.commit()
    finally:
        conn.close()

    return {
        "id": doc_id,
        "filename": fname,
        "size": len(content),
        "text_length": len(text),
        "status": "ok",
    }


@app.get("/api/v1/repos/{owner}/{repo}/docs")
def list_user_docs(owner: str, repo: str, user: dict = Depends(get_current_user)):
    repo_name = f"{owner}/{repo}"
    if not _user_owns_repo(user["sub"], repo_name):
        raise HTTPException(404, "仓库不存在或无权访问")
    conn = _get_db()
    try:
        rows = conn.execute(
            "SELECT id, filename, size_bytes, text_length, uploaded_at FROM user_docs "
            "WHERE user_id=? AND repo_name=? ORDER BY uploaded_at DESC",
            (user["sub"], repo_name),
        ).fetchall()
    finally:
        conn.close()
    return {
        "docs": [
            {
                "id": r["id"],
                "filename": r["filename"],
                "size": r["size_bytes"],
                "text_length": r["text_length"],
                "uploaded_at": r["uploaded_at"],
            }
            for r in rows
        ]
    }


@app.delete("/api/v1/repos/{owner}/{repo}/docs/{doc_id}")
def delete_user_doc(owner: str, repo: str, doc_id: int, user: dict = Depends(get_current_user)):
    repo_name = f"{owner}/{repo}"
    if not _user_owns_repo(user["sub"], repo_name):
        raise HTTPException(404, "仓库不存在或无权访问")
    conn = _get_db()
    try:
        cur = conn.execute(
            "DELETE FROM user_docs WHERE id=? AND user_id=? AND repo_name=?",
            (doc_id, user["sub"], repo_name),
        )
        conn.commit()
        if cur.rowcount == 0:
            raise HTTPException(404, "文档不存在")
    finally:
        conn.close()
    return {"status": "ok"}


# ---------------------------------------------------------------------------
# QA endpoint (auth required)
# ---------------------------------------------------------------------------

def _registry_for(gq, repo, repo_root=None):
    """用 Web 后端(内存图 gq + pgvector)建统一工具 registry —— 与 MCP/CLI 共用一份工具实现。"""
    from codedoc.tools.registry import build_registry
    from codedoc.index import pg_vectors
    if repo_root is None:
        info = indexed_repos.get(repo)
        repo_root = (info or {}).get("path")
    return build_registry(getattr(gq, "cfg", None), gq,
                          vec_search=lambda q, k: pg_vectors.query(repo, q, top_k=k),
                          repo_root=repo_root)


def _retrieve_context(store, gq, question: str, max_nodes: int = 25, repo: str | None = None, smem=None):
    """Multi-strategy retrieval for QA context. Returns list of dicts: {node, score, reason}."""
    import re as _re
    hits = []
    seen = set()

    def add(node, score, reason):
        if not node:
            return
        nid = getattr(node, "id", None)
        if not nid or nid in seen:
            return
        seen.add(nid)
        hits.append({"node": node, "score": score, "reason": reason})

    # Strategy 0: 语义召回经统一工具层 search(混合 向量+全文);失败回退直连 pgvector
    if repo:
        _used_tool = False
        try:
            _reg = _registry_for(gq, repo)
            _res = _reg.call("search", {"query": question, "top_k": 25})
            if _res.get("ok"):
                for it in _res.get("items", []):
                    n = gq.get_node(it["node_id"])
                    if n:
                        add(n, 5.0 + float(it.get("score") or 0), "tool:search")
                _used_tool = True
        except Exception:
            _used_tool = False
        if not _used_tool:
            try:
                from codedoc.index import pg_vectors
                for vh in pg_vectors.query(repo, question, top_k=25):
                    n = gq.get_node(vh["node_id"])
                    if n:
                        add(n, 5.0 + float(vh.get("score") or 0), "vector")
            except Exception:
                pass

    # 记忆双路:短追问(指代/省略主语)+ 有锚点 -> 补锚点的 callers/callees(关系答案)
    if repo and smem is not None and len((question or '').strip()) < 16:
        _a = smem.current_anchor()
        if _a and _a.get('id'):
            try:
                _mreg = _registry_for(gq, repo)
                for _rel in ('callers', 'callees'):
                    for _it in _mreg.call(_rel, {'node_id': _a['id']}).get('items', [])[:8]:
                        _mn = gq.get_node(_it['node_id'])
                        if _mn:
                            add(_mn, 8.0, 'memory:' + _rel)  # 置顶:高于搜索(<=6.0),防被 top-N 截断
            except Exception:
                pass

    # Strategy 1: fulltext over the whole question
    try:
        for n, s in gq.fulltext_search(question, limit=10):
            add(n, s * 2.0, "fulltext")
    except Exception:
        pass

    # Strategy 2: extract identifiers from question and search each
    STOP = {"什么", "怎么", "如何", "为什么",
            "what", "how", "the", "and", "for", "with", "from", "this", "that",
            "are", "have", "has", "use", "used", "using", "click", "flask"}
    tokens = _re.findall(r"[A-Z][a-zA-Z0-9]+|[a-z_]+_[a-z_]+|[A-Za-z]{4,}", question or "")
    for tok in set(tokens):
        if len(tok) < 3 or tok.lower() in STOP:
            continue
        try:
            for n, s in gq.fulltext_search(tok, limit=3):
                add(n, s, f"keyword:{tok}")
        except Exception:
            pass

    # Strategy 3: type-targeted sample of major kinds
    kinds_to_pull = ["class", "route_handler", "method", "function"]
    for kind in kinds_to_pull:
        count = 0
        for n_dict in store.nodes.values():
            if n_dict.get("kind") != kind:
                continue
            try:
                n = gq.get_node(n_dict["id"])
            except Exception:
                n = None
            if not n:
                continue
            add(n, 0.3, f"sample:{kind}")
            count += 1
            if count >= 8:
                break

    hits.sort(key=lambda h: -h["score"])
    hits = hits[:max_nodes]
    # BGE-reranker 精排
    if repo and len(hits) > 1:
        try:
            from codedoc.index import pg_vectors
            docs = [((getattr(h['node'], 'qualified_name', '') or '') + ' ' +
                     (getattr(h['node'], 'signature', '') or '') + ' ' +
                     (getattr(h['node'], 'docstring', '') or ''))[:500] for h in hits]
            order = dict(pg_vectors.rerank(question, docs))
            if order:
                hits = [hits[i] for i in sorted(range(len(hits)), key=lambda i: -order.get(i, -1.0))]
        except Exception:
            pass
    return hits


# ---- 关系关键词 -> 图谱边类型(关系级核验) ----
_REL_KEYWORDS = [
    (("继承", "extends", "派生自", "基类", "子类化"), "extends"),
    (("实现", "implements"), "implements"),
    (("调用", "calls", "调了", "会调", "invoke"), "calls"),
    (("注入", "autowired", "inject"), "bean_inject"),
    (("导入", "imports", "引入"), "imports"),
    (("包含", "定义了", "contains"), "contains"),
]

# 既是英文常用词又可能是符号名:当英文用时不误判为命中(降假阳性)
_COMMON_WORDS = {
    "add", "get", "set", "put", "pop", "run", "map", "key", "new", "old", "app", "name",
    "open", "close", "copy", "main", "test", "data", "item", "list", "dict", "type",
    "value", "index", "process", "push", "clear", "load", "save", "call", "next", "read",
    "write", "init", "true", "false", "none", "self", "args", "kwargs", "print", "object",
}


def _build_graph_index(stores):
    """建核验所需索引:短名->限定名集、名->节点id集、限定名集、文件集、id->节点、边表。"""
    name_to_qnames, name_to_ids = {}, {}
    qname_set, file_set, id_to_node, edges = set(), set(), {}, []
    for _st in stores:
        for nid, n in _st.nodes.items():
            id_to_node[nid] = n
            nm = n.get("name") or ""
            qn = n.get("qualified_name") or ""
            if qn:
                qname_set.add(qn)
                short = qn.split(".")[-1]
                name_to_qnames.setdefault(short, set()).add(qn)
                name_to_ids.setdefault(short, set()).add(nid)
                name_to_ids.setdefault(qn, set()).add(nid)
            if nm:
                name_to_qnames.setdefault(nm, set()).add(qn or nm)
                name_to_ids.setdefault(nm, set()).add(nid)
            f = n.get("file") or ""
            if f:
                file_set.add(f); file_set.add(f.split("/")[-1])
        edges.extend(_st.edges)
    return {"name_to_qnames": name_to_qnames, "name_to_ids": name_to_ids,
            "qname_set": qname_set, "file_set": file_set,
            "id_to_node": id_to_node, "edges": edges}


def _classify_ref(ref, idx):
    """单个引用分档:exact(精确限定命中) / unique(唯一短名) / ambiguous(一名多义) / file / miss。"""
    if ref in idx["qname_set"]:
        return "exact"
    if "." in ref and any(q == ref or q.endswith("." + ref) for q in idx["qname_set"]):
        return "exact"
    if ref in idx["file_set"]:
        return "file"
    short = ref.split(".")[-1]
    qs = {q for q in (idx["name_to_qnames"].get(short) or set()) if q}
    if not qs:
        return "file" if short in idx["file_set"] else "miss"
    return "unique" if len(qs) <= 1 else "ambiguous"


def _verify_relationships(text, idx):
    """关系级核验(确定性、启发式):文本里『A <关系词> B』-> 查图谱这条边在不在。"""
    import re as _re
    checks = []
    edge_pairs = {}
    for e in idx["edges"]:
        edge_pairs.setdefault(e.get("kind"), set()).add((e.get("src"), e.get("dst")))
    for kws, kind in _REL_KEYWORDS:
        for kw in kws:
            pat = _re.compile(r"`([\w.]+)`[^`]{0,18}" + _re.escape(kw) + r"[^`]{0,18}`([\w.]+)`")
            for m in pat.finditer(text or ""):
                ra, rb = m.group(1), m.group(2)
                if ra.split(".")[-1] == rb.split(".")[-1]:
                    continue  # 跳过自指(A 关系 A 无意义)
                ida = (idx["name_to_ids"].get(ra) or set()) | (idx["name_to_ids"].get(ra.split(".")[-1]) or set())
                idb = (idx["name_to_ids"].get(rb) or set()) | (idx["name_to_ids"].get(rb.split(".")[-1]) or set())
                if not ida or not idb:
                    continue  # 符号本身不在,交给符号级核验,关系级不重复罚
                pairs = edge_pairs.get(kind, set())
                ok = any((x, y) in pairs for x in ida for y in idb)
                checks.append({"a": ra, "b": rb, "rel": kw, "kind": kind, "supported": ok})
    return checks


def _verify_response(answer: str, store, extra_stores=None) -> dict:
    """符号级 + 关系级 grounding 核验(确定性,向后兼容旧字段)。
    符号级:反引号引用分 exact/unique/ambiguous/file/miss,exact+unique+file 记命中,
            歧义(一名多义)单列不算干净命中,英文常用词当英文用时不计入分母;
    关系级:文本里『A 继承/调用/... B』去图谱查边,未命中=编造关系。"""
    import re as _re
    refs = _re.findall(r"`([A-Za-z_][A-Za-z0-9_.]*)`", answer or "")
    idx = _build_graph_index([store] + list(extra_stores or []))
    tiers = {"exact": [], "unique": [], "ambiguous": [], "file": [], "miss": [], "skipped": []}
    for r in refs:
        short = r.split(".")[-1].lower()
        if "." not in r and r.islower() and short in _COMMON_WORDS:
            tiers["skipped"].append(r)
            continue
        tiers[_classify_ref(r, idx)].append(r)
    countable = sum(len(tiers[k]) for k in ("exact", "unique", "ambiguous", "file", "miss"))
    exists = len(tiers["exact"]) + len(tiers["unique"]) + len(tiers["file"]) + len(tiers["ambiguous"])
    pinned = len(tiers["exact"]) + len(tiers["unique"]) + len(tiers["file"])
    groundedness = (exists / countable) if countable else 1.0   # 存在率:引用的符号真实存在(没编造)
    specificity = (pinned / countable) if countable else 1.0    # 可精确定位率:能唯一对上一个符号
    rel_checks = _verify_relationships(answer or "", idx)
    rel_bad = [c for c in rel_checks if not c["supported"]]
    words = max(1, len((answer or "").split()))
    return {
        "total_refs": countable,
        "valid_refs": tiers["exact"] + tiers["unique"] + tiers["file"],
        "invalid_refs": tiers["miss"],
        "ambiguous_refs": tiers["ambiguous"],
        "groundedness": groundedness,
        "specificity": specificity,
        "tiers": {k: len(v) for k, v in tiers.items()},
        "relationship_total": len(rel_checks),
        "relationship_unsupported": rel_bad,
        "cite_density": round(countable / words * 100, 2),
    }


def _verify_prose_deep(text, stores, llm, max_symbols=12, extra_evidence=""):
    """深度语义核验(LLM 裁判):取引用符号的真实 docstring/签名当证据,判论断是否被支持。
    返回 {judged, unsupported:[{claim,why}]}。非确定,失败安全返回 0。"""
    import re as _re, json as _json
    if not text or llm is None:
        return {"judged": 0, "unsupported": []}
    try:
        from codedoc.agents.llm import ChatMessage as _CM
    except Exception:
        return {"judged": 0, "unsupported": []}
    idx = _build_graph_index(list(stores))
    refs = _re.findall(r"`([A-Za-z_][A-Za-z0-9_.]*)`", text)
    evid, seen = [], set()
    for r in refs:
        ids = idx["name_to_ids"].get(r) or idx["name_to_ids"].get(r.split(".")[-1])
        if not ids:
            continue
        nid = sorted(ids)[0]
        if nid in seen:
            continue
        seen.add(nid)
        n = idx["id_to_node"][nid]
        ev = (n.get("docstring") or n.get("signature") or "(无 docstring)").strip().replace("\n", " ")
        evid.append("- `%s` (%s): %s" % (n.get("qualified_name") or n.get("name"), n.get("kind", ""), ev[:160]))
        if len(evid) >= max_symbols:
            break
    if extra_evidence:
        evid.insert(0, "(确定性事实,已由代码确定性抽取、视为真实可信,不要标记被它支持的论断)" + extra_evidence)
    if not evid:
        return {"judged": 0, "unsupported": []}
    prompt = ("下面是一段技术说明,以及它引用符号的真实定义(证据)。逐条找出说明里"
              "**与证据矛盾、或证据无法支持**的断言。只返回 JSON 数组,每项 "
              "{\"claim\":\"...\",\"why\":\"...\"};没有问题就返回 []。\n\n"
              "【说明】\n" + text[:1800] + "\n\n【证据】\n" + "\n".join(evid))
    try:
        resp = llm.chat([_CM("system", "你是严格的技术事实核查员,只依据给定证据判断,不放过编造,也不无中生有。"),
                         _CM("user", prompt)], max_tokens=700) or ""
        m = _re.search(r"\[.*\]", resp, _re.DOTALL)
        arr = _json.loads(m.group(0)) if m else []
        bad = [x for x in arr if isinstance(x, dict) and x.get("claim")]
        return {"judged": len(evid), "unsupported": bad}
    except Exception:
        return {"judged": len(evid), "unsupported": []}


def _ensure_convmem_table(conn):
    conn.execute("CREATE TABLE IF NOT EXISTS conversation_memory ("
                 "user_id TEXT, repo TEXT, mem_json TEXT, updated_at TEXT, "
                 "version INTEGER NOT NULL DEFAULT 0, "
                 "PRIMARY KEY(user_id, repo))")


def _merge_mem(a, b):
    """并集合并两份会话记忆——只做加法、绝不漏节点。b(本轮请求)叠在 a(并发别人写的)之上。
    用于乐观锁冲突后的 reload→union→重试。"""
    from codedoc.memory.manager import MemoryManager
    m = MemoryManager(getattr(b, "anchor_history", 16), getattr(b, "semantic_focus", 8))
    # focus:按 name 取并集,重名取较高分(避免重复累加)
    m.focus = {n: dict(d) for n, d in (getattr(a, "focus", None) or {}).items()}
    for n, d in (getattr(b, "focus", None) or {}).items():
        if n in m.focus:
            m.focus[n]["score"] = max(m.focus[n].get("score", 0.0), d.get("score", 0.0))
            if not m.focus[n].get("node_id"):
                m.focus[n]["node_id"] = d.get("node_id")
        else:
            m.focus[n] = dict(d)
    # anchors:a 在前 b 在后(b 的最后一个=最新 current),去相邻重复,保留末 N 个
    merged_anchors = []
    for an in (getattr(a, "anchors", None) or []) + (getattr(b, "anchors", None) or []):
        if merged_anchors and merged_anchors[-1].get("id") == an.get("id"):
            continue
        merged_anchors.append(an)
    m.anchors = merged_anchors[-m.anchor_history:]
    # notes:并集保序去重,保留末 8 条
    return m


def _load_session_mem(user_id, repo):
    from codedoc.memory.manager import MemoryManager
    conn = _get_db()
    try:
        _ensure_convmem_table(conn)
        row = conn.execute("SELECT mem_json, version FROM conversation_memory WHERE user_id=? AND repo=?",
                           (str(user_id), repo)).fetchone()
        if row and row[0]:
            m = MemoryManager.from_json(json.loads(row[0]))
            m._loaded_version = int(row[1] or 0)   # 记下加载版本,保存时做 CAS
            return m
    except Exception:
        pass
    finally:
        conn.close()
    m = MemoryManager()
    m._loaded_version = 0
    return m


def _save_session_mem(user_id, repo, mem, max_retries=5):
    """乐观锁 + CAS 保存,防多 worker / 多并发请求丢更新。
    UPDATE ... WHERE version=expected;冲突(rowcount=0 且行已存在)→ reload 当前→union 合并→重试。"""
    import datetime as _dt
    from codedoc.memory.manager import MemoryManager
    expected = int(getattr(mem, "_loaded_version", 0) or 0)
    cur_mem = mem
    conn = _get_db()
    try:
        _ensure_convmem_table(conn)
        for _ in range(max_retries):
            payload = json.dumps(cur_mem.to_json(), ensure_ascii=False)
            now = _dt.datetime.utcnow().isoformat()
            c = conn.execute(
                "UPDATE conversation_memory SET mem_json=?, version=version+1, updated_at=? "
                "WHERE user_id=? AND repo=? AND version=?",
                (payload, now, str(user_id), repo, expected))
            if c.rowcount and c.rowcount > 0:
                conn.commit()
                try:
                    mem._loaded_version = expected + 1
                except Exception:
                    pass
                return True
            # rowcount=0:要么没这行(首次写),要么 version 被并发推进了
            row = conn.execute(
                "SELECT version, mem_json FROM conversation_memory WHERE user_id=? AND repo=?",
                (str(user_id), repo)).fetchone()
            if row is None:
                try:
                    conn.execute(
                        "INSERT INTO conversation_memory (user_id, repo, mem_json, updated_at, version) "
                        "VALUES (?,?,?,?,1)", (str(user_id), repo, payload, now))
                    conn.commit()
                    try:
                        mem._loaded_version = 1
                    except Exception:
                        pass
                    return True
                except Exception:
                    try:
                        conn.rollback()
                    except Exception:
                        pass
                    expected = 0
                    continue   # 并发抢插,下一轮走 UPDATE 路径
            else:
                # 冲突:别人写了更新的版本 → 取当前→union 合并→以新版本重试
                cur_version = int(row[0] or 0)
                try:
                    other = MemoryManager.from_json(json.loads(row[1]) if row[1] else {})
                except Exception:
                    other = MemoryManager()
                cur_mem = _merge_mem(other, mem)
                expected = cur_version
                continue
        return False   # 极端高竞争重试耗尽:本轮不落,绝不静默覆盖别人
    except Exception:
        try:
            conn.rollback()
        except Exception:
            pass
        return False
    finally:
        conn.close()


_SUMMARY_WINDOW = int(os.environ.get("CODEDOC_SUMMARY_WINDOW", "8"))   # 保留最近 N 轮在窗口;更老的折叠进滚动摘要


def _summarize_overflow(user_id, repo):
    """把超出最近窗口的老对话轮折叠进会话记忆的 rolling summary(异步 worker 调用)。"""
    conn = _get_db()
    try:
        rows = conn.execute(
            "SELECT question, answer FROM qa_history WHERE user_id=? AND repo=? ORDER BY id ASC",
            (str(user_id), repo)).fetchall()
    except Exception:
        rows = []
    finally:
        conn.close()
    if not rows or len(rows) <= _SUMMARY_WINDOW:
        return False   # 没有溢出轮
    overflow = rows[:-_SUMMARY_WINDOW]                # 除最近窗口外的老轮
    info = _ensure_repo_indexed(repo)
    if not info:
        return False
    from codedoc.agents.llm import ChatMessage, build_llm
    llm = build_llm(info["cfg"])
    prior = (getattr(_load_session_mem(user_id, repo), "summary", "") or "").strip()
    convo = "\n".join("Q: %s\nA: %s" % (r[0], (r[1] or "")[:400]) for r in overflow[-30:])
    prompt = ("已有摘要:\n%s\n\n新增更早的对话轮:\n%s\n\n"
              "请把两者融合成一段不超过 200 字的滚动摘要,只保留对理解后续追问有用的话题、"
              "涉及的代码符号与结论,中文输出,不要逐条罗列。" % (prior or "(无)", convo))
    try:
        resp = llm.chat([
            ChatMessage("system", "你是会话摘要助手,产出简洁、信息密度高的滚动摘要。"),
            ChatMessage("user", prompt),
        ])
        new_summary = (getattr(resp, "content", None) or getattr(resp, "text", None) or str(resp)).strip()
    except Exception:
        return False
    if not new_summary:
        return False
    # 用 CAS 保存:重载→设 summary→存(与并发写安全共存、不丢更新)
    mem2 = _load_session_mem(user_id, repo)
    mem2.summary = new_summary[:1000]
    return _save_session_mem(user_id, repo, mem2)


def _start_summary_worker():
    """后台 worker:领取 memory_summary_jobs(FOR UPDATE SKIP LOCKED,多 worker 安全),
    折叠溢出轮为滚动摘要。启动先做崩溃孤儿重建。env CODEDOC_SUMMARY_WORKER=0 可关。"""
    if os.environ.get("CODEDOC_SUMMARY_WORKER", "1") != "1":
        return
    import threading
    import time as _time
    from codedoc.memory import summary_queue as _sq

    def _loop():
        try:
            _sq.ensure_schema()
            n = _sq.reclaim_orphans(300)
            if n:
                print("[summary] reclaimed %d orphan job(s)" % n)
        except Exception:
            pass
        while True:
            job = None
            try:
                job = _sq.claim_one()
            except Exception:
                job = None
            if not job:
                _time.sleep(5)
                continue
            jid, uid, repo = job[0], job[1], job[2]
            try:
                _summarize_overflow(uid, repo)
                _sq.complete(jid)
                print("[summary] done job=%s user=%s repo=%s" % (jid, uid, repo))
            except Exception as _e:
                _sq.fail(jid, str(_e))

    threading.Thread(target=_loop, name="summary-worker", daemon=True).start()
    print("[summary] worker started")


def _mem_context(mem):
    parts = []
    _s = (getattr(mem, "summary", "") or "").strip()
    if _s:
        parts.append("早前对话滚动摘要:" + _s)
    a = mem.current_anchor()
    if a:
        parts.append("当前话题锚点:`%s`" % a.get("name"))
    chain = [x.get("name") for x in mem.anchors[-4:] if x.get("name")]
    if len(chain) > 1:
        parts.append("话题链:" + " -> ".join(chain))
    tf = mem.top_focus(5, floor=0.3)
    if tf:
        parts.append("最近活跃符号:" + "、".join("%s(%.2f)" % (n, sc) for n, sc in tf))
    if not parts:
        return ""
    return "## 会话记忆(用于理解追问里的指代/省略主语,不要照抄这些符号)\n" + "\n".join(parts)


def _history_window(user_id, repo, budget=2000):
    conn = _get_db()
    try:
        rows = conn.execute("SELECT question, answer FROM qa_history WHERE user_id=? AND repo=? "
                            "ORDER BY id DESC LIMIT 12", (user_id, repo)).fetchall()
    except Exception:
        rows = []
    finally:
        conn.close()
    if not rows:
        return ""
    try:
        import tiktoken
        _enc = tiktoken.get_encoding("cl100k_base")
        _tok = lambda t: len(_enc.encode(t or ""))
    except Exception:
        _tok = lambda t: len(t or "") // 3
    picked, total = [], 0
    for r in rows:
        q, a = r[0], r[1]
        seg = "Q: %s\nA: %s" % (q, (a or "")[:300])
        c = _tok(seg)
        if total + c > budget and picked:
            break
        picked.append(seg); total += c
    picked.reverse()
    return "## 最近对话(token 预算窗口,仅供理解上下文)\n" + "\n".join(picked)


def _run_multi_agent(question, repos):
    """多仓问答走 LangGraph StateGraph(Planner→RepoAgent→Merger→Synthesiser)。"""
    from codedoc.graph.memory_backend import MemoryGraphQuery
    from codedoc.agents.qa_agent import QAAgent
    from codedoc.agents.llm import build_llm
    regs, stores, cfg0 = {}, [], None
    for r in repos:
        info = _ensure_repo_indexed(r)
        if not info:
            continue
        cfg0 = cfg0 or info["cfg"]
        gq = MemoryGraphQuery(info["cfg"], info["store"])
        regs[r] = _registry_for(gq, r, repo_root=info.get("path"))
        stores.append(info["store"])
    if len(regs) < 2:
        return None
    llm = build_llm(cfg0)
    import uuid as _uuid, concurrent.futures as _cf
    _tid = "qa-" + _uuid.uuid4().hex[:16]  # per-run thread_id:每次问答一条 checkpoint 线
    _budget = float(os.environ.get("CODEDOC_MULTI_BUDGET", "45"))  # 整 run 延迟硬预算(秒)
    _agent = QAAgent(cfg0, regs, llm)
    try:
        with _cf.ThreadPoolExecutor(max_workers=1) as _ex:
            out = _ex.submit(_agent.ask, question, list(regs), _tid).result(timeout=_budget)
    except _cf.TimeoutError:
        return None  # 超预算 → 退回单仓直连合成(/ask fallback),不让多仓拖死请求
    ver = _verify_response(out["answer"], stores[0], extra_stores=stores[1:])
    return {"answer": out["answer"], "agent_trace": out.get("trace", []),
            "selected": out.get("selected", []), "groundedness": ver["groundedness"],
            "total_refs": ver["total_refs"], "invalid_refs": ver["invalid_refs"],
            "status": "ok", "mode": "multi-agent", "thread_id": _tid}


@app.post("/api/v1/ask")
def ask(req: AskRequest, user: dict = Depends(get_current_user)):
    _repos = [r for r in (req.repos or []) if r]
    if len(_repos) > 1:
        _ma = _run_multi_agent(req.question, _repos)
        if _ma:
            return _ma
        if not req.repo:
            req.repo = _repos[0]  # 多Agent超预算/失败 -> 退回首仓有据单仓,不做无上下文直答
    info = _ensure_repo_indexed(req.repo) if req.repo else None
    if not info:
        # Fallback to direct LLM
        try:
            from codedoc.agents.llm import ChatMessage, build_llm
            from codedoc.config import load_config
            cfg = load_config(Path("/home/ubuntu/apps/codedoc"))
            llm = build_llm(cfg)
            resp = llm.chat([
                ChatMessage("system", "你是代码文档助手，简洁专业地回答用户关于代码的问题。"),
                ChatMessage("user", req.question),
            ], max_tokens=800)
            return {"answer": resp, "status": "ok"}
        except Exception as e:
            return {"answer": "", "status": f"error: {e}"}

    try:
        from codedoc.agents.llm import ChatMessage, build_llm
        from codedoc.graph.memory_backend import MemoryGraphQuery
        store = info["store"]
        cfg = info["cfg"]
        gq = MemoryGraphQuery(cfg, store)

        # GraphRAG: 全局/架构问题走 Louvain 社区摘要(map-reduce),首问算+缓存
        try:
            from codedoc import graphrag
            if graphrag.is_global_question(req.question):
                if "_comm_summaries" not in info:
                    info["_comm_summaries"] = graphrag.community_summaries(store, build_llm(cfg))
                if info.get("_comm_summaries"):
                    _ga = graphrag.answer_global(req.question, info["_comm_summaries"], build_llm(cfg))
                    _gv = _verify_response(_ga, store)
                    return {"answer": _ga, "groundedness": _gv["groundedness"],
                            "mode": "graphrag", "communities": len(info["_comm_summaries"]),
                            "status": "ok"}
        except Exception as _ge:
            print("graphrag failed:", _ge)

        # Layer 1: multi-strategy retrieval
        _smem = None
        try:
            _smem = _load_session_mem(user["sub"], req.repo)
        except Exception:
            _smem = None
        hits = _retrieve_context(store, gq, req.question, max_nodes=25, repo=req.repo, smem=_smem)
        context = ""
        allowed_set = set()
        for h in hits:
            node = h["node"]
            context += f"- {node.kind}: {node.qualified_name}"
            if getattr(node, "signature", None):
                context += f" ({node.signature})"
            if getattr(node, "docstring", None):
                context += f"\n  {node.docstring[:150]}"
            try:
                callers = gq.callers(node.id, depth=1)
            except Exception:
                callers = []
            if callers:
                context += f"\n  被调用: {', '.join(c.qualified_name for c in callers[:3])}"
            try:
                callees = gq.callees(node.id, depth=1)
            except Exception:
                callees = []
            if callees:
                context += f"\n  调用: {', '.join(c.qualified_name for c in callees[:3])}"
            context += "\n"
            qn = getattr(node, "qualified_name", "") or ""
            if qn:
                allowed_set.add(qn.split(".")[-1])
            nm = getattr(node, "name", "") or ""
            if nm:
                allowed_set.add(nm)

        allowed_symbols = ", ".join(sorted(allowed_set))

        doc_context = ""
        if "doc" in info and info["doc"]:
            doc_context = "\n\n## 已生成的代码文档（作为参考）\n" + info["doc"][:3000]

        user_doc_context = ""
        user_doc_names: list[str] = []
        try:
            _conn = _get_db()
            try:
                _rows = _conn.execute(
                    "SELECT filename, content FROM user_docs WHERE user_id=? AND repo_name=? "
                    "ORDER BY uploaded_at DESC LIMIT 3",
                    (user["sub"], req.repo),
                ).fetchall()
            finally:
                _conn.close()
            for _d in _rows:
                user_doc_names.append(_d["filename"])
                user_doc_context += f"\n\n## 用户上传文档：{_d['filename']}\n" + (_d["content"] or "")[:2000]
        except Exception:
            user_doc_context = ""

        # Layer 2: strict grounding prompt
        system_prompt = (
            f"你是 codedoc 代码知识问答助手。**严格基于以下提供的代码上下文回答**。\n\n"
            f"## 上下文（仓库 {req.repo}）\n{context}{doc_context}{user_doc_context}\n\n"
            "## 严格规则\n"
            "1. **只能引用上方上下文中明确出现的代码符号**（类名、方法名、模块名）。\n"
            "2. 引用代码符号时必须用反引号包裹：`ClassName` 或 `method_name`。\n"
            "3. **不要编造类名、方法名、文件路径**。如果上下文中没有相关信息，明确告诉用户「上下文中未找到相关代码」。\n"
            "4. 如果问题超出代码理解范畴（比如询问业务用法），简要回答即可，不要硬编代码引用。\n"
            "5. 用中文，简洁专业。\n\n"
            f"可引用的符号列表：\n{allowed_symbols}"
        )

        # 会话记忆 + 最近对话窗口(token 预算)注入(增强理解,不替代检索)
        try:
            _mc = _mem_context(_smem) if _smem else ""
            _hc = _history_window(user["sub"], req.repo)
            if _mc or _hc:
                system_prompt = system_prompt + "\n\n" + _hc + ("\n" if _hc else "") + _mc
        except Exception:
            pass

        from codedoc.llm_router import build_routed_llm
        llm = build_routed_llm(cfg, getattr(req, "model", None) or "auto")
        resp = llm.chat([
            ChatMessage("system", system_prompt),
            ChatMessage("user", req.question),
        ], max_tokens=1000)
        _model_used = getattr(llm, "last_model", None)

        # Layer 3: verify and annotate
        verify = _verify_response(resp, store)
        answer = resp or ""
        if verify["invalid_refs"]:
            answer += (
                f"\n\n*注：响应中检测到 {len(verify['invalid_refs'])} "
                f"个符号未在代码库中找到（可能为示意性引用）。*"
            )

        # REFLEXION_SINGLE 自我反思:hedge 或 groundedness 低 -> 自评缺口 -> 定向补检索 -> 重答(有界1轮)
        try:
            _RHEDGE = ("未找到", "无法回答", "没有足够",
                       "上下文中未找到", "没有相关", "无法确定")
            if any(w in (resp or "") for w in _RHEDGE) or (verify or {}).get("groundedness", 1.0) < 0.5:
                import json as _json, re as _re2
                _rsys = ("你是审查 agent，判断答案是否充分且有据。"
                         "只输出 JSON: {\"sufficient\": true/false, \"gap\": \"还缺哪个具体符号/信息\"}")
                _rr = llm.chat([ChatMessage("system", _rsys),
                                ChatMessage("user", "Q: %s\n\nA: %s" % (req.question, (resp or "")[:1500]))],
                               max_tokens=180)
                _mm = _re2.search(r"\{.*\}", _rr or "", _re2.DOTALL)
                try:
                    _rj = _json.loads(_mm.group(0)) if _mm else {}
                except Exception:
                    _rj = {}
                if not _rj.get("sufficient", True):
                    _gap = str(_rj.get("gap") or req.question)
                    _gh = _retrieve_context(store, gq, _gap, max_nodes=12, repo=req.repo, smem=_smem)
                    _gctx = ""
                    for _h in _gh:
                        _n = _h["node"]
                        _gctx += "- %s: %s" % (_n.kind, _n.qualified_name)
                        if getattr(_n, "signature", None):
                            _gctx += " (%s)" % _n.signature
                        _gctx += "\n"
                        _qn = (getattr(_n, "qualified_name", "") or "").split(".")[-1]
                        if _qn:
                            allowed_set.add(_qn)
                    if _gctx.strip():
                        _sp2 = system_prompt + "\n\n## 反思补检索(gap: %s)\n%s\n补充符号: %s" % (
                            _gap[:60], _gctx, ", ".join(sorted(allowed_set)))
                        resp = llm.chat([ChatMessage("system", _sp2), ChatMessage("user", req.question)], max_tokens=1000)
                        verify = _verify_response(resp, store)
                        answer = resp or ""
                        if verify["invalid_refs"]:
                            answer += "\n\n*反思重答后仍有 %d 个符号未找到。*" % len(verify["invalid_refs"])
        except Exception:
            pass

        # 选择性更新会话记忆(衰减 + 强化真命中 + 门控换锚点)
        try:
            if _smem is not None:
                _smem.decay(0.7)
                for _h in (hits or [])[:5]:
                    _nd = _h["node"]
                    _deg = len(gq._out.get(_nd.id, [])) + len(gq._in.get(_nd.id, []))
                    _smem.touch_focus(_nd.name, _nd.id, degree=_deg)
                if hits and len(req.question) >= 16:
                    _t0 = hits[0]["node"]
                    _smem.push_anchor(_t0.id, _t0.name)
                _save_session_mem(user["sub"], req.repo, _smem)
        except Exception:
            pass

        # Persist to qa_history (best-effort; do not fail the request).
        try:
            _conn2 = _get_db()
            try:
                _conn2.execute(
                    "INSERT INTO qa_history (user_id, repo, question, answer, groundedness, "
                    "total_refs, valid_refs_count, invalid_refs, context_nodes, user_docs) "
                    "VALUES (?,?,?,?,?,?,?,?,?,?)",
                    (
                        user["sub"], req.repo, req.question, answer,
                        verify.get("groundedness"),
                        verify.get("total_refs", 0),
                        len(verify.get("valid_refs", [])),
                        json.dumps(verify.get("invalid_refs", []), ensure_ascii=False),
                        len(hits) if hits else 0,
                        json.dumps(user_doc_names or [], ensure_ascii=False),
                    ),
                )
                _conn2.commit()
            finally:
                _conn2.close()
        except Exception as _qa_persist_err:
            import logging as _lg
            _lg.getLogger("codedoc").warning("Failed to persist QA: %s", _qa_persist_err)

        # 对话轮数超出 token 窗口阈值 → 投递异步滚动摘要作业(去重、非阻塞、失败不影响主请求)
        try:
            from codedoc.memory import summary_queue as _sq
            _cc = _get_db()
            try:
                _cnt = _cc.execute("SELECT COUNT(*) FROM qa_history WHERE user_id=? AND repo=?",
                                   (user["sub"], req.repo)).fetchone()[0]
            finally:
                _cc.close()
            if _cnt and int(_cnt) > _SUMMARY_WINDOW:
                _sq.enqueue(user["sub"], req.repo)
        except Exception:
            pass

        return {
            "answer": answer,
            "status": "ok",
            "model_used": _model_used,
            "context_nodes": len(hits),
            "user_docs": user_doc_names,
            "groundedness": round(verify["groundedness"], 3),
            "total_refs": verify["total_refs"],
            "valid_refs_count": len(verify["valid_refs"]),
            "invalid_refs": verify["invalid_refs"],
        }
    except Exception as e:
        return {"answer": "", "status": f"error: {e}"}


# ---------------------------------------------------------------------------
# Diagram generators (PlantUML + legacy Mermaid)
# ---------------------------------------------------------------------------

def _sanitize_mermaid_id(s):
    """Make safe node id (also used for PlantUML aliases)."""
    import re as _re_sm
    return _re_sm.sub(r'[^A-Za-z0-9]', '_', s)[:30]


def _plantuml_encode64(data: bytes) -> str:
    """PlantUML's custom base64-like alphabet for URL-safe fragment."""
    alphabet = "0123456789ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz-_"
    result = []
    i = 0
    while i < len(data):
        b1 = data[i]
        b2 = data[i + 1] if i + 1 < len(data) else 0
        b3 = data[i + 2] if i + 2 < len(data) else 0
        result.append(alphabet[b1 >> 2])
        result.append(alphabet[((b1 & 0x3) << 4) | (b2 >> 4)])
        result.append(alphabet[((b2 & 0xF) << 2) | (b3 >> 6)])
        result.append(alphabet[b3 & 0x3F])
        i += 3
    return ''.join(result)


def _plantuml_encode(text: str) -> str:
    """Encode PlantUML source for plantuml.com URL: deflate raw + custom base64."""
    import zlib
    # deflate raw (no zlib header / adler32 trailer)
    compressed = zlib.compress(text.encode('utf-8'), 9)[2:-4]
    return _plantuml_encode64(compressed)


def plantuml_url(text: str, fmt: str = "svg") -> str:
    """Return rendered diagram URL via plantuml.com.

    fmt: svg | png
    """
    encoded = _plantuml_encode(text)
    return f"https://www.plantuml.com/plantuml/{fmt}/{encoded}"


def _find_parent_module(store, node_id):
    """Walk up parent_id chain to find the module."""
    seen = set()
    cur = node_id
    while cur and cur not in seen:
        seen.add(cur)
        n = store.nodes.get(cur)
        if not n:
            return None
        if n["kind"] == "module":
            return cur
        cur = n.get("parent_id")
    return None


def gen_module_arch_plantuml(store, gq) -> str:
    """PlantUML component diagram of module-level imports (Apple-clean style)."""
    modules = {nid: n for nid, n in store.nodes.items() if n["kind"] == "module"}

    # Filter noise: tests, examples, demos, __pycache__
    def is_relevant(m):
        q = (m.get("qualified_name") or "").lower()
        if any(skip in q for skip in ["test", "__pycache__", "example", "demo"]):
            return False
        return True

    modules = {k: v for k, v in modules.items() if is_relevant(v)}
    if not modules:
        return "@startuml\n!theme plain\ncomponent Empty\n@enduml"

    # Build edges (imports), skip unresolved/external
    def _is_real_module(nid):
        if nid not in modules:
            return False
        n = modules[nid]
        if n.get("kind") == "unresolved":
            return False
        name = (n.get("name") or "")
        qn = (n.get("qualified_name") or "")
        if name.startswith("?") or qn.startswith("?") or qn.startswith("external"):
            return False
        return True

    edges = set()
    for e in store.edges:
        if e["kind"] != "imports":
            continue
        src_mod = _find_parent_module(store, e["src"])
        dst_mod = _find_parent_module(store, e["dst"])
        if not src_mod or not dst_mod or src_mod == dst_mod:
            continue
        if not _is_real_module(src_mod) or not _is_real_module(dst_mod):
            continue
        edges.add((src_mod, dst_mod))

    # Remove self-loops (defensive)
    edges = {(s, d) for s, d in edges if s != d}

    deg = {}
    for s, d in edges:
        deg[s] = deg.get(s, 0) + 1
        deg[d] = deg.get(d, 0) + 1

    top = sorted([m for m in modules if m in deg], key=lambda m: -deg[m])[:8]
    if not top:
        top = list(modules.keys())[:8]
    top_set = set(top)

    # Score edges by endpoint degree; keep top 12 between top nodes
    relevant_edges = [(s, d) for s, d in edges if s in top_set and d in top_set]
    relevant_edges.sort(key=lambda sd: -(deg.get(sd[0], 0) + deg.get(sd[1], 0)))
    kept_edges = relevant_edges[:12]

    # Assign unique counter-based IDs (m1, m2, ...) — avoids hash-collision label leaks
    id_map = {}
    counter = [0]

    def get_uid(m):
        if m in id_map:
            return id_map[m]
        counter[0] += 1
        uid = f"m{counter[0]}"
        id_map[m] = uid
        return uid

    # Group by top-level package
    groups = {}  # package_name -> list of (uid, label)
    for m in top:
        uid = get_uid(m)
        qn = modules[m].get("qualified_name") or modules[m].get("name") or m
        parts_q = qn.split(".")
        pkg = parts_q[0] if len(parts_q) > 1 else "root"
        label = parts_q[-1][:25]
        groups.setdefault(pkg, []).append((uid, label))

    lines = [
        "@startuml",
        "!theme plain",
        "skinparam backgroundColor #FFFFFF",
        "skinparam shadowing false",
        "skinparam roundCorner 12",
        "skinparam padding 10",
        'skinparam DefaultFontName "SF Pro Text, -apple-system, sans-serif"',
        "skinparam DefaultFontColor #1D1D1F",
        "skinparam ArrowColor #86868B",
        "skinparam ArrowThickness 1.2",
        "skinparam linetype ortho",
        "skinparam component {",
        "  BackgroundColor #FFFFFF",
        "  BorderColor #0071E3",
        "  FontColor #1D1D1F",
        "}",
        "skinparam package {",
        "  BackgroundColor #FAFAFA",
        "  BorderColor #D2D2D7",
        "  FontColor #86868B",
        "  Style rectangle",
        "}",
        "left to right direction",
        "",
        "title 系统架构图",
    ]

    for pkg, items in groups.items():
        if len(items) == 1 or pkg == "root":
            for uid, label in items:
                lines.append(f'component "{label}" as {uid}')
        else:
            lines.append(f'package "{pkg}" {{')
            for uid, label in items:
                lines.append(f'  component "{label}" as {uid}')
            lines.append("}")

    for s, d in kept_edges:
        lines.append(f"{id_map[s]} --> {id_map[d]}")
    lines.append("@enduml")
    return "\n".join(lines)


def gen_class_uml_plantuml(store, gq) -> str:
    """PlantUML class diagram for top classes (Apple-clean style)."""
    classes = {nid: n for nid, n in store.nodes.items() if n["kind"] == "class"}

    # Filter noise: test classes, mocks, fixtures
    def is_relevant(c):
        q = (c.get("qualified_name") or "").lower()
        if any(skip in q for skip in ["test", "mock", "fixture"]):
            return False
        return True

    classes = {k: v for k, v in classes.items() if is_relevant(v)}
    if not classes:
        return "@startuml\n!theme plain\nclass 暂无可显示类\n@enduml"

    degree = {nid: 0 for nid in classes}
    for e in store.edges:
        if e["src"] in degree:
            degree[e["src"]] += 1
        if e["dst"] in degree:
            degree[e["dst"]] += 1

    top = sorted(classes.keys(), key=lambda c: -degree[c])[:5]
    top_set = set(top)

    lines = [
        "@startuml",
        "!theme plain",
        "skinparam backgroundColor #FFFFFF",
        "skinparam shadowing false",
        "skinparam roundCorner 8",
        "skinparam padding 10",
        'skinparam DefaultFontName "SF Pro Text, -apple-system, sans-serif"',
        "skinparam DefaultFontColor #1D1D1F",
        "skinparam ArrowColor #86868B",
        "skinparam ArrowThickness 1.2",
        "skinparam classAttributeIconSize 0",
        "skinparam class {",
        "  BackgroundColor #FFFFFF",
        "  BorderColor #0071E3",
        "  HeaderBackgroundColor #F5F5F7",
        "  FontColor #1D1D1F",
        "  AttributeFontColor #424245",
        "}",
        "hide circle",
        "",
        "title 核心类 UML",
    ]
    for cid in top:
        c = classes[cid]
        name = _sanitize_mermaid_id(c["name"])
        methods_all, fields_all = [], []
        for e in store.edges:
            if e["src"] == cid and e["kind"] == "contains":
                child = store.nodes.get(e["dst"])
                if not child:
                    continue
                if child["kind"] in ("method", "function"):
                    methods_all.append(child["name"])
                elif child["kind"] == "field":
                    fields_all.append(child["name"])

        non_dunder = sorted({m for m in methods_all if not m.startswith("__")})
        dunder = sorted({m for m in methods_all if m.startswith("__")})
        methods = non_dunder[:5] if non_dunder else dunder[:3]
        fields = sorted(set(fields_all))[:3]

        lines.append(f"class {name} {{")
        for f in fields:
            lines.append(f"  {_sanitize_mermaid_id(f)[:20]}")
        if fields and methods:
            lines.append("  --")
        for m in methods:
            lines.append(f"  {_sanitize_mermaid_id(m)[:20]}()")
        lines.append("}")

    for e in store.edges:
        if e["src"] in top_set and e["dst"] in top_set:
            if e["kind"] == "extends":
                child = _sanitize_mermaid_id(classes[e["src"]]["name"])
                parent = _sanitize_mermaid_id(classes[e["dst"]]["name"])
                lines.append(f"{parent} <|-- {child}")
            elif e["kind"] == "implements":
                child = _sanitize_mermaid_id(classes[e["src"]]["name"])
                parent = _sanitize_mermaid_id(classes[e["dst"]]["name"])
                lines.append(f"{parent} <|.. {child}")
    lines.append("@enduml")
    return "\n".join(lines)


def gen_call_flow_plantuml(store, gq) -> str:
    """PlantUML flow diagram for call chains (Apple-clean style)."""

    def is_external(node):
        if not node:
            return True
        if node.get("kind") == "unresolved":
            return True
        nm = node.get("name") or ""
        qn = node.get("qualified_name") or ""
        if nm.startswith("?") or qn.startswith("?"):
            return True
        if nm in {"print", "len", "str", "int", "list", "dict", "input", "open", "range",
                  "tuple", "set", "bool", "float", "type", "isinstance", "hasattr",
                  "getattr", "setattr", "repr", "format"}:
            return True
        return False

    routes = [n for n in store.nodes.values() if n["kind"] == "route_handler"]
    if not routes:
        call_count = {}
        for e in store.edges:
            if e["kind"] == "calls":
                call_count[e["src"]] = call_count.get(e["src"], 0) + 1
        methods = sorted(
            [n for n in store.nodes.values()
             if n["kind"] in ("method", "function") and not is_external(n)],
            key=lambda m: -call_count.get(m["id"], 0),
        )[:3]
        if not methods:
            return "@startuml\n!theme plain\nrectangle \"暂无调用关系\" as A\n@enduml"
        seeds = methods
    else:
        seeds = routes[:3]

    lines = [
        "@startuml",
        "!theme plain",
        "skinparam backgroundColor #FFFFFF",
        "skinparam shadowing false",
        "skinparam roundCorner 10",
        "skinparam padding 10",
        'skinparam DefaultFontName "SF Pro Text, -apple-system, sans-serif"',
        "skinparam DefaultFontColor #1D1D1F",
        "skinparam ArrowColor #86868B",
        "skinparam ArrowThickness 1.2",
        "skinparam linetype polyline",
        "left to right direction",
        "skinparam rectangle {",
        "  BackgroundColor #FFFFFF",
        "  BorderColor #0071E3",
        "  FontColor #1D1D1F",
        "  roundCorner 12",
        "}",
        "",
        "title 调用关系图",
    ]
    NODE_CAP = 12
    id_map = {}  # original node id -> unique sanitized id (n1, n2, ...)
    drawn_edges = set()
    counter = [0]

    def get_or_add(node):
        nid = node["id"]
        if nid in id_map:
            return id_map[nid]
        if len(id_map) >= NODE_CAP:
            return None
        counter[0] += 1
        uid = f"n{counter[0]}"
        id_map[nid] = uid
        label = (node["name"] or "")[:22].replace('"', "'")
        kind_hint = {
            "route_handler": "<<route>>",
            "method": "<<method>>",
            "function": "<<func>>",
            "class": "<<class>>",
        }.get(node.get("kind", ""), "")
        if kind_hint:
            lines.append(f'rectangle "{label}\\n{kind_hint}" as {uid}')
        else:
            lines.append(f'rectangle "{label}" as {uid}')
        return uid

    def expand(node, depth):
        if depth <= 0 or len(id_map) >= NODE_CAP:
            return
        src_uid = get_or_add(node)
        if not src_uid:
            return
        for e in store.edges:
            if e["src"] != node["id"]:
                continue
            if e["kind"] not in ("calls", "route_handler"):
                continue
            target = store.nodes.get(e["dst"])
            if not target or is_external(target):
                continue
            tgt_uid = get_or_add(target)
            if not tgt_uid:
                continue
            if src_uid == tgt_uid:
                continue
            key = (src_uid, tgt_uid)
            if key in drawn_edges:
                continue
            drawn_edges.add(key)
            lines.append(f"{src_uid} --> {tgt_uid}")
            expand(target, depth - 1)

    for seed in seeds:
        expand(seed, 2)
    lines.append("@enduml")
    return "\n".join(lines)


# ---------------------------------------------------------------------------
# Doc templates
# ---------------------------------------------------------------------------

# ---------------------------------------------------------------------------
# Deterministic extractors for deliverable-quality docs (Confluence-ready).
# These pull max value from the parsed graph + raw source files in repo_dir.
# ---------------------------------------------------------------------------

_SKIP_PATH_PARTS = ("__pycache__", ".venv", "node_modules", ".git", ".tox", "dist", "build", ".mypy_cache", ".pytest_cache")


def _iter_py_files(repo_dir: Path, cap: int = 400):
    """Yield up to cap .py files, excluding vendored / generated paths."""
    count = 0
    for py in repo_dir.rglob("*.py"):
        s = str(py)
        if any(p in s for p in _SKIP_PATH_PARTS):
            continue
        yield py
        count += 1
        if count >= cap:
            return


def _read_source_slice(repo_dir: Path, rel_file: str, start_line: int | None, end_line: int | None) -> str:
    """Return the source text between [start_line, end_line] (1-indexed inclusive)."""
    if not rel_file:
        return ""
    fpath = repo_dir / rel_file
    if not fpath.exists() or not fpath.is_file():
        return ""
    try:
        text = fpath.read_text(encoding="utf-8", errors="replace")
    except OSError:
        return ""
    lines = text.splitlines()
    s = max(0, (start_line or 1) - 1)
    e = end_line or (s + 1)
    return "\n".join(lines[s:e])


def extract_entry_points(repo_dir: Path, store) -> list[dict]:
    """Find entry points: __main__ blocks, console_scripts, route handlers, click commands."""
    entries: list[dict] = []

    # 1) pyproject.toml [project.scripts] / [tool.poetry.scripts]
    pyproject = repo_dir / "pyproject.toml"
    if pyproject.exists():
        try:
            content = pyproject.read_text(encoding="utf-8", errors="replace")
            for header in (r"\[project\.scripts\]", r"\[tool\.poetry\.scripts\]"):
                m = re.search(header + r"\s*\n(.*?)(?=\n\[|\Z)", content, re.DOTALL)
                if not m:
                    continue
                for line in m.group(1).strip().splitlines():
                    line = line.strip()
                    if not line or line.startswith("#") or "=" not in line:
                        continue
                    name, target = line.split("=", 1)
                    entries.append({
                        "type": "console_script",
                        "name": name.strip().strip('"\''),
                        "target": target.strip().strip('"\''),
                    })
        except OSError:
            pass

    # 2) setup.py entry_points
    setup_py = repo_dir / "setup.py"
    if setup_py.exists():
        try:
            content = setup_py.read_text(encoding="utf-8", errors="replace")
            m = re.search(r"console_scripts\s*['\"]?\s*:\s*\[(.*?)\]", content, re.DOTALL)
            if m:
                for raw in re.findall(r"['\"]([^'\"]+=[^'\"]+)['\"]", m.group(1)):
                    if "=" in raw:
                        nm, tgt = raw.split("=", 1)
                        entries.append({
                            "type": "console_script",
                            "name": nm.strip(),
                            "target": tgt.strip(),
                        })
        except OSError:
            pass

    # 3) __main__ blocks in source files
    main_files: list[str] = []
    for py_file in _iter_py_files(repo_dir, cap=400):
        try:
            src = py_file.read_text(encoding="utf-8", errors="replace")
        except OSError:
            continue
        if 'if __name__ == "__main__"' in src or "if __name__ == '__main__'" in src:
            try:
                rel = str(py_file.relative_to(repo_dir)).replace("\\", "/")
            except ValueError:
                rel = py_file.name
            main_files.append(rel)
    main_files.sort(key=lambda p: (p.count("/"), len(p)))
    for rel in main_files[:10]:
        entries.append({"type": "__main__", "file": rel})

    # 4) From graph: route_handler + click commands (detected via decorators in signature)
    route_count = 0
    cli_count = 0
    for n in store.nodes.values():
        if n.get("kind") == "route_handler" and route_count < 8:
            entries.append({
                "type": "route",
                "name": n.get("name", ""),
                "qualified_name": n.get("qualified_name", ""),
                "file": n.get("file", ""),
            })
            route_count += 1
        sig = (n.get("signature") or "") + " " + (n.get("docstring") or "")
        if cli_count < 8 and n.get("kind") in ("function", "method"):
            if "@click.command" in sig or "@click.group" in sig or "@cli.command" in sig:
                entries.append({
                    "type": "cli_command",
                    "name": n.get("name", ""),
                    "qualified_name": n.get("qualified_name", ""),
                    "file": n.get("file", ""),
                })
                cli_count += 1

    return entries[:25]


def _node_calls(store, src_id):
    """Yield (target_node) for each 'calls' edge from src_id, skipping unresolved."""
    for e in store.edges:
        if e.get("src") != src_id or e.get("kind") != "calls":
            continue
        tgt = store.nodes.get(e.get("dst"))
        if not tgt:
            continue
        if tgt.get("kind") == "unresolved":
            continue
        nm = tgt.get("name") or ""
        if not nm or nm.startswith("?"):
            continue
        yield tgt


def gen_sequence_diagram_plantuml(store, seed_node, depth: int = 3, max_calls: int = 14) -> str:
    """Build an Apple-clean PlantUML sequence diagram following call chain from seed_node."""
    if not seed_node:
        return ""
    header = [
        "@startuml",
        "!theme plain",
        "skinparam backgroundColor #FFFFFF",
        "skinparam shadowing false",
        "skinparam roundCorner 8",
        'skinparam DefaultFontName "SF Pro Text, -apple-system, sans-serif"',
        "skinparam DefaultFontColor #1D1D1F",
        "skinparam ArrowColor #86868B",
        "skinparam ArrowThickness 1.2",
        "skinparam sequence {",
        "  ParticipantBorderColor #0071E3",
        "  ParticipantBackgroundColor #FFFFFF",
        "  ParticipantFontColor #1D1D1F",
        "  LifeLineBorderColor #D2D2D7",
        "  ActorBorderColor #0071E3",
        "  ActorBackgroundColor #FFFFFF",
        "}",
        "",
        f"title 入口调用时序: {seed_node.get('name', '')[:40]}",
        "actor User",
    ]
    body: list[str] = []
    participants: dict[str, str] = {}
    counter = [0]

    def get_participant(node) -> str:
        nid = node["id"]
        if nid in participants:
            return participants[nid]
        counter[0] += 1
        pid = f"p{counter[0]}"
        participants[nid] = pid
        name = (node.get("name") or "?")[:22].replace('"', "'")
        header.append(f'participant "{name}" as {pid}')
        return pid

    drawn: set[tuple[str, str]] = set()
    call_count = [0]
    visited_in_chain: set[str] = set()

    def expand(node, current_depth: int):
        if current_depth <= 0 or call_count[0] >= max_calls:
            return
        if node["id"] in visited_in_chain:
            return
        visited_in_chain.add(node["id"])
        src_pid = get_participant(node)
        for target in _node_calls(store, node["id"]):
            tgt_pid = get_participant(target)
            key = (src_pid, tgt_pid)
            if key in drawn:
                continue
            drawn.add(key)
            label = (target.get("name") or "?")[:18]
            body.append(f"{src_pid} -> {tgt_pid}: {label}")
            call_count[0] += 1
            if call_count[0] >= max_calls:
                return
            expand(target, current_depth - 1)

    root_pid = get_participant(seed_node)
    body.insert(0, f"User -> {root_pid}: invoke")
    expand(seed_node, depth)
    if not any(line for line in body if "->" in line and not line.startswith("User")):
        body.append(f"{root_pid} --> User: (无下游调用)")
    else:
        body.append(f"{root_pid} --> User: result")

    return "\n".join(header + [""] + body + ["@enduml"])


def _scan_raises(repo_dir: Path, node) -> list[str]:
    body = _read_source_slice(repo_dir, node.get("file", ""), node.get("start_line"), node.get("end_line"))
    if not body:
        return []
    raises = re.findall(r'\braise\s+([A-Z][A-Za-z0-9_]+)', body)
    return sorted(set(raises))[:8]


def _scan_catches(repo_dir: Path, node) -> list[str]:
    body = _read_source_slice(repo_dir, node.get("file", ""), node.get("start_line"), node.get("end_line"))
    if not body:
        return []
    catches: list[str] = []
    for m in re.findall(r'\bexcept\s+([A-Z][A-Za-z0-9_.]+|\([^)]+\))', body):
        if m.startswith("("):
            for sub in re.findall(r'[A-Z][A-Za-z0-9_.]+', m):
                catches.append(sub)
        else:
            catches.append(m)
    return sorted(set(catches))[:8]


def _scan_self_writes(repo_dir: Path, node) -> list[str]:
    body = _read_source_slice(repo_dir, node.get("file", ""), node.get("start_line"), node.get("end_line"))
    if not body:
        return []
    writes = re.findall(r'\bself\.([A-Za-z_][A-Za-z0-9_]*)\s*=', body)
    return sorted(set(writes))[:6]


def extract_api_contracts(repo_dir: Path, store, top_class_ids: list[str]) -> list[dict]:
    """For each class id, extract its public methods + signatures + raise/catch info."""
    contracts: list[dict] = []
    for cid in top_class_ids:
        c = store.nodes.get(cid)
        if not c:
            continue
        methods: list[dict] = []
        for e in store.edges:
            if e.get("src") != cid or e.get("kind") != "contains":
                continue
            m = store.nodes.get(e.get("dst"))
            if not m or m.get("kind") not in ("method", "function"):
                continue
            name = m.get("name", "")
            if not name or name.startswith("_"):
                continue
            sig = m.get("signature") or ""
            doc = (m.get("docstring") or "").strip()
            if len(doc) > 240:
                doc = doc[:240].rstrip() + "..."
            methods.append({
                "name": name,
                "signature": sig,
                "docstring": doc,
                "file": m.get("file", ""),
                "raises": _scan_raises(repo_dir, m),
                "catches": _scan_catches(repo_dir, m),
                "writes_self": _scan_self_writes(repo_dir, m),
            })
        methods.sort(key=lambda x: x["name"])
        contracts.append({
            "class": c.get("qualified_name") or c.get("name", ""),
            "docstring": ((c.get("docstring") or "").strip()[:200]),
            "file": c.get("file", ""),
            "methods": methods[:8],
        })
    return contracts


def extract_exception_hierarchy(store) -> dict:
    """Find custom exception classes via name heuristics + 'extends' edges."""
    exceptions: list[dict] = []
    name_set: set[str] = set()
    for nid, n in store.nodes.items():
        if n.get("kind") != "class":
            continue
        nm = n.get("name", "") or ""
        if not any(nm.endswith(s) for s in ("Error", "Exception", "Warning", "Fail", "Invalid")):
            continue
        parents: list[str] = []
        for e in store.edges:
            if e.get("src") == nid and e.get("kind") == "extends":
                p = store.nodes.get(e.get("dst"))
                if p:
                    parents.append(p.get("name") or "?")
        exceptions.append({
            "name": nm,
            "qualified_name": n.get("qualified_name", ""),
            "parents": parents,
            "file": n.get("file", ""),
        })
        name_set.add(nm)
    exceptions.sort(key=lambda x: x["name"])
    return {"count": len(exceptions), "list": exceptions, "name_set": name_set}


def gen_exception_uml_plantuml(exceptions: list[dict]) -> str:
    """PlantUML class diagram of the custom exception hierarchy."""
    if not exceptions:
        return ""
    lines = [
        "@startuml",
        "!theme plain",
        "skinparam backgroundColor #FFFFFF",
        "skinparam shadowing false",
        "skinparam roundCorner 8",
        'skinparam DefaultFontName "SF Pro Text, -apple-system, sans-serif"',
        "skinparam DefaultFontColor #1D1D1F",
        "skinparam ArrowColor #86868B",
        "skinparam class {",
        "  BackgroundColor #FFFFFF",
        "  BorderColor #D70015",
        "  HeaderBackgroundColor #FFF5F5",
        "  FontColor #1D1D1F",
        "}",
        "hide circle",
        "hide members",
        "",
        "title 异常类继承关系",
    ]
    drawn_classes: set[str] = set()
    for exc in exceptions[:18]:
        nm = _sanitize_mermaid_id(exc["name"])
        if nm in drawn_classes:
            continue
        drawn_classes.add(nm)
        lines.append(f'class {nm}')
        for p in exc["parents"]:
            pnm = _sanitize_mermaid_id(p)
            if pnm not in drawn_classes:
                lines.append(f'class {pnm}')
                drawn_classes.add(pnm)
            lines.append(f"{pnm} <|-- {nm}")
    lines.append("@enduml")
    return "\n".join(lines)


def count_raise_catch_sites(repo_dir: Path, exception_names: set[str]) -> dict:
    """Tally raise X / except X sites across the repo for known exception names."""
    raise_counts: dict[str, int] = {n: 0 for n in exception_names}
    catch_counts: dict[str, int] = {n: 0 for n in exception_names}
    if not exception_names:
        return {"raise": raise_counts, "catch": catch_counts}
    raise_re = re.compile(r'\braise\s+([A-Z][A-Za-z0-9_]+)')
    except_re = re.compile(r'\bexcept\s+([A-Z][A-Za-z0-9_.]+|\([^)]+\))')
    for py_file in _iter_py_files(repo_dir, cap=400):
        try:
            src = py_file.read_text(encoding="utf-8", errors="replace")
        except OSError:
            continue
        for m in raise_re.finditer(src):
            nm = m.group(1)
            if nm in raise_counts:
                raise_counts[nm] += 1
        for m in except_re.finditer(src):
            grp = m.group(1)
            if grp.startswith("("):
                for sub in re.findall(r'[A-Z][A-Za-z0-9_.]+', grp):
                    short = sub.rsplit(".", 1)[-1]
                    if short in catch_counts:
                        catch_counts[short] += 1
            else:
                short = grp.rsplit(".", 1)[-1]
                if short in catch_counts:
                    catch_counts[short] += 1
    return {"raise": raise_counts, "catch": catch_counts}


def extract_config(repo_dir: Path, store) -> dict:
    """Scan source for env vars + locate Settings/Config classes."""
    env_vars: dict[str, str | None] = {}  # name -> default value (or None)
    config_files: set[str] = set()
    env_re_get = re.compile(r'os\.environ\.get\(\s*["\']([A-Z_][A-Z0-9_]*)["\']\s*(?:,\s*([^)]+))?\)')
    env_re_getenv = re.compile(r'os\.getenv\(\s*["\']([A-Z_][A-Z0-9_]*)["\']\s*(?:,\s*([^)]+))?\)')
    env_re_idx = re.compile(r'os\.environ\[\s*["\']([A-Z_][A-Z0-9_]*)["\']\s*\]')
    yaml_re = re.compile(r'[\'"]([^\'"\s]+\.ya?ml)[\'"]')
    json_cfg_re = re.compile(r'[\'"]([^\'"\s]*(?:config|settings)[^\'"\s]*\.json)[\'"]', re.IGNORECASE)
    constants: list[tuple[str, str, str]] = []  # (module, name, value)
    const_re = re.compile(r'^([A-Z][A-Z0-9_]{2,})\s*=\s*(.+?)$', re.MULTILINE)

    for py_file in _iter_py_files(repo_dir, cap=400):
        try:
            src = py_file.read_text(encoding="utf-8", errors="replace")
        except OSError:
            continue
        for m in env_re_get.finditer(src):
            name, default = m.group(1), (m.group(2) or "").strip()
            if name not in env_vars:
                env_vars[name] = default[:60] if default else None
        for m in env_re_getenv.finditer(src):
            name, default = m.group(1), (m.group(2) or "").strip()
            if name not in env_vars:
                env_vars[name] = default[:60] if default else None
        for m in env_re_idx.finditer(src):
            name = m.group(1)
            if name not in env_vars:
                env_vars[name] = None
        for m in yaml_re.finditer(src):
            config_files.add(m.group(1))
        for m in json_cfg_re.finditer(src):
            config_files.add(m.group(1))
        # Module-top-level constants — quick + dirty
        try:
            rel = str(py_file.relative_to(repo_dir)).replace("\\", "/")
        except ValueError:
            rel = py_file.name
        if "/__init__.py" in rel or rel.endswith("settings.py") or rel.endswith("config.py") or rel.endswith("constants.py"):
            for m in const_re.finditer(src):
                nm, val = m.group(1), m.group(2).strip()
                if nm in {"True", "False", "None"}:
                    continue
                if len(val) > 80:
                    val = val[:80] + "..."
                constants.append((rel, nm, val))

    config_classes: list[dict] = []
    for n in store.nodes.values():
        if n.get("kind") != "class":
            continue
        nm = n.get("name", "") or ""
        if any(k in nm for k in ("Settings", "Config", "Configuration", "BaseSettings")):
            config_classes.append({
                "name": nm,
                "qualified_name": n.get("qualified_name", ""),
                "file": n.get("file", ""),
            })

    return {
        "env_vars": sorted(env_vars.items()),
        "config_files": sorted(config_files)[:15],
        "config_classes": config_classes[:10],
        "constants": constants[:25],
    }


_DEP_CATEGORIES = {
    "web": {"fastapi", "flask", "django", "starlette", "uvicorn", "aiohttp", "tornado", "sanic", "bottle", "quart", "werkzeug", "jinja2"},
    "db": {"sqlalchemy", "psycopg2", "psycopg", "pymongo", "redis", "mysqlclient", "aiopg", "asyncpg", "peewee", "alembic", "sqlmodel"},
    "test": {"pytest", "unittest", "tox", "nox", "coverage", "mock", "hypothesis", "pytest-cov", "pytest-asyncio"},
    "cli": {"click", "typer", "argparse", "fire", "docopt", "rich"},
    "data": {"pandas", "numpy", "scipy", "pyarrow", "polars", "dask"},
    "ml": {"torch", "tensorflow", "scikit-learn", "transformers", "sklearn", "keras"},
    "http": {"requests", "httpx", "urllib3", "aiohttp"},
    "auth": {"jwt", "pyjwt", "bcrypt", "passlib", "cryptography", "authlib"},
    "config": {"pydantic", "pydantic-settings", "python-dotenv", "dynaconf", "omegaconf", "hydra-core"},
    "async": {"asyncio", "anyio", "trio", "celery", "kombu"},
    "serialize": {"pyyaml", "toml", "msgpack", "orjson", "marshmallow"},
    "lint": {"ruff", "black", "mypy", "flake8", "pylint", "isort", "pre-commit"},
}


def _categorize_dep(name: str) -> str:
    low = name.lower().replace("_", "-")
    for cat, members in _DEP_CATEGORIES.items():
        if low in members or low.split("-")[0] in members:
            return cat
    return "other"


def extract_dependencies(repo_dir: Path, store) -> dict:
    """Pull dependencies from pyproject.toml / requirements.txt / setup.py."""
    deps: dict[str, str] = {}  # name -> source
    extras: dict[str, list[str]] = {}

    pyproject = repo_dir / "pyproject.toml"
    if pyproject.exists():
        try:
            content = pyproject.read_text(encoding="utf-8", errors="replace")
        except OSError:
            content = ""
        # PEP 621 dependencies = [...]
        m = re.search(r'^dependencies\s*=\s*\[(.*?)\]', content, re.MULTILINE | re.DOTALL)
        if m:
            for raw in re.findall(r'["\']([^"\']+)["\']', m.group(1)):
                name = re.split(r'[<>=!~;\[]', raw, 1)[0].strip()
                if name:
                    deps.setdefault(name, "pyproject.toml")
        # poetry [tool.poetry.dependencies]
        m = re.search(r'\[tool\.poetry\.dependencies\]\s*\n(.*?)(?=\n\[|\Z)', content, re.DOTALL)
        if m:
            for line in m.group(1).strip().splitlines():
                if "=" in line and not line.strip().startswith("#"):
                    nm = line.split("=", 1)[0].strip().strip('"\'')
                    if nm and nm.lower() != "python":
                        deps.setdefault(nm, "pyproject.toml")
        # [project.optional-dependencies] — each subkey is a list
        m = re.search(r'\[project\.optional-dependencies\]\s*\n(.*?)(?=\n\[|\Z)', content, re.DOTALL)
        if m:
            for arr_m in re.finditer(r'^([a-zA-Z0-9_-]+)\s*=\s*\[(.*?)\]', m.group(1), re.MULTILINE | re.DOTALL):
                group_name = arr_m.group(1)
                for raw in re.findall(r'["\']([^"\']+)["\']', arr_m.group(2)):
                    name = re.split(r'[<>=!~;\[]', raw, 1)[0].strip()
                    if name:
                        deps.setdefault(name, f"pyproject.toml [optional:{group_name}]")
        # PEP 735 [dependency-groups] (used by uv / pip's group manager)
        m = re.search(r'\[dependency-groups\]\s*\n(.*?)(?=\n\[|\Z)', content, re.DOTALL)
        if m:
            for arr_m in re.finditer(r'^([a-zA-Z0-9_-]+)\s*=\s*\[(.*?)\]', m.group(1), re.MULTILINE | re.DOTALL):
                group_name = arr_m.group(1)
                for raw in re.findall(r'["\']([^"\']+)["\']', arr_m.group(2)):
                    name = re.split(r'[<>=!~;\[]', raw, 1)[0].strip()
                    if name:
                        deps.setdefault(name, f"pyproject.toml [group:{group_name}]")

    req_files = ["requirements.txt", "requirements-dev.txt", "requirements/base.txt"]
    for rf in req_files:
        p = repo_dir / rf
        if not p.exists():
            continue
        try:
            for line in p.read_text(encoding="utf-8", errors="replace").splitlines():
                line = line.strip()
                if not line or line.startswith("#") or line.startswith("-"):
                    continue
                name = re.split(r'[<>=!~;\[]', line, 1)[0].strip()
                if name:
                    deps.setdefault(name, rf)
        except OSError:
            pass

    setup_py = repo_dir / "setup.py"
    if setup_py.exists() and not deps:
        try:
            content = setup_py.read_text(encoding="utf-8", errors="replace")
            m = re.search(r'install_requires\s*=\s*\[(.*?)\]', content, re.DOTALL)
            if m:
                for raw in re.findall(r'["\']([^"\']+)["\']', m.group(1)):
                    name = re.split(r'[<>=!~;\[]', raw, 1)[0].strip()
                    if name:
                        deps.setdefault(name, "setup.py")
        except OSError:
            pass

    categorized: dict[str, list[str]] = {}
    for name, src in sorted(deps.items()):
        cat = _categorize_dep(name)
        categorized.setdefault(cat, []).append(name)

    return {"deps": sorted(deps.items()), "categorized": categorized, "count": len(deps)}


def extract_test_coverage(store) -> dict:
    """Collect test files & functions, group by module, infer human description from name."""
    def is_test_file(rel: str) -> bool:
        if not rel:
            return False
        low = rel.lower().replace("\\", "/")
        return (
            low.startswith("test")
            or low.startswith("tests/")
            or "/tests/" in low
            or "/test/" in low
            or "_test.py" in low
            or "test_" in low.rsplit("/", 1)[-1]
        )

    test_funcs: list[dict] = []
    tested_class_refs: set[str] = set()
    for n in store.nodes.values():
        rel = n.get("file", "") or ""
        if not is_test_file(rel):
            continue
        if n.get("kind") not in ("method", "function"):
            continue
        name = n.get("name", "") or ""
        if not name.startswith("test_") and not name.startswith("test"):
            continue
        # Build a humanized description
        rest = name[5:] if name.startswith("test_") else name[4:]
        desc = rest.replace("_", " ").strip().capitalize()
        test_funcs.append({
            "name": name,
            "file": rel,
            "description": desc or name,
            "qualified_name": n.get("qualified_name", ""),
        })

    # Group by module file
    by_module: dict[str, list[dict]] = {}
    for t in test_funcs:
        by_module.setdefault(t["file"], []).append(t)

    # Identify untested classes (rough heuristic: class name not mentioned in any test
    # function's qualified_name or test source file path)
    all_classes = [n for n in store.nodes.values() if n.get("kind") == "class"]
    test_blob = " ".join(t["qualified_name"] + " " + t["name"] for t in test_funcs).lower()
    untested: list[dict] = []
    for c in all_classes:
        nm = c.get("name") or ""
        rel = c.get("file", "") or ""
        if is_test_file(rel):
            continue
        if not nm or nm.startswith("_"):
            continue
        if nm.lower() not in test_blob:
            untested.append({"name": nm, "qualified_name": c.get("qualified_name", ""), "file": rel})

    return {
        "total_tests": len(test_funcs),
        "module_count": len(by_module),
        "by_module": by_module,
        "untested_classes": untested[:25],
        "untested_total": len(untested),
        "class_total": len(all_classes),
    }


def extract_project_metadata(repo_dir: Path) -> dict:
    """README intro + pyproject info + LICENSE detection + CI workflow names + Dockerfile/Makefile."""
    meta: dict = {}

    # README first paragraph
    for cand in ["README.md", "README.rst", "README", "readme.md", "Readme.md"]:
        f = repo_dir / cand
        if not f.exists():
            continue
        try:
            content = f.read_text(encoding="utf-8", errors="replace")[:8000]
        except OSError:
            continue
        meta["readme_file"] = cand
        # Find the first non-trivial paragraph (skip title/badges)
        paragraph: list[str] = []
        for line in content.splitlines():
            stripped = line.strip()
            if not stripped:
                if paragraph:
                    break
                continue
            if stripped.startswith("#"):
                if paragraph:
                    break
                continue
            if set(stripped) <= {"=", "-", "~", "*", "_"}:
                # underlines / horizontal rules
                if paragraph:
                    break
                continue
            if stripped.startswith("[!["):
                # badges
                continue
            if stripped.startswith("<") and ">" in stripped and not paragraph:
                # HTML wrapper lines
                continue
            paragraph.append(stripped)
            if sum(len(p) for p in paragraph) > 600:
                break
        intro = " ".join(paragraph).strip()
        # strip simple markdown links to text
        intro = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', intro)
        meta["readme_intro"] = intro[:700]
        break

    # pyproject.toml
    pp = repo_dir / "pyproject.toml"
    if pp.exists():
        try:
            content = pp.read_text(encoding="utf-8", errors="replace")
            for key in ["name", "version", "description", "requires-python", "license", "homepage", "repository"]:
                m = re.search(rf'^{key}\s*=\s*["\']([^"\']+)["\']', content, re.MULTILINE)
                if m:
                    meta[f"pyproject_{key.replace('-', '_')}"] = m.group(1)
            # keywords / authors
            m = re.search(r'^keywords\s*=\s*\[(.*?)\]', content, re.MULTILINE | re.DOTALL)
            if m:
                kws = re.findall(r'["\']([^"\']+)["\']', m.group(1))
                if kws:
                    meta["pyproject_keywords"] = kws[:10]
            m = re.search(r'^authors\s*=\s*\[(.*?)\]', content, re.MULTILINE | re.DOTALL)
            if m:
                authors = re.findall(r'name\s*=\s*["\']([^"\']+)["\']', m.group(1))
                if authors:
                    meta["pyproject_authors"] = authors[:5]
        except OSError:
            pass

    # LICENSE
    for cand in ["LICENSE", "LICENSE.txt", "LICENSE.md", "COPYING", "LICENCE"]:
        f = repo_dir / cand
        if not f.exists():
            continue
        try:
            content = f.read_text(encoding="utf-8", errors="replace")[:400]
            for lic in ["MIT", "Apache License", "BSD", "GPL", "LGPL", "MPL", "ISC"]:
                if lic in content:
                    meta["license_file"] = cand
                    meta["license_type"] = "Apache 2.0" if lic == "Apache License" else lic
                    break
        except OSError:
            pass
        if "license_type" in meta:
            break

    # Dockerfile / docker-compose
    for cand in ["Dockerfile", "Dockerfile.dev", "docker-compose.yml", "docker-compose.yaml"]:
        if (repo_dir / cand).exists():
            meta.setdefault("docker_files", []).append(cand)

    # Makefile targets
    mk = repo_dir / "Makefile"
    if mk.exists():
        try:
            content = mk.read_text(encoding="utf-8", errors="replace")
            targets = re.findall(r'^([a-zA-Z][a-zA-Z0-9_\-]*):', content, re.MULTILINE)
            meta["makefile_targets"] = sorted(set(t for t in targets if not t.startswith(".")))[:15]
        except OSError:
            pass

    # GitHub Actions workflows
    wf_dir = repo_dir / ".github" / "workflows"
    if wf_dir.exists():
        workflows: list[str] = []
        try:
            for yml in wf_dir.iterdir():
                if not yml.is_file() or yml.suffix not in (".yml", ".yaml"):
                    continue
                try:
                    content = yml.read_text(encoding="utf-8", errors="replace")[:600]
                    m = re.search(r'^name:\s*(.+)$', content, re.MULTILINE)
                    nm = m.group(1).strip().strip('"\'') if m else yml.stem
                    workflows.append(nm)
                except OSError:
                    workflows.append(yml.stem)
        except OSError:
            pass
        meta["ci_workflows"] = workflows[:15]

    return meta



def extract_project_background(repo_dir: Path, store) -> dict:
    """Mine project background from README, docs/, package docstrings, top class docs,
    examples/, CHANGELOG, pyproject classifiers/urls, and GitHub workflows."""
    bg: dict = {
        "readme_full": "",
        "readme_features": [],
        "readme_examples": [],
        "docs_files": [],
        "package_docstring": "",
        "top_class_docs": [],
        "examples_dir": [],
        "changelog_recent": [],
        "classifiers": [],
        "urls": {},
        "workflows": [],
    }

    # README — full content + extracted features + code blocks
    for cand in ["README.md", "README.rst", "README", "readme.md", "Readme.md"]:
        f = repo_dir / cand
        if not f.exists():
            continue
        try:
            content = f.read_text(encoding="utf-8", errors="replace")[:30000]
        except OSError:
            continue
        bg["readme_full"] = content
        # Features section
        features_m = re.search(
            r'(?:^##\s*(?:Features|Highlights|Why\s+\S+|Key\s+Features)\b[^\n]*\n|'
            r'^(?:Features|Highlights|Key\s+Features)\s*\n[=\-~]{3,}\n)'
            r'(.*?)(?=\n##\s|\n[A-Z][^\n]{0,60}\n[=\-~]{3,}\n|\Z)',
            content, re.IGNORECASE | re.DOTALL | re.MULTILINE,
        )
        if features_m:
            feature_lines = []
            for ln in features_m.group(1).split('\n'):
                s = ln.strip()
                if s.startswith(('-', '*', '+', '•')):
                    feature_lines.append(s.lstrip('-*+• ').strip())
            bg["readme_features"] = [x for x in feature_lines if x][:15]
        # Python code blocks (first 3)
        code_blocks = re.findall(r'```(?:python|py)?\n(.*?)\n```', content, re.DOTALL)
        bg["readme_examples"] = [b for b in code_blocks if b.strip()][:3]
        break

    # docs/ folder
    docs_dir = repo_dir / "docs"
    if docs_dir.is_dir():
        try:
            md_files = sorted(docs_dir.rglob("*.md"))[:8]
            rst_files = sorted(docs_dir.rglob("*.rst"))[:8]
        except OSError:
            md_files, rst_files = [], []
        all_docs = (md_files + rst_files)[:10]
        for f in all_docs:
            try:
                content = f.read_text(encoding="utf-8", errors="replace")[:5000]
                rel = f.relative_to(repo_dir).as_posix()
                bg["docs_files"].append({"name": rel, "preview": content[:1500]})
            except OSError:
                pass

    # Main package __init__.py docstring
    try:
        init_candidates = list(repo_dir.glob("*/__init__.py")) + list(repo_dir.glob("src/*/__init__.py"))
    except OSError:
        init_candidates = []
    for init in init_candidates[:5]:
        try:
            content = init.read_text(encoding="utf-8", errors="replace")
        except OSError:
            continue
        m = re.match(r'\s*(?:r|u|b)?(""".*?"""|\'\'\'.*?\'\'\')', content, re.DOTALL)
        if m:
            raw = m.group(1)
            doc = raw.strip('"\'').strip()
            if len(doc) > 30:
                bg["package_docstring"] = doc[:2000]
                break

    # Top class docstrings (by degree)
    classes = [n for n in store.nodes.values() if n.get("kind") == "class" and n.get("docstring")]
    if classes:
        degree: dict = {}
        for c in classes:
            d = sum(1 for e in store.edges if e["src"] == c["id"] or e["dst"] == c["id"])
            degree[c["id"]] = d
        classes.sort(key=lambda c: -degree.get(c["id"], 0))
        for c in classes[:3]:
            bg["top_class_docs"].append({
                "name": c["name"],
                "qualified_name": c.get("qualified_name", c["name"]),
                "doc": (c.get("docstring") or "")[:800],
            })

    # examples/ directory
    ex_dir = repo_dir / "examples"
    if ex_dir.is_dir():
        try:
            py_files = sorted(ex_dir.rglob("*.py"))[:6]
        except OSError:
            py_files = []
        for f in py_files:
            try:
                content = f.read_text(encoding="utf-8", errors="replace")[:1200]
            except OSError:
                continue
            func_m = re.search(r'def\s+(\w+)\s*\(', content)
            try:
                rel = f.relative_to(repo_dir).as_posix()
            except ValueError:
                rel = f.name
            bg["examples_dir"].append({
                "file": rel,
                "preview": content[:600],
                "first_func": func_m.group(1) if func_m else None,
            })

    # CHANGELOG
    for cand in ["CHANGELOG.md", "CHANGELOG.rst", "CHANGES.md", "CHANGES.rst",
                 "HISTORY.md", "HISTORY.rst", "CHANGELOG"]:
        f = repo_dir / cand
        if not f.exists():
            continue
        try:
            content = f.read_text(encoding="utf-8", errors="replace")[:20000]
        except OSError:
            continue
        # Split on version headings
        entries = re.split(
            r'(?=^##?\s*(?:Version\s+)?v?\d+\.\d+|\n\d+\.\d+\.\d+\s*\(\d{4})',
            content, flags=re.MULTILINE,
        )
        for ent in entries[:6]:
            ent = ent.strip()
            if ent and len(ent) > 50:
                bg["changelog_recent"].append(ent[:600])
            if len(bg["changelog_recent"]) >= 5:
                break
        break

    # pyproject classifiers + urls
    pp = repo_dir / "pyproject.toml"
    if pp.exists():
        try:
            content = pp.read_text(encoding="utf-8", errors="replace")
        except OSError:
            content = ""
        cls_m = re.search(r'classifiers\s*=\s*\[(.*?)\]', content, re.DOTALL)
        if cls_m:
            bg["classifiers"] = re.findall(r'["\']([^"\']{5,100})["\']', cls_m.group(1))[:15]
        urls_m = re.search(r'\[project\.urls\]\s*\n(.*?)(?=\n\[|\Z)', content, re.DOTALL)
        if urls_m:
            for line in urls_m.group(1).strip().split('\n'):
                line = line.strip()
                if not line or line.startswith('#'):
                    continue
                if '=' in line:
                    k, v = line.split('=', 1)
                    bg["urls"][k.strip()] = v.strip().strip('"\'')

    # Workflows
    wf_dir = repo_dir / ".github" / "workflows"
    if wf_dir.is_dir():
        try:
            wf_files = sorted(list(wf_dir.glob("*.yml")) + list(wf_dir.glob("*.yaml")))[:10]
        except OSError:
            wf_files = []
        for f in wf_files:
            try:
                content = f.read_text(encoding="utf-8", errors="replace")[:1500]
            except OSError:
                continue
            name_m = re.search(r'^name:\s*(.+)', content, re.MULTILINE)
            bg["workflows"].append({
                "file": f.name,
                "name": (name_m.group(1).strip().strip('"\'') if name_m else f.stem),
            })

    return bg


def extract_directory_tree(repo_dir: Path, max_depth: int = 3, max_entries: int = 120) -> str:
    """Build a simple `tree -L N` style listing, skipping noise."""
    skip_names = {"__pycache__", ".git", "node_modules", ".venv", ".tox", "dist", "build",
                  ".mypy_cache", ".pytest_cache", ".idea", ".vscode", "venv", "env"}
    lines: list[str] = [repo_dir.name + "/"]
    entries = [0]

    def walk(path: Path, prefix: str, depth: int):
        if depth > max_depth or entries[0] >= max_entries:
            return
        try:
            items = sorted(path.iterdir(), key=lambda p: (not p.is_dir(), p.name.lower()))
        except OSError:
            return
        items = [it for it in items if it.name not in skip_names and not it.name.endswith(".pyc")]
        # also skip noisy hidden files except the few we care about
        items = [it for it in items if not (it.name.startswith(".") and it.name not in {".github", ".gitignore", ".env.example"})]
        for i, it in enumerate(items):
            if entries[0] >= max_entries:
                lines.append(prefix + "└── ...")
                return
            is_last = (i == len(items) - 1)
            connector = "└── " if is_last else "├── "
            label = it.name + ("/" if it.is_dir() else "")
            lines.append(prefix + connector + label)
            entries[0] += 1
            if it.is_dir():
                walk(it, prefix + ("    " if is_last else "│   "), depth + 1)

    walk(repo_dir, "", 1)
    return "\n".join(lines)


# ---------------------------------------------------------------------------


DOC_TEMPLATES = {
    "default": {
        "name": "完整文档",
        "description": "Confluence 级交付文档，包含全部图表和契约",
        "sections": [
            "overview", "background", "metadata", "directory", "arch_diagram", "entry_points",
            "sequence_diagram", "class_uml", "api_contracts", "call_flow",
            "exception_hierarchy", "config", "modules", "routes", "tests",
            "dependencies", "recommendations",
        ],
    },
    "architecture": {
        "name": "架构聚焦",
        "description": "突出系统架构和类设计，适合架构评审",
        "sections": [
            "overview", "background", "metadata", "arch_diagram", "class_uml",
            "sequence_diagram", "modules", "dependencies", "recommendations",
        ],
    },
    "api": {
        "name": "API 文档",
        "description": "聚焦接口、契约和调用链路",
        "sections": [
            "overview", "background", "metadata", "entry_points", "api_contracts", "routes",
            "exception_hierarchy", "config", "recommendations",
        ],
    },
    "minimal": {
        "name": "极简",
        "description": "只保留概览、架构图和推荐问答",
        "sections": ["overview", "arch_diagram", "modules", "recommendations"],
    },
}

SECTION_LABELS = {
    "overview": "项目概览",
    "background": "项目背景",
    "metadata": "项目信息",
    "directory": "目录结构",
    "arch_diagram": "系统架构图",
    "entry_points": "入口点与启动流程",
    "sequence_diagram": "关键时序图",
    "class_uml": "核心类 UML",
    "api_contracts": "API 契约",
    "call_flow": "调用关系图",
    "exception_hierarchy": "错误处理",
    "config": "配置项",
    "modules": "关键模块详解",
    "routes": "接口文档",
    "tests": "测试覆盖",
    "dependencies": "外部依赖",
    "recommendations": "推荐问答",
}


# Legacy Mermaid generators (kept for backward compat — unused after PlantUML migration)
def gen_module_arch_diagram(store, gq) -> str:
    """Mermaid graph TD of module dependencies via imports."""
    modules = {nid: n for nid, n in store.nodes.items() if n["kind"] == "module"}
    if not modules:
        return "graph TD\n  A[\"根模块\"]"

    edges = set()
    for e in store.edges:
        if e["kind"] != "imports":
            continue
        src_mod = _find_parent_module(store, e["src"])
        dst_mod = _find_parent_module(store, e["dst"])
        if src_mod and dst_mod and src_mod != dst_mod:
            edges.add((src_mod, dst_mod))

    edge_counts = {}
    for s, d in edges:
        edge_counts[s] = edge_counts.get(s, 0) + 1
        edge_counts[d] = edge_counts.get(d, 0) + 1
    if not edge_counts:
        # No import edges; fall back to listing top modules
        top_modules = list(modules.keys())[:8]
        lines = ["graph TD"]
        for m in top_modules:
            label = modules[m].get("qualified_name", m).split(".")[-1][:30]
            lines.append(f'  M_{_sanitize_mermaid_id(m)}["{label}"]')
        return "\n".join(lines)
    top_modules = sorted(edge_counts.keys(), key=lambda m: -edge_counts[m])[:15]
    top_set = set(top_modules)

    lines = ["graph TD"]

    def safe_id(mid):
        return "M_" + _sanitize_mermaid_id(mid)

    def label(mid):
        return modules.get(mid, {}).get("qualified_name", mid).split(".")[-1][:30]

    for m in top_modules:
        lines.append(f'  {safe_id(m)}["{label(m)}"]')
    for s, d in edges:
        if s in top_set and d in top_set:
            lines.append(f"  {safe_id(s)} --> {safe_id(d)}")
    return "\n".join(lines)


def gen_class_uml_diagram(store, gq) -> str:
    """Mermaid classDiagram for top classes with methods + relationships."""
    classes = {nid: n for nid, n in store.nodes.items() if n["kind"] == "class"}
    if not classes:
        return "classDiagram\n  class Empty"

    degree = {nid: 0 for nid in classes}
    for e in store.edges:
        if e["src"] in degree:
            degree[e["src"]] += 1
        if e["dst"] in degree:
            degree[e["dst"]] += 1

    top = sorted(classes.keys(), key=lambda c: -degree[c])[:8]
    top_set = set(top)

    lines = ["classDiagram"]
    for cid in top:
        c = classes[cid]
        name = c["name"][:30]
        methods, fields = [], []
        for e in store.edges:
            if e["src"] == cid and e["kind"] == "contains":
                child = store.nodes.get(e["dst"])
                if not child:
                    continue
                if child["kind"] in ("method", "function"):
                    methods.append(child["name"][:25])
                elif child["kind"] == "field":
                    fields.append(child["name"][:25])
        lines.append(f"  class {_sanitize_mermaid_id(name)} {{")
        for f in fields[:5]:
            lines.append(f"    +{_sanitize_mermaid_id(f)}")
        for m in methods[:8]:
            lines.append(f"    +{_sanitize_mermaid_id(m)}()")
        lines.append("  }")

    for e in store.edges:
        if e["kind"] in ("extends", "implements") and e["src"] in top_set and e["dst"] in top_set:
            child = _sanitize_mermaid_id(classes[e["src"]]["name"])
            parent = _sanitize_mermaid_id(classes[e["dst"]]["name"])
            arrow = "<|--" if e["kind"] == "extends" else "<|.."
            lines.append(f"  {parent} {arrow} {child}")
    return "\n".join(lines)


def gen_call_flow_diagram(store, gq) -> str:
    """Mermaid graph LR showing call chains from route_handlers (or top methods)."""
    routes = [n for n in store.nodes.values() if n["kind"] == "route_handler"]
    if not routes:
        call_count = {}
        for e in store.edges:
            if e["kind"] == "calls":
                call_count[e["src"]] = call_count.get(e["src"], 0) + 1
        methods = sorted(
            [n for n in store.nodes.values() if n["kind"] in ("method", "function")],
            key=lambda m: -call_count.get(m["id"], 0),
        )[:5]
        if not methods:
            return "graph LR\n  A[\"无调用关系\"]"
        seeds = methods
    else:
        seeds = routes[:5]

    lines = ["graph LR"]
    visited = set()
    edges_drawn = set()

    def add_node(node):
        nid = _sanitize_mermaid_id(node["id"])
        if nid in visited:
            return nid
        visited.add(nid)
        label = (node["name"] or "")[:25].replace('"', "'")
        lines.append(f'  {nid}["{label}"]')
        return nid

    def expand(node, depth):
        if depth <= 0 or len(visited) > 20:
            return
        nid = add_node(node)
        for e in store.edges:
            if e["src"] == node["id"] and e["kind"] in ("calls", "route_handler"):
                target = store.nodes.get(e["dst"])
                if not target:
                    continue
                tid = add_node(target)
                key = (nid, tid)
                if key not in edges_drawn:
                    edges_drawn.add(key)
                    lines.append(f"  {nid} --> {tid}")
                expand(target, depth - 1)

    for seed in seeds:
        expand(seed, 3)
    return "\n".join(lines)


# ---------------------------------------------------------------------------
# Docgen endpoint (auth required)
# ---------------------------------------------------------------------------

@app.get("/api/v1/models")
def list_models_ep(user: dict = Depends(get_current_user)):
    """多模型网关:列出 7 个模型 + 健康/降级/实时统计。"""
    from codedoc.llm_router import list_models, default_model_id
    models = list_models()
    healthy = [m["id"] for m in models if m["healthy"]]
    return {"models": models, "default": default_model_id(),
            "healthy": healthy, "auto": "auto", "total": len(models)}


@app.get("/api/v1/explore")
def explore_ep(anchor: str, repo: str, user: dict = Depends(get_current_user)):
    """邻域子图(图谱交互):走统一工具层 explore。"""
    info = _ensure_repo_indexed(repo)
    if not info:
        raise HTTPException(404, "repo 未索引")
    from codedoc.graph.memory_backend import MemoryGraphQuery
    gq = MemoryGraphQuery(info["cfg"], info["store"])
    return _registry_for(gq, repo, repo_root=info.get("path")).call("explore", {"anchor": anchor})


@app.get("/api/v1/impact")
def impact_ep(node_id: str, repo: str, user: dict = Depends(get_current_user)):
    """改动影响面(反向传递闭包):走统一工具层 impact。"""
    info = _ensure_repo_indexed(repo)
    if not info:
        raise HTTPException(404, "repo 未索引")
    from codedoc.graph.memory_backend import MemoryGraphQuery
    gq = MemoryGraphQuery(info["cfg"], info["store"])
    return _registry_for(gq, repo, repo_root=info.get("path")).call("impact", {"node_id": node_id})


# ---- 内部 server-to-server 工具端点(供 devbot 等复用 codedoc 检索能力,内部 key 鉴权)----
_INTERNAL_KEY = os.environ.get("CODEDOC_INTERNAL_KEY", "CHANGE_ME_INTERNAL_KEY")


def _check_internal(request: Request):
    if request.headers.get("x-internal-key") != _INTERNAL_KEY:
        raise HTTPException(403, "invalid internal key")


from codedoc.graph.query_cache import QueryCache
def _build_gq(cfg, store):
    from codedoc.graph.memory_backend import MemoryGraphQuery
    return MemoryGraphQuery(cfg, store)
_GQ_CACHE = QueryCache(_build_gq)


def _gq_info_for(repo: str):
    info = _ensure_repo_indexed(repo)
    if not info:
        raise HTTPException(404, "repo 未索引: %s" % repo)
    from codedoc.graph.memory_backend import MemoryGraphQuery
    return _GQ_CACHE.get(repo, info["cfg"], info["store"]), info


class ToolSearchReq(BaseModel):
    repo: str
    query: str
    top_k: int = 8


class ToolBodyReq(BaseModel):
    repo: str
    node_id: str
    max_lines: int = 50


class ToolImpactReq(BaseModel):
    repo: str = ""
    files: list[str] = []


@app.post("/tools/search")
def tools_search(req: ToolSearchReq, request: Request):
    """语义+全文混合检索(供外部服务接地用)。"""
    _check_internal(request)
    gq, info = _gq_info_for(req.repo)
    return _registry_for(gq, req.repo, repo_root=info.get("path")).call(
        "search", {"query": req.query, "top_k": req.top_k})


@app.post("/tools/get_body")
def tools_get_body(req: ToolBodyReq, request: Request):
    """按 node_id 取真实函数体(file+行号 deref)。"""
    _check_internal(request)
    gq, info = _gq_info_for(req.repo)
    return _registry_for(gq, req.repo, repo_root=info.get("path")).call(
        "get_body", {"id": req.node_id, "max_lines": req.max_lines})


@app.post("/tools/impact")
def tools_impact(req: ToolImpactReq, request: Request):
    """变更影响:files → 命中文件的节点 → impact 反向闭包聚合。"""
    _check_internal(request)
    if not req.repo or not req.files:
        return {"items": [], "edges": [], "summary": "no repo/files"}
    gq, info = _gq_info_for(req.repo)
    reg = _registry_for(gq, req.repo, repo_root=info.get("path"))
    store = info["store"]
    seeds = [nid for nid, n in store.nodes.items()
             if (n.get("file") and any(n["file"].endswith(f) or f.endswith(n["file"]) for f in req.files))]
    affected = {}
    for nid in seeds[:20]:
        try:
            r = reg.call("impact", {"node_id": nid})
        except Exception:
            continue
        for it in r.get("items", []):
            affected[it.get("node_id") or it.get("qualified_name") or str(len(affected))] = it
    items = list(affected.values())
    return {"items": items, "edges": [], "summary": "%d affected nodes from %d files" % (len(items), len(req.files))}


@app.get("/api/v1/docgen/templates")
def docgen_templates(user: dict = Depends(get_current_user)):
    """Return predefined doc templates and the section catalog for the UI picker."""
    return {"templates": DOC_TEMPLATES, "sections": SECTION_LABELS}


# ---------- Async docgen worker -------------------------------------------
#
# Persistence model: task records + queue live in Postgres (`docgen_tasks`,见
# codedoc/docgen_queue.py)。每个 uvicorn worker 跑一个领取循环,经 FOR UPDATE
# SKIP LOCKED 抢单 —— 任意 worker 可领任意任务(真正跨 worker),崩溃由 reclaim_orphans
# 打回 queued 重跑。_docgen_tasks 仅作处理中任务的本地进度缓存。

def _start_docgen_worker():
    """启动本进程的 docgen 领取循环(若未在跑)。每个 uvicorn worker 一个,
    经 PG FOR UPDATE SKIP LOCKED 抢单 → 真正跨 worker 队列。env CODEDOC_DOCGEN_WORKER=0 可关。"""
    global _docgen_worker_thread
    if os.environ.get("CODEDOC_DOCGEN_WORKER", "1") != "1":
        return
    if _docgen_worker_thread and _docgen_worker_thread.is_alive():
        return
    _docgen_worker_thread = threading.Thread(target=_docgen_worker_loop, daemon=True)
    _docgen_worker_thread.start()


def _docgen_worker_loop():
    """持续从 PG 队列领取 docgen 任务(多 worker 安全),逐个生成。启动先做崩溃孤儿重建。"""
    import time as _time
    from codedoc import docgen_queue as _dq
    try:
        _dq.ensure_schema()
        n = _dq.reclaim_orphans(600)
        if n:
            print("[docgen] reclaimed %d orphan task(s)" % n)
    except Exception:
        pass
    while True:
        row = None
        try:
            row = _dq.claim_one(time.time())
        except Exception:
            row = None
        if not row:
            _time.sleep(3)
            continue
        task_id = row["task_id"]
        try:
            sections = json.loads(row["sections"]) if row.get("sections") else []
        except Exception:
            sections = []
        task = {"task_id": task_id, "user_id": row["user_id"], "repo": row["repo"],
                "sections": sections, "template": row["template"], "status": "running",
                "stage": "running", "progress": "", "position": 0,
                "submitted_at": row.get("submitted_at") or time.time()}
        with _docgen_lock:
            _docgen_tasks[task_id] = dict(task)
        try:
            _run_docgen_task(task_id, task)
        except Exception as e:  # noqa: BLE001
            import traceback as _tb
            err_msg = str(e)[:500]
            with _docgen_lock:
                if task_id in _docgen_tasks:
                    _docgen_tasks[task_id]["status"] = "error"
                    _docgen_tasks[task_id]["stage"] = "error"
                    _docgen_tasks[task_id]["error"] = err_msg
                    _docgen_tasks[task_id]["finished_at"] = time.time()
            try:
                _docgen_db_update(task_id, status="error", stage="error",
                                  progress="", error=err_msg,
                                  finished_at=time.time())
            except Exception:  # noqa: BLE001
                pass
            try:
                import logging as _lg
                _lg.getLogger("codedoc").error("docgen task %s failed: %s\n%s",
                                               task_id, err_msg, _tb.format_exc()[-1500:])
            except Exception:  # noqa: BLE001
                pass


def _docgen_db_insert(task_id: str, user_id: int, repo: str, sections: list[str],
                      template: str, submitted_at: float):
    from codedoc import docgen_queue as _dq
    _dq.insert(task_id, user_id, repo, json.dumps(sections), template, submitted_at)


def _docgen_db_update(task_id: str, **fields):
    """Persist a state delta to the PG docgen queue. keys must match column names."""
    if not fields:
        return
    from codedoc import docgen_queue as _dq
    _dq.update(task_id, fields)


def _docgen_db_fetch(task_id: str) -> dict | None:
    from codedoc import docgen_queue as _dq
    d = _dq.fetch(task_id)
    if not d:
        return None
    d = dict(d)
    try:
        d["sections"] = json.loads(d["sections"]) if d.get("sections") else []
    except Exception:
        d["sections"] = []
    try:
        d["recommendations"] = json.loads(d["recommendations"]) if d.get("recommendations") else []
    except Exception:
        d["recommendations"] = []
    return d


def _docgen_db_queue_position(task_id: str, submitted_at: float) -> int:
    """Return 1-based position of a queued task among all currently queued tasks."""
    from codedoc import docgen_queue as _dq
    return _dq.queue_position(task_id, submitted_at)


def _docgen_db_queue_total() -> int:
    from codedoc import docgen_queue as _dq
    return _dq.queue_total()


def _update_stage(task_id: str, stage: str, progress: str = ""):
    """Update current stage of a docgen task and mark it running."""
    started_at = None
    with _docgen_lock:
        t = _docgen_tasks.get(task_id)
        if t:
            t["stage"] = stage
            t["progress"] = progress
            t["position"] = 0
            if t["status"] == "queued":
                t["status"] = "running"
                started_at = time.time()
                t["started_at"] = started_at
    # Mirror to sqlite (best-effort)
    update = {"stage": stage, "progress": progress}
    if started_at is not None:
        update["status"] = "running"
        update["started_at"] = started_at
    try:
        _docgen_db_update(task_id, **update)
    except Exception:  # noqa: BLE001
        pass


def _run_docgen_task(task_id: str, task: dict):
    """Run the actual doc generation in stages, updating progress as we go."""
    repo_name = task["repo"]
    if repo_name.startswith("multi://"):
        return _run_multi_docgen_task(task_id, task)
    requested = task["sections"]
    template = task["template"]
    section_set = set(requested)

    # Stage: load index
    _update_stage(task_id, "loading", "加载仓库索引...")
    info = _ensure_repo_indexed(repo_name)
    if not info:
        finished_at = time.time()
        with _docgen_lock:
            _docgen_tasks[task_id]["status"] = "error"
            _docgen_tasks[task_id]["stage"] = "error"
            _docgen_tasks[task_id]["error"] = "仓库未索引"
            _docgen_tasks[task_id]["finished_at"] = finished_at
        try:
            _docgen_db_update(task_id, status="error", stage="error",
                              error="仓库未索引", finished_at=finished_at)
        except Exception:  # noqa: BLE001
            pass
        return

    from codedoc.agents.llm import ChatMessage, build_llm
    from codedoc.graph.memory_backend import MemoryGraphQuery
    store = info["store"]
    cfg = info["cfg"]
    gq = MemoryGraphQuery(cfg, store)
    llm = build_llm(cfg)
    repo_dir = Path(info.get("path") or (REPOS_DIR / repo_name.split("/")[-1]))

    # Pre-compute shared data
    modules = [n for n in store.nodes.values() if n["kind"] == "module"]
    classes = [n for n in store.nodes.values() if n["kind"] == "class"]
    routes = [n for n in store.nodes.values() if n["kind"] == "route_handler"]

    def _degree(nid):
        return sum(1 for e in store.edges if e["src"] == nid or e["dst"] == nid)

    def _out(nid):
        return sum(1 for e in store.edges if e["src"] == nid)

    # Project metadata is needed by both overview and metadata sections
    project_meta: dict = {}
    if any(s in section_set for s in ("overview", "metadata", "dependencies", "config")):
        try:
            project_meta = extract_project_metadata(repo_dir)
        except Exception:  # noqa: BLE001
            project_meta = {}

    # Build context summary (used by overview + recommendations)
    langs = sorted({n.get("language", "") for n in store.nodes.values() if n.get("language")})
    summary = f"仓库: {repo_name}\n"
    summary += f"语言: {', '.join(langs)}\n"
    summary += f"模块数: {len(modules)} | 类: {len(classes)} | 路由: {len(routes)}\n"
    if project_meta.get("pyproject_description"):
        summary += f"pyproject 描述: {project_meta['pyproject_description']}\n"
    if project_meta.get("readme_intro"):
        summary += f"README 摘要: {project_meta['readme_intro'][:300]}\n"
    summary += "\n核心类列表:\n"
    for c in sorted(classes, key=lambda x: -_degree(x["id"]))[:15]:
        line = f"  - {c['qualified_name']}"
        if c.get("docstring"):
            line += f": {c['docstring'][:80]}"
        summary += line + "\n"
    summary += "\n核心路由:\n"
    for r in routes[:10]:
        summary += f"  - {r['qualified_name']}\n"

    # Generate sections
    overview = ""
    if "overview" in section_set:
        _update_stage(task_id, "overview", "正在生成项目概览...")
        overview = llm.chat([
            ChatMessage("system",
                "你是技术文档作者。基于以下代码结构 + README/pyproject 信息，写一段"
                "项目概览（200-300 字），包括：项目定位、技术栈、主要能力、典型使用场景。"
                "用 Markdown 段落，不要包含标题，不要照搬 README 原文。"),
            ChatMessage("user", summary),
        ], max_tokens=600)

    arch_uml = ""
    if "arch_diagram" in section_set:
        _update_stage(task_id, "arch_diagram", "正在绘制系统架构图...")
        arch_uml = gen_module_arch_plantuml(store, gq)

    class_uml = ""
    if "class_uml" in section_set:
        _update_stage(task_id, "class_uml", "正在绘制类 UML 图...")
        class_uml = gen_class_uml_plantuml(store, gq)

    call_uml = ""
    if "call_flow" in section_set:
        _update_stage(task_id, "call_flow", "正在绘制调用关系图...")
        call_uml = gen_call_flow_plantuml(store, gq)

    module_details = ""
    if "modules" in section_set:
        _update_stage(task_id, "modules", "正在分析关键模块...")
        top_modules = sorted(modules, key=lambda x: -_out(x["id"]))[:5]
        for i, m in enumerate(top_modules):
            _update_stage(task_id, "modules", f"分析模块 {i+1}/{len(top_modules)}: {m['qualified_name'][:40]}")
            children = [n for n in store.nodes.values() if n.get("parent_id") == m["id"]]
            class_names = [c["name"] for c in children if c["kind"] == "class"][:8]
            method_count = len([c for c in children if c["kind"] in ("method", "function")])
            detail_summary = (
                f"模块: {m['qualified_name']}\n"
                f"类: {class_names}\n"
                f"方法数: {method_count}\n"
            )
            mod_text = llm.chat([
                ChatMessage("system",
                    "为这个代码模块写一段简短的设计说明（100-150 字）。"
                    "包含：职责、核心类、与其他模块的关系。"),
                ChatMessage("user", detail_summary),
            ], max_tokens=300)
            module_details += f"### {m['qualified_name']}\n\n{mod_text}\n\n"

    questions: list[str] = []
    if "recommendations" in section_set:
        _update_stage(task_id, "recommendations", "正在生成推荐问答...")
        top_class_names = [c["name"] for c in classes[:10]]
        rec_prompt = (
            f"根据这个项目（{repo_name}），生成 6 个推荐的代码问答问题，"
            f"让用户深入了解。问题应该具体、实用。\n"
            f"项目核心类: {top_class_names}\n"
            f"路由数: {len(routes)}\n"
            f"返回 JSON 数组，格式: [\"问题1\", \"问题2\", ...]"
        )
        rec_resp = llm.chat([
            ChatMessage("system",
                "生成代码问答推荐问题。只返回 JSON 数组，不要其他文字。"),
            ChatMessage("user", rec_prompt),
        ], max_tokens=400)
        import json as _json
        import re as _re
        try:
            m = _re.search(r'\[.*?\]', rec_resp, _re.DOTALL)
            if m:
                questions = _json.loads(m.group(0))
                questions = [q for q in questions if isinstance(q, str)]
        except Exception:
            pass
        if not questions:
            questions = [
                f"{repo_name} 的核心功能有哪些？",
                f"项目的入口在哪里？",
                f"主要类之间的依赖关系是什么？",
                f"有哪些重要的公共 API？",
                f"项目是如何组织代码的？",
                f"哪些模块被依赖最多？",
            ]

    # ---- New deterministic extractors ----
    entry_points: list[dict] = []
    if "entry_points" in section_set or "sequence_diagram" in section_set:
        _update_stage(task_id, "entry_points", "扫描入口点...")
        try:
            entry_points = extract_entry_points(repo_dir, store)
        except Exception:  # noqa: BLE001
            entry_points = []

    sequence_diagrams: list[dict] = []
    if "sequence_diagram" in section_set:
        _update_stage(task_id, "sequence_diagram", "绘制时序图...")
        # Pick up to 3 seed nodes: prefer route_handlers, then graph nodes matching entry __main__ files
        seeds: list[dict] = []
        seen_ids: set[str] = set()
        for r in routes[:3]:
            seeds.append(r)
            seen_ids.add(r["id"])
        if len(seeds) < 3:
            # Use most-out-degree functions/methods as fallback seeds
            cands = [n for n in store.nodes.values() if n.get("kind") in ("function", "method")
                     and not (n.get("name") or "").startswith("_")]
            cands.sort(key=lambda n: -_out(n["id"]))
            for c in cands:
                if c["id"] in seen_ids:
                    continue
                seeds.append(c)
                seen_ids.add(c["id"])
                if len(seeds) >= 3:
                    break
        for s in seeds[:3]:
            try:
                puml = gen_sequence_diagram_plantuml(store, s, depth=4, max_calls=14)
            except Exception:  # noqa: BLE001
                puml = ""
            if puml:
                sequence_diagrams.append({"node": s, "puml": puml})

    api_contracts_data: list[dict] = []
    if "api_contracts" in section_set:
        _update_stage(task_id, "api_contracts", "提取 API 契约...")
        # Pick top 8 public classes by degree
        public_classes = [c for c in classes if not (c.get("name") or "").startswith("_")]
        public_classes.sort(key=lambda c: -_degree(c["id"]))
        top_class_ids = [c["id"] for c in public_classes[:8]]
        try:
            api_contracts_data = extract_api_contracts(repo_dir, store, top_class_ids)
        except Exception:  # noqa: BLE001
            api_contracts_data = []

    exception_data: dict = {}
    exception_uml = ""
    exception_counts: dict = {}
    if "exception_hierarchy" in section_set:
        _update_stage(task_id, "exception_hierarchy", "梳理异常体系...")
        try:
            exception_data = extract_exception_hierarchy(store)
            exception_uml = gen_exception_uml_plantuml(exception_data.get("list", []))
            exception_counts = count_raise_catch_sites(repo_dir, exception_data.get("name_set", set()))
        except Exception:  # noqa: BLE001
            exception_data = {"count": 0, "list": []}
            exception_counts = {"raise": {}, "catch": {}}

    config_data: dict = {}
    if "config" in section_set:
        _update_stage(task_id, "config", "扫描配置项...")
        try:
            config_data = extract_config(repo_dir, store)
        except Exception:  # noqa: BLE001
            config_data = {"env_vars": [], "config_files": [], "config_classes": [], "constants": []}

    test_data: dict = {}
    if "tests" in section_set:
        _update_stage(task_id, "tests", "统计测试覆盖...")
        try:
            test_data = extract_test_coverage(store)
        except Exception:  # noqa: BLE001
            test_data = {"total_tests": 0, "module_count": 0, "by_module": {}, "untested_classes": []}

    deps_data: dict = {}
    if "dependencies" in section_set:
        _update_stage(task_id, "dependencies", "解析外部依赖...")
        try:
            deps_data = extract_dependencies(repo_dir, store)
        except Exception:  # noqa: BLE001
            deps_data = {"deps": [], "categorized": {}, "count": 0}

    directory_tree = ""
    if "directory" in section_set:
        _update_stage(task_id, "directory", "构建目录树...")
        try:
            directory_tree = extract_directory_tree(repo_dir)
        except Exception:  # noqa: BLE001
            directory_tree = ""

    # Assemble doc in requested order
    _update_stage(task_id, "finalizing", "整理文档...")
    parts = [f"# {repo_name} 设计文档", ""]
    idx = 0

    def _h(title: str) -> str:
        nonlocal idx
        idx += 1
        return f"## {idx}. {title}"

    for sec in requested:
        if sec == "overview":
            parts.append(_h("项目概览"))
            parts.append("")
            # 1) README first paragraph verbatim — highest fidelity description
            if project_meta.get("readme_intro"):
                parts.append(f"> **README 摘要** (`{project_meta.get('readme_file', 'README')}`)：")
                parts.append(f"> {project_meta['readme_intro']}")
                parts.append("")
            # 2) LLM-synthesized overview that reconciles README + code structure
            if overview:
                parts.append(overview)
                parts.append("")
            # 3) Structural stats
            parts.append(
                f"**结构统计**: {len(modules)} 个模块 | {len(classes)} 个类 | "
                f"{len(routes)} 个路由 | {len(store.nodes)} 个节点 | {len(store.edges)} 条关系"
            )
            parts.append("")
        elif sec == "background":
            parts.append(_h("项目背景"))
            parts.append("")
            try:
                bg = extract_project_background(repo_dir, store)
            except Exception:  # noqa: BLE001
                bg = {}
            # 1.1 项目定位 — first README paragraph
            positioning_text = ""
            if bg.get("readme_full"):
                para = ""
                for line in bg["readme_full"].split('\n'):
                    stripped = line.strip()
                    if stripped.startswith('#') or stripped.startswith('[!['):
                        continue
                    if set(stripped) <= {"=", "-", "~", "*", "_"} and stripped:
                        continue
                    if not stripped:
                        if para:
                            break
                        continue
                    para += stripped + " "
                    if len(para) > 400:
                        break
                positioning_text = para.strip()
            if not positioning_text and project_meta.get("readme_intro"):
                positioning_text = project_meta["readme_intro"]
            if positioning_text:
                parts.append("### 项目定位")
                parts.append("")
                parts.append(positioning_text)
                parts.append("")
            # 1.2 核心特性
            if bg.get("readme_features"):
                parts.append("### 核心特性")
                parts.append("")
                for feat in bg["readme_features"]:
                    parts.append(f"- {feat}")
                parts.append("")
            # 1.3 设计目标
            if bg.get("package_docstring") or bg.get("top_class_docs"):
                parts.append("### 设计目标")
                parts.append("")
                if bg.get("package_docstring"):
                    parts.append("**主包说明** (`__init__.py` docstring)：")
                    parts.append("")
                    parts.append("```")
                    parts.append(bg["package_docstring"][:600])
                    parts.append("```")
                    parts.append("")
                if bg.get("top_class_docs"):
                    parts.append("**核心抽象的设计意图**：")
                    parts.append("")
                    for c in bg["top_class_docs"]:
                        doc_snip = (c.get("doc") or "").strip().replace("\n", " ")
                        parts.append(f"- `{c['name']}` — {doc_snip[:200]}")
                    parts.append("")
            # 1.4 典型使用场景
            if bg.get("readme_examples") or bg.get("examples_dir"):
                parts.append("### 典型使用场景")
                parts.append("")
                for i, ex in enumerate(bg.get("readme_examples", [])[:2]):
                    parts.append(f"**示例 {i+1}**（来自 README）：")
                    parts.append("")
                    parts.append("```python")
                    parts.append(ex[:600])
                    parts.append("```")
                    parts.append("")
                if bg.get("examples_dir"):
                    parts.append("**examples/ 目录**：")
                    parts.append("")
                    for ex in bg["examples_dir"][:5]:
                        line = f"- `{ex['file']}`"
                        if ex.get("first_func"):
                            line += f" — `{ex['first_func']}()`"
                        parts.append(line)
                    parts.append("")
            # 1.5 项目演进
            if bg.get("changelog_recent"):
                parts.append("### 项目演进")
                parts.append("")
                parts.append(f"以下为 CHANGELOG 中最近 {min(len(bg['changelog_recent']), 3)} 条记录：")
                parts.append("")
                for entry in bg["changelog_recent"][:3]:
                    parts.append("```")
                    parts.append(entry[:400])
                    parts.append("```")
                    parts.append("")
            # 1.6 目标受众
            if bg.get("classifiers"):
                parts.append("### 目标受众与适用环境")
                parts.append("")
                parts.append("根据 `pyproject.toml` 的 PyPI Trove classifiers：")
                parts.append("")
                for c in bg["classifiers"]:
                    parts.append(f"- {c}")
                parts.append("")
            # 1.7 项目链接
            if bg.get("urls"):
                parts.append("### 项目链接")
                parts.append("")
                for k, v in bg["urls"].items():
                    parts.append(f"- **{k}**: {v}")
                parts.append("")
            # 1.8 CI 关注点
            if bg.get("workflows"):
                parts.append("### CI 关注点")
                parts.append("")
                parts.append("项目通过以下 GitHub Actions workflow 保证质量：")
                parts.append("")
                for wf in bg["workflows"][:8]:
                    parts.append(f"- `{wf['file']}` — {wf.get('name', wf['file'])}")
                parts.append("")
            # Empty fallback
            has_any = any([
                positioning_text, bg.get("readme_features"), bg.get("package_docstring"),
                bg.get("top_class_docs"), bg.get("readme_examples"), bg.get("examples_dir"),
                bg.get("changelog_recent"), bg.get("classifiers"), bg.get("urls"),
                bg.get("workflows"),
            ])
            if not has_any:
                parts.append("_未在仓库中发现 README、CHANGELOG、examples/、docs/ 等可挖掘的背景信息_")
                parts.append("")
        elif sec == "metadata":
            parts.append(_h("项目信息"))
            parts.append("")
            if project_meta:
                parts.append("| 字段 | 值 |")
                parts.append("|------|----|")
                rows = [
                    ("名称", project_meta.get("pyproject_name")),
                    ("版本", project_meta.get("pyproject_version")),
                    ("描述", project_meta.get("pyproject_description")),
                    ("Python 版本", project_meta.get("pyproject_requires_python")),
                    ("License", project_meta.get("license_type") or project_meta.get("pyproject_license")),
                    ("Homepage", project_meta.get("pyproject_homepage")),
                    ("Repository", project_meta.get("pyproject_repository")),
                ]
                for label, val in rows:
                    if val:
                        parts.append(f"| {label} | {val} |")
                if project_meta.get("pyproject_keywords"):
                    parts.append(f"| Keywords | {', '.join(project_meta['pyproject_keywords'])} |")
                if project_meta.get("pyproject_authors"):
                    parts.append(f"| Authors | {', '.join(project_meta['pyproject_authors'])} |")
                if project_meta.get("docker_files"):
                    parts.append(f"| Docker | {', '.join(f'`{x}`' for x in project_meta['docker_files'])} |")
                if project_meta.get("makefile_targets"):
                    parts.append(f"| Make 目标 | {', '.join(f'`{x}`' for x in project_meta['makefile_targets'][:10])} |")
                if project_meta.get("ci_workflows"):
                    parts.append(f"| CI 工作流 | {', '.join(project_meta['ci_workflows'][:8])} |")
            else:
                parts.append("_未检测到 `pyproject.toml` 等项目元数据文件_")
            parts.append("")
        elif sec == "directory":
            parts.append(_h("目录结构"))
            parts.append("")
            if directory_tree:
                parts.append("```")
                parts.append(directory_tree)
                parts.append("```")
            else:
                parts.append("_目录扫描失败_")
            parts.append("")
        elif sec == "entry_points":
            parts.append(_h("入口点与启动流程"))
            parts.append("")
            if entry_points:
                # Group by type
                groups: dict[str, list[dict]] = {}
                for ep in entry_points:
                    groups.setdefault(ep["type"], []).append(ep)
                type_labels = {
                    "console_script": "命令行入口 (console_scripts)",
                    "__main__": "`__main__` 入口脚本",
                    "route": "HTTP 路由入口",
                    "cli_command": "CLI 命令 (click/typer)",
                }
                for t, items in groups.items():
                    parts.append(f"### {type_labels.get(t, t)}")
                    parts.append("")
                    if t == "console_script":
                        parts.append("| 命令 | 目标 |")
                        parts.append("|------|------|")
                        for ep in items:
                            parts.append(f"| `{ep['name']}` | `{ep['target']}` |")
                    elif t == "__main__":
                        for ep in items:
                            parts.append(f"- `{ep['file']}`")
                    elif t == "route":
                        parts.append("| 路径 / 名称 | Handler | 文件 |")
                        parts.append("|------|---------|------|")
                        for ep in items:
                            parts.append(f"| `{ep['name']}` | `{ep['qualified_name']}` | `{ep.get('file', '')}` |")
                    elif t == "cli_command":
                        for ep in items:
                            parts.append(f"- `{ep['qualified_name']}` ({ep.get('file', '')})")
                    parts.append("")
            else:
                parts.append("_未检测到显式入口点（无 console_scripts、无 `__main__` 块、无路由处理器）_")
                parts.append("")
        elif sec == "sequence_diagram":
            parts.append(_h("关键时序图"))
            parts.append("")
            if sequence_diagrams:
                parts.append(f"以下时序图基于代码 `calls` 关系图自动追踪（深度 4），共选取 {len(sequence_diagrams)} 条典型路径：")
                parts.append("")
                for i, sd in enumerate(sequence_diagrams):
                    node = sd["node"]
                    parts.append(f"### {i+1}. {node.get('qualified_name') or node.get('name')}")
                    parts.append("")
                    seq_url = plantuml_url(sd["puml"])
                    parts.append(f"![时序图 {i+1}]({seq_url})")
                    parts.append("")
                    parts.append(
                        "<details><summary>查看 PlantUML 源码</summary>\n\n```\n"
                        + sd["puml"] + "\n```\n</details>"
                    )
                    parts.append("")
            else:
                parts.append("_未提取到可绘制的调用链（图中缺少 `calls` 边）_")
                parts.append("")
        elif sec == "arch_diagram":
            parts.append(_h("系统架构图"))
            parts.append("")
            parts.append("下图展示模块间的依赖关系：")
            parts.append("")
            arch_url = plantuml_url(arch_uml)
            parts.append(f"![模块架构图]({arch_url})")
            parts.append("")
            parts.append(
                "<details><summary>查看 PlantUML 源码</summary>\n\n```\n"
                + arch_uml + "\n```\n</details>"
            )
            parts.append("")
        elif sec == "class_uml":
            parts.append(_h("核心类 UML"))
            parts.append("")
            class_url = plantuml_url(class_uml)
            parts.append(f"![核心类 UML]({class_url})")
            parts.append("")
            parts.append(
                "<details><summary>查看 PlantUML 源码</summary>\n\n```\n"
                + class_uml + "\n```\n</details>"
            )
            parts.append("")
        elif sec == "call_flow":
            parts.append(_h("调用关系图"))
            parts.append("")
            call_url = plantuml_url(call_uml)
            parts.append(f"![调用关系图]({call_url})")
            parts.append("")
            parts.append(
                "<details><summary>查看 PlantUML 源码</summary>\n\n```\n"
                + call_uml + "\n```\n</details>"
            )
            parts.append("")
        elif sec == "modules":
            parts.append(_h("关键模块详解"))
            parts.append("")
            parts.append(module_details)
            parts.append("")
        elif sec == "routes":
            parts.append(_h("接口文档"))
            parts.append("")
            if routes:
                parts.append("| 路径 | Handler | 文件 |")
                parts.append("|------|---------|------|")
                for r in routes[:20]:
                    parts.append(
                        f"| `{r.get('name', '')}` | `{r['qualified_name']}` | `{r.get('file', '')}` |"
                    )
            else:
                parts.append("_此项目未检测到 HTTP 路由_")
            parts.append("")
        elif sec == "api_contracts":
            parts.append(_h("API 契约"))
            parts.append("")
            if api_contracts_data:
                parts.append(f"以下展示 {len(api_contracts_data)} 个核心公开类的方法契约：签名、文档、可能抛出/捕获的异常、自身状态写入：")
                parts.append("")
                for cls in api_contracts_data:
                    parts.append(f"### `{cls['class']}`")
                    parts.append("")
                    if cls.get("docstring"):
                        parts.append(f"> {cls['docstring']}")
                        parts.append("")
                    if cls.get("file"):
                        parts.append(f"_声明位置_: `{cls['file']}`")
                        parts.append("")
                    if not cls["methods"]:
                        parts.append("_无公开方法_")
                        parts.append("")
                        continue
                    for m in cls["methods"]:
                        sig = m.get("signature") or "()"
                        # signatures sometimes include the function name; strip if present
                        if not sig.startswith("("):
                            # try to extract just the parameter portion
                            mm = re.search(r"\((.*)\)\s*(->.*)?", sig)
                            if mm:
                                ret = mm.group(2) or ""
                                sig_clean = "(" + mm.group(1) + ")" + (" " + ret.strip() if ret else "")
                            else:
                                sig_clean = sig
                        else:
                            sig_clean = sig
                        parts.append(f"#### `{m['name']}{sig_clean}`")
                        parts.append("")
                        if m.get("docstring"):
                            parts.append(m["docstring"])
                            parts.append("")
                        attrs: list[str] = []
                        if m.get("raises"):
                            attrs.append(f"- **可能抛出**: {', '.join('`'+r+'`' for r in m['raises'])}")
                        if m.get("catches"):
                            attrs.append(f"- **捕获异常**: {', '.join('`'+r+'`' for r in m['catches'])}")
                        if m.get("writes_self"):
                            attrs.append(f"- **修改实例状态**: {', '.join('`self.'+w+'`' for w in m['writes_self'])}")
                        if m.get("file"):
                            attrs.append(f"- **来源**: `{m['file']}`")
                        if attrs:
                            parts.extend(attrs)
                            parts.append("")
            else:
                parts.append("_未提取到可展示的 API 契约_")
                parts.append("")
        elif sec == "exception_hierarchy":
            parts.append(_h("错误处理"))
            parts.append("")
            count = (exception_data or {}).get("count", 0)
            lst = (exception_data or {}).get("list", [])
            if count > 0:
                parts.append(f"项目自定义异常类共 **{count}** 个：")
                parts.append("")
                if exception_uml:
                    exc_url = plantuml_url(exception_uml)
                    parts.append(f"![异常类继承关系]({exc_url})")
                    parts.append("")
                    parts.append(
                        "<details><summary>查看 PlantUML 源码</summary>\n\n```\n"
                        + exception_uml + "\n```\n</details>"
                    )
                    parts.append("")
                parts.append("| 异常类 | 父类 | raise 次数 | except 次数 | 文件 |")
                parts.append("|--------|------|-----------|------------|------|")
                rc = (exception_counts or {}).get("raise", {})
                cc = (exception_counts or {}).get("catch", {})
                for exc in lst[:25]:
                    nm = exc["name"]
                    parents = ", ".join(exc.get("parents") or []) or "—"
                    parts.append(
                        f"| `{nm}` | {parents} | {rc.get(nm, 0)} | {cc.get(nm, 0)} | `{exc.get('file', '')}` |"
                    )
            else:
                parts.append("_未发现自定义异常类_")
            parts.append("")
        elif sec == "config":
            parts.append(_h("配置项"))
            parts.append("")
            if config_data:
                env_vars = config_data.get("env_vars", [])
                if env_vars:
                    parts.append(f"### 环境变量 (共 {len(env_vars)} 项)")
                    parts.append("")
                    parts.append("| 变量名 | 默认值 |")
                    parts.append("|--------|--------|")
                    for name, default in env_vars[:40]:
                        d = default if default else "—"
                        parts.append(f"| `{name}` | {d} |")
                    parts.append("")
                cfg_classes = config_data.get("config_classes", [])
                if cfg_classes:
                    parts.append("### 配置类")
                    parts.append("")
                    parts.append("| 类名 | 文件 |")
                    parts.append("|------|------|")
                    for c in cfg_classes:
                        parts.append(f"| `{c['qualified_name']}` | `{c.get('file', '')}` |")
                    parts.append("")
                cfg_files = config_data.get("config_files", [])
                if cfg_files:
                    parts.append("### 配置文件引用")
                    parts.append("")
                    for cf in cfg_files:
                        parts.append(f"- `{cf}`")
                    parts.append("")
                consts = config_data.get("constants", [])
                if consts:
                    parts.append("### 模块级常量 (设置/配置模块)")
                    parts.append("")
                    parts.append("| 文件 | 常量 | 值 |")
                    parts.append("|------|------|----|")
                    for fn, nm, val in consts[:20]:
                        v_safe = val.replace("|", "\\|")
                        parts.append(f"| `{fn}` | `{nm}` | `{v_safe}` |")
                    parts.append("")
                if not (env_vars or cfg_classes or cfg_files or consts):
                    parts.append("_未发现可识别的配置项_")
                    parts.append("")
            else:
                parts.append("_配置扫描失败_")
                parts.append("")
        elif sec == "tests":
            parts.append(_h("测试覆盖"))
            parts.append("")
            if test_data and test_data.get("total_tests"):
                parts.append(
                    f"**总测试数**: {test_data['total_tests']} | "
                    f"**测试模块数**: {test_data['module_count']} | "
                    f"**类总数**: {test_data.get('class_total', 0)} | "
                    f"**未被引用类（估算）**: {test_data.get('untested_total', 0)}"
                )
                parts.append("")
                parts.append("### 测试模块分布")
                parts.append("")
                parts.append("| 测试文件 | 用例数 | 示例用例 |")
                parts.append("|----------|--------|----------|")
                for mod, tests in list(test_data.get("by_module", {}).items())[:25]:
                    sample = ", ".join(f"`{t['name']}`" for t in tests[:3])
                    parts.append(f"| `{mod}` | {len(tests)} | {sample} |")
                parts.append("")
                # Show some humanized test descriptions
                first_mod = next(iter(test_data.get("by_module", {}).values()), [])
                if first_mod:
                    parts.append("### 示例用例描述（从测试名推断）")
                    parts.append("")
                    for t in first_mod[:10]:
                        parts.append(f"- `{t['name']}` — 测试 {t['description']}")
                    parts.append("")
                untested = test_data.get("untested_classes", [])
                if untested:
                    parts.append(f"### 可能未被直接测试的类 (示例 {min(len(untested), 15)} 个)")
                    parts.append("")
                    for c in untested[:15]:
                        parts.append(f"- `{c['qualified_name']}` — `{c.get('file', '')}`")
                    parts.append("")
            else:
                parts.append("_未检测到测试文件 / 测试用例_")
                parts.append("")
        elif sec == "dependencies":
            parts.append(_h("外部依赖"))
            parts.append("")
            if deps_data and deps_data.get("count"):
                parts.append(f"共声明 **{deps_data['count']}** 项外部依赖（来源：`pyproject.toml` / `requirements.txt` / `setup.py`）：")
                parts.append("")
                categorized = deps_data.get("categorized", {})
                cat_labels = {
                    "web": "Web / HTTP 服务", "db": "数据存储", "test": "测试",
                    "cli": "CLI / 终端", "data": "数据处理", "ml": "机器学习",
                    "http": "HTTP 客户端", "auth": "认证 / 加密",
                    "config": "配置 / 校验", "async": "异步 / 任务",
                    "serialize": "序列化", "lint": "代码质量", "other": "其它",
                }
                for cat, label in cat_labels.items():
                    items = categorized.get(cat)
                    if not items:
                        continue
                    parts.append(f"- **{label}**: {', '.join('`'+x+'`' for x in items)}")
                parts.append("")
                parts.append("<details><summary>完整依赖清单</summary>")
                parts.append("")
                parts.append("| 包名 | 来源 |")
                parts.append("|------|------|")
                for name, src in deps_data["deps"]:
                    parts.append(f"| `{name}` | `{src}` |")
                parts.append("")
                parts.append("</details>")
            else:
                parts.append("_未在常见位置 (`pyproject.toml`、`requirements*.txt`、`setup.py`) 发现声明依赖_")
            parts.append("")
        elif sec == "recommendations":
            parts.append("## 推荐问答")
            parts.append("")
            parts.append("想深入了解？尝试以下问题（点击直接提问）：")
            parts.append("")
            for q in questions[:6]:
                parts.append(f"- ASK::{q}")
            parts.append("")

    # 抗幻觉:对 LLM 叙述章(overview + 模块详解)做 groundedness 核验,未命中标红 + 附脚注
    try:
        _prose = (overview or "") + "\n" + (module_details or "")
        _dv = _verify_response(_prose, store)
        _facts_s = "本仓: %d 模块 / %d 类 / %d 路由" % (len(modules), len(classes), len(routes))
        _deep = _verify_prose_deep(_prose, [store], llm, extra_evidence=_facts_s)
        if _dv["invalid_refs"]:
            for _r in sorted(set(_dv["invalid_refs"]), key=len, reverse=True):
                parts = [p.replace("`%s`" % _r, "`%s`<sup>⚠️未在图谱中找到</sup>" % _r) for p in parts]
        parts.append("## 正确性核验(叙述层抗幻觉)")
        parts.append("")
        _ti = _dv["tiers"]; _rb = _dv["relationship_unsupported"]
        parts.append("叙述章反引号符号去本仓图谱核验,分档:精确 %d / 唯一 %d / 歧义 %d / 未命中 %d,"
                     "groundedness(存在率)**%.2f** / 可精确定位 **%.2f**(未命中已标红);"
                     "关系级核验 %d 条、未命中 %d 条。图表 / 结构事实不经 LLM、无需核验。"
                     % (_ti["exact"], _ti["unique"], _ti["ambiguous"], _ti["miss"],
                        _dv["groundedness"], _dv["specificity"], _dv["relationship_total"], len(_rb)))
        if _rb:
            parts.append("")
            parts.append("⚠️ 图谱中查不到的关系(疑似编造):%s" %
                         "; ".join("`%s` %s `%s`" % (c["a"], c["rel"], c["b"]) for c in _rb))
        parts.append("")
        parts.append("**深度语义核验(LLM 裁判,非确定、仅离线)**:判了 %d 个符号相关论断、"
                     "标出 %d 条证据不支持。" % (_deep["judged"], len(_deep["unsupported"])))
        for _u in _deep["unsupported"][:6]:
            parts.append("- ⚠️ %s — %s" % (str(_u.get("claim", ""))[:90], str(_u.get("why", ""))[:90]))
        parts.append("")
    except Exception:
        pass

    doc = "\n".join(parts)

    # Cache on the repo (so other endpoints can read it)
    if repo_name in indexed_repos:
        indexed_repos[repo_name]["doc"] = doc
        indexed_repos[repo_name]["recommendations"] = questions[:6]

    finished_at = time.time()
    with _docgen_lock:
        _docgen_tasks[task_id]["status"] = "done"
        _docgen_tasks[task_id]["stage"] = "done"
        _docgen_tasks[task_id]["progress"] = "文档生成完成"
        _docgen_tasks[task_id]["document"] = doc
        _docgen_tasks[task_id]["recommendations"] = questions[:6]
        _docgen_tasks[task_id]["sections"] = requested
        _docgen_tasks[task_id]["template"] = template
        _docgen_tasks[task_id]["finished_at"] = finished_at
    try:
        _docgen_db_update(
            task_id,
            status="done", stage="done", progress="文档生成完成",
            document=doc, recommendations=json.dumps(questions[:6]),
            finished_at=finished_at,
        )
    except Exception:  # noqa: BLE001
        pass



# ===========================================================================
# 多仓文档生成(系统级文档):确定性流水线,事实全从图谱/源码确定性抽取,
# 只 2 次 LLM 写散文(总览 + 协作分析);复用 docgen PG 队列做异步与崩溃重跑。
# repo 字段用 "multi://a,b,c" 哨兵编码,无需改队列表结构。
# ===========================================================================

# 多仓文档的章节键(仅用于状态展示/计数;真正结构在 _run_multi_docgen_task 里)
MULTI_DOC_SECTIONS = [
    "cross_overview", "repo_profiles", "cross_symbols",
    "cross_imports", "cross_arch", "cross_collab",
]


def _scan_cross_imports(repo_dir, other_pkgs):
    """扫一个仓的 .py 源码,找它 import 了哪些"其它仓的顶层包名"。
    确定性、可解释:这是跨仓真实依赖(如 flask 源码 import werkzeug/click)。"""
    found = set()
    if not other_pkgs:
        return found
    pat = re.compile(r"^\s*(?:from|import)\s+([A-Za-z_][A-Za-z0-9_]*)", re.MULTILINE)
    pkgset = set(other_pkgs)
    scanned = 0
    try:
        for p in Path(repo_dir).rglob("*.py"):
            if scanned >= 600:
                break
            # 跳过测试/虚拟环境噪声
            sp_parts = set(p.parts)
            if {".venv", "venv", "site-packages", "node_modules"} & sp_parts:
                continue
            scanned += 1
            try:
                text = p.read_text(encoding="utf-8", errors="replace")
            except Exception:
                continue
            for m in pat.finditer(text):
                top = m.group(1)
                if top in pkgset:
                    found.add(top)
            if found == pkgset:
                break
    except Exception:
        pass
    return found


def gen_cross_repo_arch_plantuml(repos_info, import_pairs, shared_pairs):
    """跨仓架构图:每个仓一个 component,实线箭头=跨仓引用(import),虚线=同名公共符号。"""
    lines = ["@startuml", "!theme plain", "skinparam componentStyle rectangle",
             "left to right direction", 'skinparam component {BackgroundColor #F2FBF5 BorderColor #1a6}']
    alias = {}
    for i, ri in enumerate(repos_info):
        a = "R%d" % i
        alias[ri["name"]] = a
        lines.append('component "%s\\n%d 类 / %d 边" as %s' % (
            ri["pkg"], len(ri["classes"]), ri["n_edges"], a))
    drawn = set()
    for (src, dst) in sorted(import_pairs):
        if src in alias and dst in alias:
            lines.append("%s --> %s : imports" % (alias[src], alias[dst]))
            drawn.add((src, dst))
    for (a_name, b_name), cnt in sorted(shared_pairs.items(), key=lambda x: -x[1]):
        if (a_name, b_name) in drawn or (b_name, a_name) in drawn:
            continue  # 已有 import 实线就不再叠虚线,避免杂乱
        if a_name in alias and b_name in alias:
            lines.append("%s .. %s : 同名x%d" % (alias[a_name], alias[b_name], cnt))
    lines.append("@enduml")
    return "\n".join(lines)


def _run_multi_docgen_task(task_id: str, task: dict):
    """多仓文档生成:加载各仓索引 -> 概况 -> 跨仓同名符号 -> 跨仓引用(扫源码) ->
    跨仓架构图 -> 2 次 LLM 写总览/协作分析 -> 拼 markdown。"""
    import itertools
    from codedoc.agents.llm import ChatMessage, build_llm

    raw = task["repo"]
    repo_names = [r for r in raw[len("multi://"):].split(",") if r]

    def _fail(msg):
        finished_at = time.time()
        with _docgen_lock:
            if task_id in _docgen_tasks:
                _docgen_tasks[task_id]["status"] = "error"
                _docgen_tasks[task_id]["stage"] = "error"
                _docgen_tasks[task_id]["error"] = msg
                _docgen_tasks[task_id]["finished_at"] = finished_at
        try:
            _docgen_db_update(task_id, status="error", stage="error", error=msg,
                              finished_at=finished_at)
        except Exception:
            pass

    _update_stage(task_id, "loading", "加载各仓索引...")
    repos_info = []
    cfg0 = None
    for r in repo_names:
        _update_stage(task_id, "loading", "加载仓索引: %s" % r)
        info = _ensure_repo_indexed(r)
        if not info:
            continue
        cfg0 = cfg0 or info["cfg"]
        store = info["store"]
        path = Path(info.get("path") or (REPOS_DIR / r.split("/")[-1]))

        def _deg(nid, st=store):
            return sum(1 for e in st.edges if e["src"] == nid or e["dst"] == nid)

        modules = [n for n in store.nodes.values() if n["kind"] == "module"]
        classes = [n for n in store.nodes.values() if n["kind"] == "class"]
        routes = [n for n in store.nodes.values() if n["kind"] == "route_handler"]
        funcs = [n for n in store.nodes.values() if n["kind"] in ("function", "method")]
        langs = sorted({n.get("language", "") for n in store.nodes.values() if n.get("language")})
        core = sorted(classes, key=lambda c: -_deg(c["id"]))[:5]
        try:
            eps = extract_entry_points(path, store)
        except Exception:
            eps = []
        repos_info.append({
            "name": r, "pkg": r.split("/")[-1], "store": store, "path": path,
            "langs": langs, "modules": modules, "classes": classes, "routes": routes,
            "funcs": funcs, "core": core, "eps": eps,
            "n_nodes": len(store.nodes), "n_edges": len(store.edges),
        })

    if len(repos_info) < 2:
        _fail("多仓文档至少需要 2 个已索引的仓(当前可用 %d 个)" % len(repos_info))
        return

    llm = build_llm(cfg0)

    # ---- 跨仓同名公共符号(复用多仓 QA Merger 的思路) ----
    _update_stage(task_id, "cross_symbols", "比对跨仓同名符号...")

    _NOISE_NAMES = {"Foo", "Bar", "Baz", "Test", "Tests", "Example", "Dummy",
                    "Mock", "Base", "Item", "Thing", "Obj", "Sample", "Demo", "Child", "Parent"}

    def _pub(n):
        nm = n.get("name") or ""
        # 只比"类":类是有意义的跨仓共享抽象;函数同名多为巧合(foo/main/wrapper),排掉
        return bool(nm) and not nm.startswith("_") and len(nm) >= 3 and \
            n["kind"] == "class" and nm not in _NOISE_NAMES

    name_to_repos = {}
    for ri in repos_info:
        names = {n["name"] for n in ri["store"].nodes.values() if _pub(n)}
        ri["pubnames"] = names
        for nm in names:
            name_to_repos.setdefault(nm, set()).add(ri["name"])
    shared = {nm: sorted(rs) for nm, rs in name_to_repos.items() if len(rs) > 1}
    shared_pairs = {}
    for a, b in itertools.combinations(repos_info, 2):
        c = len(a["pubnames"] & b["pubnames"])
        if c > 0:
            shared_pairs[(a["name"], b["name"])] = c

    # ---- 跨仓真实引用:扫源码 import(强信号、确定性) ----
    _update_stage(task_id, "cross_imports", "扫描跨仓引用(import)...")
    pkg_to_repo = {ri["pkg"]: ri["name"] for ri in repos_info}
    import_pairs = set()
    for ri in repos_info:
        others = [p for p in pkg_to_repo if p != ri["pkg"]]
        for p in _scan_cross_imports(ri["path"], others):
            import_pairs.add((ri["name"], pkg_to_repo[p]))

    # ---- 跨仓架构图 ----
    _update_stage(task_id, "cross_arch", "绘制跨仓架构图...")
    cross_uml = gen_cross_repo_arch_plantuml(repos_info, import_pairs, shared_pairs)

    # ---- LLM 1:系统级总览 ----
    _update_stage(task_id, "cross_overview", "生成跨仓总览...")
    sum_lines = []
    for ri in repos_info:
        sum_lines.append(
            "- %s:语言 %s,%d 模块 / %d 类 / %d 函数 / %d 路由,核心类 %s" % (
                ri["name"], ",".join(ri["langs"]) or "?", len(ri["modules"]),
                len(ri["classes"]), len(ri["funcs"]), len(ri["routes"]),
                ", ".join(c["name"] for c in ri["core"][:3]) or "—"))
    overview = ""
    try:
        overview = llm.chat([
            ChatMessage("system",
                "你是架构文档作者。下面是同一业务线下多个代码仓的结构统计。"
                "写一段 200-300 字的系统级总览:这几个仓各自承担什么、整体像什么系统、"
                "技术栈共性。用 Markdown 段落,不要加标题,不要编造仓里没有的东西。"
                "提到代码里的类名/函数名时用反引号括起来;仓名/包名直接写、不要加反引号。"),
            ChatMessage("user", "\n".join(sum_lines)),
        ], max_tokens=600) or ""
    except Exception as e:
        overview = "_(总览生成失败:%s)_" % e

    # ---- LLM 2:协作 / 依赖分析 ----
    _update_stage(task_id, "cross_collab", "生成跨仓协作分析...")
    rel_lines = []
    for (s, d) in sorted(import_pairs):
        rel_lines.append("- %s 的源码 import 了 %s(跨仓依赖,方向 %s → %s)" % (s, d, s, d))
    for (a, b), c in sorted(shared_pairs.items(), key=lambda x: -x[1])[:15]:
        rel_lines.append("- %s 与 %s 有 %d 个同名公共符号" % (a, b, c))
    if not rel_lines:
        rel_lines.append("- (未发现跨仓 import,也无显著同名符号)")
    collab = ""
    try:
        collab = llm.chat([
            ChatMessage("system",
                "你是架构师。下面是多个仓之间的真实依赖(import)与同名符号统计。"
                "分析:谁依赖谁、分层关系(谁是底座/谁是上层)、可能的协作点;"
                "再给 3 条务实的协作/治理建议。只基于给定事实,不要编造。"
                "提到的类名/函数名用反引号括起来;仓名直接写、不要加反引号。"),
            ChatMessage("user", "\n".join(rel_lines)),
        ], max_tokens=700) or ""
    except Exception as e:
        collab = "_(协作分析生成失败:%s)_" % e

    # ---- 正确性核验:对 2 段 LLM 散文做 groundedness 核验(同单仓抗幻觉) ----
    _update_stage(task_id, "verify", "核验散文引用...")
    _all_stores = [ri["store"] for ri in repos_info]

    def _verify_prose(text):
        try:
            return _verify_response(text or "", _all_stores[0], extra_stores=_all_stores[1:])
        except Exception:
            return {"total_refs": 0, "valid_refs": [], "invalid_refs": [], "groundedness": 1.0}

    def _flag(text, invalid):
        for r in sorted(set(invalid), key=len, reverse=True):
            text = text.replace("`%s`" % r, "`%s`<sup>⚠️未在图谱中找到</sup>" % r)
        return text

    ov_ver = _verify_prose(overview)
    co_ver = _verify_prose(collab)
    _facts = "各仓: " + "; ".join("%s %d类/%d函数/%d路由" % (ri["name"], len(ri["classes"]), len(ri["funcs"]), len(ri["routes"])) for ri in repos_info)
    _facts += " || 跨仓import依赖: " + ("; ".join("%s->%s" % (s, d) for s, d in sorted(import_pairs)) or "无")
    _facts += " || 同名共享类对: " + ("; ".join("%s&%s=%d" % (a.split("/")[-1], b.split("/")[-1], c) for (a, b), c in shared_pairs.items()) or "无")
    _deep = _verify_prose_deep((overview or "") + "\n\n" + (collab or ""), _all_stores, llm, extra_evidence=_facts)
    overview = _flag(overview, ov_ver["invalid_refs"])
    collab = _flag(collab, co_ver["invalid_refs"])

    # ---- 拼装 Markdown ----
    _update_stage(task_id, "finalizing", "拼装文档...")
    P = []
    title = " + ".join(ri["pkg"] for ri in repos_info)
    P.append("# 多仓系统文档:%s" % title)
    P.append("")
    P.append("> 本文档由 codedoc 跨 **%d** 个代码仓自动生成。事实层(各仓统计、同名符号、"
             "跨仓 import、架构图)全部从代码图谱与源码确定性抽取;仅\"总览\"与\"协作分析\""
             "两段为 LLM 基于上述事实撰写。" % len(repos_info))
    P.append("")

    P.append("## 1. 跨仓总览")
    P.append("")
    P.append(overview)
    P.append("")

    P.append("## 2. 各仓概况对比")
    P.append("")
    P.append("| 仓 | 语言 | 模块 | 类 | 函数 | 路由 | 节点 | 关系边 | 核心类(按度数) |")
    P.append("|----|------|------|----|------|------|------|--------|----------------|")
    for ri in repos_info:
        P.append("| `%s` | %s | %d | %d | %d | %d | %d | %d | %s |" % (
            ri["name"], ",".join(ri["langs"]) or "?", len(ri["modules"]), len(ri["classes"]),
            len(ri["funcs"]), len(ri["routes"]), ri["n_nodes"], ri["n_edges"],
            ", ".join("`%s`" % c["name"] for c in ri["core"][:3]) or "—"))
    P.append("")

    P.append("## 3. 跨仓架构图")
    P.append("")
    P.append("实线箭头 = 跨仓真实引用(一个仓源码 import 另一个仓);虚线 = 同名公共符号关联。")
    P.append("")
    P.append("![跨仓架构图](%s)" % plantuml_url(cross_uml))
    P.append("")
    P.append("<details><summary>查看 PlantUML 源码</summary>\n\n```\n" + cross_uml + "\n```\n</details>")
    P.append("")

    P.append("## 4. 跨仓真实引用(import 扫描)")
    P.append("")
    if import_pairs:
        P.append("扫描各仓源码 `import` 语句,命中其它仓顶层包名的依赖如下(这是真实的跨仓依赖方向):")
        P.append("")
        P.append("| 引用方(importer) | 被引用方 |")
        P.append("|------------------|----------|")
        for (s, d) in sorted(import_pairs):
            P.append("| `%s` | `%s` |" % (s, d))
    else:
        P.append("_未在源码中发现这几个仓之间的直接 import 依赖(可能彼此独立,或通过第三方包间接关联)。_")
    P.append("")

    P.append("## 5. 跨仓同名公共类")
    P.append("")
    if shared:
        items = sorted(shared.items(), key=lambda kv: (-len(kv[1]), kv[0]))[:40]
        P.append("以下公共**类**(排除下划线私有与测试桩名)在 **多个仓**都出现,可能是 re-export、"
                 "共享接口或同名实现(共 %d 个,展示前 %d 个):" % (len(shared), len(items)))
        P.append("")
        P.append("| 符号 | 出现于 |")
        P.append("|------|--------|")
        for nm, rs in items:
            P.append("| `%s` | %s |" % (nm, ", ".join("`%s`" % x for x in rs)))
    else:
        P.append("_未发现跨仓同名公共符号。_")
    P.append("")

    P.append("## 6. 各仓核心入口")
    P.append("")
    for ri in repos_info:
        P.append("### `%s`" % ri["name"])
        P.append("")
        if ri["eps"]:
            for ep in ri["eps"][:8]:
                label = ep.get("name") or ep.get("qualified_name") or ep.get("file") or "?"
                P.append("- [%s] `%s`" % (ep.get("type", ""), label))
        else:
            P.append("- _未检测到显式入口点_")
        P.append("")

    P.append("## 7. 跨仓协作与依赖分析")
    P.append("")
    P.append(collab)
    P.append("")

    P.append("## 8. 正确性核验(散文层抗幻觉)")
    P.append("")
    P.append("对两段 LLM 散文做确定性 grounding 核验。①**符号级**:反引号符号去 %d 仓图谱并集查,"
             "分『精确/唯一/歧义/未命中』四档:groundedness=存在率(精确+唯一+歧义,即没编造),"
             "精确+唯一=可精确定位,歧义=一名多义需限定,未命中=编造、已正文标红。②**关系级**:"
             "文本里『A 继承/调用 B』去查图谱这条边在不在,"
             "未命中即编造关系。事实层(表格/import/架构图)不经 LLM、无需核验。" % len(repos_info))
    P.append("")
    P.append("| 段落 | 精确 | 唯一 | 歧义 | 未命中 | groundedness | 关系核验(命中/总) |")
    P.append("|------|------|------|------|--------|--------------|------------------|")

    def _vrow(label, v):
        ti = v["tiers"]
        rel = "%d/%d" % (v["relationship_total"] - len(v["relationship_unsupported"]), v["relationship_total"])
        return "| %s | %d | %d | %d | %d | %.2f | %s |" % (
            label, ti["exact"], ti["unique"], ti["ambiguous"], ti["miss"], v["groundedness"], rel)

    P.append(_vrow("跨仓总览", ov_ver))
    P.append(_vrow("协作分析", co_ver))
    _bad = sorted(set(ov_ver["invalid_refs"]) | set(co_ver["invalid_refs"]))
    _relbad = ov_ver["relationship_unsupported"] + co_ver["relationship_unsupported"]
    if _bad:
        P.append("")
        P.append("⚠️ 符号未命中、已标红:%s" % ", ".join("`%s`" % x for x in _bad))
    if _relbad:
        P.append("")
        P.append("⚠️ 图谱中查不到的关系(疑似编造):%s" %
                 "; ".join("`%s` %s `%s`" % (c["a"], c["rel"], c["b"]) for c in _relbad))
    P.append("")
    P.append("**③ 深度语义核验(LLM 裁判,非确定、仅离线 docgen)**:取上面引用符号的真实 "
             "docstring 当证据,判了 %d 个符号相关论断,标出 %d 条证据不支持。" %
             (_deep["judged"], len(_deep["unsupported"])))
    for _u in _deep["unsupported"][:6]:
        P.append("- ⚠️ %s — %s" % (str(_u.get("claim", ""))[:90], str(_u.get("why", ""))[:90]))
    P.append("")

    recs = [
        "%s 和 %s 之间是怎么协作的?" % (repos_info[0]["pkg"], repos_info[1]["pkg"]),
        "这几个仓里哪个是最底层的依赖?",
        "跨仓共享的核心抽象有哪些?",
    ]
    P.append("## 推荐问答")
    P.append("")
    for q in recs:
        P.append("- ASK::%s" % q)
    P.append("")

    doc = "\n".join(P)

    finished_at = time.time()
    with _docgen_lock:
        _docgen_tasks[task_id]["status"] = "done"
        _docgen_tasks[task_id]["stage"] = "done"
        _docgen_tasks[task_id]["progress"] = "多仓文档生成完成"
        _docgen_tasks[task_id]["document"] = doc
        _docgen_tasks[task_id]["recommendations"] = recs
        _docgen_tasks[task_id]["finished_at"] = finished_at
    try:
        _docgen_db_update(task_id, status="done", stage="done",
                          progress="多仓文档生成完成", document=doc,
                          recommendations=json.dumps(recs), finished_at=finished_at)
    except Exception:
        pass


@app.post("/api/v1/docgen")
def submit_docgen(req: DocgenRequest, user: dict = Depends(get_current_user)):
    """Submit a doc-gen job. Returns immediately with a task_id; poll status endpoint."""
    _multi = [r for r in (req.repos or []) if r]
    if len(_multi) > 1:
        repo_field = "multi://" + ",".join(_multi)
        requested = list(MULTI_DOC_SECTIONS)
    else:
        repo_field = req.repo
        # Resolve which sections will be rendered
        if req.sections:
            requested = [s for s in req.sections if s in SECTION_LABELS]
        else:
            tmpl = DOC_TEMPLATES.get(req.template) or DOC_TEMPLATES["default"]
            requested = list(tmpl["sections"])
        if not requested:
            requested = list(DOC_TEMPLATES["default"]["sections"])

    task_id = str(uuid.uuid4())
    user_id = int(user["sub"])
    submitted_at = time.time()

    # Persist first so all uvicorn workers can see it
    try:
        _docgen_db_insert(task_id, user_id, repo_field, requested, req.template, submitted_at)
    except Exception as e:  # noqa: BLE001
        raise HTTPException(500, f"无法创建任务: {e}")

    with _docgen_lock:
        _docgen_tasks[task_id] = {
            "task_id": task_id,
            "user_id": user["sub"],
            "repo": repo_field,
            "sections": requested,
            "template": req.template,
            "status": "queued",
            "stage": "queued",
            "progress": "",
            "position": 0,
            "submitted_at": submitted_at,
        }
    _start_docgen_worker()   # 确保本 worker 领取循环在跑;任意 worker 都会经 SKIP LOCKED 抢到此任务

    # Compute initial position from the global queue
    try:
        position = _docgen_db_queue_position(task_id, submitted_at)
    except Exception:  # noqa: BLE001
        position = 1
    queue_total = _docgen_db_queue_total() or position
    return {"task_id": task_id, "status": "queued", "position": position,
            "queue_total": queue_total, "sections": requested, "template": req.template}


@app.get("/api/v1/docgen/tasks")
def list_docgen_tasks(user: dict = Depends(get_current_user)):
    """List all docgen tasks for the current user, newest first."""
    from codedoc import docgen_queue as _dq
    rows = _dq.list_for_user(int(user["sub"]), 50)

    tasks: list[dict] = []
    for r in rows:
        d = dict(r)
        # Compute live queue position for queued tasks
        if d["status"] == "queued":
            try:
                d["position"] = _docgen_db_queue_position(d["task_id"], d["submitted_at"])
            except Exception:  # noqa: BLE001
                d["position"] = 0
        else:
            d["position"] = 0
        tasks.append(d)
    return {"tasks": tasks}


@app.get("/api/v1/docgen/{task_id}/document")
def get_docgen_document(task_id: str, user: dict = Depends(get_current_user)):
    """Fetch the rendered document for a completed task (from sqlite cache)."""
    from codedoc import docgen_queue as _dq
    row = _dq.fetch(task_id)
    if not row or int(row.get("user_id")) != int(user["sub"]):
        raise HTTPException(404, "任务不存在")
    d = dict(row)
    if d["status"] != "done":
        return {
            "status": d["status"],
            "stage": d.get("stage") or "",
            "progress": d.get("progress") or "",
            "error": d.get("error") or "",
            "document": None,
            "recommendations": [],
        }
    try:
        recs = json.loads(d["recommendations"]) if d.get("recommendations") else []
    except Exception:  # noqa: BLE001
        recs = []
    return {
        "status": "done",
        "stage": "done",
        "document": d.get("document") or "",
        "recommendations": recs,
        "repo": d.get("repo"),
        "template": d.get("template"),
        "finished_at": d.get("finished_at"),
    }


@app.get("/api/v1/docgen/{task_id}/status")
def docgen_status(task_id: str, user: dict = Depends(get_current_user)):
    """Poll the status of an async docgen job. Reads from sqlite for cross-worker consistency."""
    # Always read latest from sqlite (single source of truth across workers)
    t = _docgen_db_fetch(task_id)
    if not t:
        # Fall back to in-process cache (race: task just submitted, db row may
        # not be visible yet — unlikely with WAL but defensive)
        with _docgen_lock:
            local = _docgen_tasks.get(task_id)
        if not local:
            raise HTTPException(404, "任务不存在")
        t = {
            "task_id": task_id,
            "user_id": int(local["user_id"]),
            "status": local["status"],
            "stage": local.get("stage", ""),
            "progress": local.get("progress", ""),
            "submitted_at": local.get("submitted_at", time.time()),
            "document": local.get("document"),
            "recommendations": local.get("recommendations", []),
            "sections": local.get("sections", []),
            "template": local.get("template", ""),
            "error": local.get("error", ""),
        }

    if int(t["user_id"]) != int(user["sub"]):
        raise HTTPException(403, "无权访问此任务")

    status = t["status"]
    position = 0
    progress = t.get("progress", "")
    if status == "queued":
        try:
            position = _docgen_db_queue_position(task_id, t["submitted_at"])
        except Exception:  # noqa: BLE001
            position = 1
        progress = f"排队中（第 {position} 位）"

    try:
        queue_total = _docgen_db_queue_total()
    except Exception:  # noqa: BLE001
        queue_total = 0

    return {
        "task_id": task_id,
        "status": status,
        "stage": t.get("stage", ""),
        "progress": progress,
        "position": position,
        "queue_total": queue_total,
        "document": t.get("document") if status == "done" else None,
        "recommendations": t.get("recommendations", []) if status == "done" else [],
        "sections": t.get("sections", []),
        "template": t.get("template", ""),
        "error": t.get("error", ""),
    }


# ---------------------------------------------------------------------------
# Health (no auth)
# ---------------------------------------------------------------------------

@app.get("/health")
def health():
    return {"status": "ok", "service": "codedoc", "version": "2.0.0"}

# ---------------------------------------------------------------------------
# Vue 3 SPA UI
# ---------------------------------------------------------------------------

FRONTEND_HTML = r"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>CodeDoc - 代码文档平台</title>
<script src="https://cdn.jsdelivr.net/npm/vue@3.4.27/dist/vue.global.prod.js"></script>
<script src="https://cdn.jsdelivr.net/npm/marked@9/marked.min.js"></script>
<style>
* { margin: 0; padding: 0; box-sizing: border-box; }
:root {
  --bg: #ffffff;
  --bg-soft: #fbfbfd;
  --panel: #ffffff;
  --border: #d2d2d7;
  --border-strong: #d2d2d7;
  --text: #1d1d1f;
  --text-2: #1d1d1f;
  --text-3: #86868b;
  --accent: #0071e3;
  --accent-hover: #0077ed;
  --hover-bg: #f5f5f7;
  --input-bg: #f5f5f7;
  --success: #34c759;
  --warn: #ff9f0a;
  --danger: #ff3b30;
  --shadow-card: 0 4px 14px rgba(0,0,0,0.05);
  --shadow-card-hover: 0 8px 24px rgba(0,0,0,0.08);
  --sidebar: 240px;
}
html, body {
  font-family: -apple-system, BlinkMacSystemFont, 'Inter', 'SF Pro Text', 'PingFang SC', 'Hiragino Sans GB', 'Microsoft YaHei', sans-serif;
  color: var(--text);
  background: #ffffff;
  -webkit-font-smoothing: antialiased;
  min-height: 100vh;
}
body { background: #ffffff; }
[v-cloak] { display: none !important; }
a { color: inherit; text-decoration: none; }
button { font-family: inherit; cursor: pointer; border: none; background: none; color: inherit; }
input, textarea { font-family: inherit; }

/* Scrollbar */
::-webkit-scrollbar { width: 8px; height: 8px; }
::-webkit-scrollbar-thumb { background: #d2d2d7; border-radius: 6px; }
::-webkit-scrollbar-thumb:hover { background: #b0b0b5; }

/* Toast */
.toast-wrap { position: fixed; top: 18px; right: 18px; z-index: 9999; display: flex; flex-direction: column; gap: 10px; pointer-events: none; }
.toast {
  padding: 12px 18px;
  border-radius: 12px;
  font-size: 14px;
  background: #1d1d1f;
  color: #fff;
  box-shadow: 0 4px 14px rgba(0,0,0,0.15);
  animation: slideIn .25s ease;
  pointer-events: auto;
  max-width: 360px;
  font-weight: 400;
}
.toast.success { background: #34c759; color: #fff; }
.toast.error   { background: #ff3b30; color: #fff; }
.toast.info    { background: #1d1d1f; color: #fff; }
@keyframes slideIn { from { transform: translateX(20px); opacity: 0; } to { transform: translateX(0); opacity: 1; } }

/* Auth */
@keyframes appleAuthFadeIn { from { opacity: 0; transform: translateY(14px); } to { opacity: 1; transform: translateY(0); } }
@keyframes appleSpin { to { transform: rotate(360deg); } }
.apple-auth-page {
  min-height: 100vh;
  display: flex; align-items: center; justify-content: center;
  background: #fbfbfd;
  position: relative;
  overflow: hidden;
}
.apple-auth-card {
  width: 100%;
  max-width: 420px;
  background: #ffffff;
  border: 1px solid #d2d2d7;
  border-radius: 18px;
  padding: 44px 40px;
  box-shadow: 0 4px 14px rgba(0,0,0,0.05);
  animation: appleAuthFadeIn 0.35s ease;
  position: relative;
  z-index: 1;
}
.apple-auth-logo {
  text-align: center;
  font-size: 32px;
  font-weight: 600;
  letter-spacing: -0.015em;
  margin-bottom: 8px;
  color: #1d1d1f;
}
.apple-auth-subtitle { text-align: center; font-size: 15px; color: #86868b; margin-bottom: 34px; font-weight: 400; }
.apple-auth-form { display: flex; flex-direction: column; gap: 12px; }
.apple-input {
  width: 100%;
  height: 48px;
  background: #f5f5f7;
  border: none;
  border-radius: 12px;
  padding: 0 16px;
  font-size: 15px;
  color: #1d1d1f;
  outline: none;
  transition: box-shadow 0.15s;
}
.apple-input::placeholder { color: #86868b; }
.apple-input:focus {
  box-shadow: 0 0 0 4px rgba(0,113,227,0.2);
}
.apple-btn {
  width: 100%;
  height: 48px;
  background: #0071e3;
  color: #fff;
  border: none;
  border-radius: 980px;
  font-size: 17px;
  font-weight: 400;
  cursor: pointer;
  transition: background 0.15s;
  display: inline-flex; align-items: center; justify-content: center; text-align: center; gap: 8px;
  margin-top: 8px;
}
.apple-btn:hover { background: #0077ed; }
.apple-btn:disabled { opacity: 0.4; cursor: not-allowed; }
.apple-spinner {
  width: 18px; height: 18px;
  border: 2px solid rgba(255,255,255,0.35);
  border-top-color: #fff;
  border-radius: 50%;
  animation: appleSpin 0.7s linear infinite;
  display: inline-block;
}
.apple-auth-toggle { text-align: center; margin-top: 8px; }
.apple-auth-toggle span {
  color: #0071e3;
  font-size: 14px;
  font-weight: 400;
  cursor: pointer;
}
.apple-auth-toggle span:hover { text-decoration: underline; }
.apple-auth-error { color: #ff3b30; font-size: 13px; margin-top: -4px; padding: 0 4px; font-weight: 400; }

/* Layout */
.shell { display: grid; grid-template-columns: var(--sidebar) 1fr; min-height: 100vh; }
.sidebar {
  background: #ffffff;
  border-right: 1px solid #d2d2d7;
  display: flex; flex-direction: column;
  height: 100vh; position: sticky; top: 0;
  overflow-y: auto;
}
.sidebar-head { padding: 20px 18px 16px; display: flex; align-items: center; gap: 10px; border-bottom: 1px solid #d2d2d7; }
.brand-mark {
  width: 32px; height: 32px;
  border-radius: 10px;
  background: #0071e3;
  display: flex; align-items: center; justify-content: center;
  color: #fff; font-weight: 600; font-size: 15px; flex-shrink: 0;
}
.brand-title {
  font-size: 16px; font-weight: 600; letter-spacing: -0.2px;
  color: #1d1d1f;
}
.brand-sub { font-size: 11px; color: #86868b; margin-top: 1px; letter-spacing: 0.2px; font-weight: 400; }
.sidebar-section-title { padding: 20px 18px 8px; font-size: 11px; font-weight: 600; color: #86868b; letter-spacing: 1px; text-transform: uppercase; }
.repo-tree { padding: 0 10px; flex: 1; }
.repo-item { border-radius: 10px; margin: 2px 0; transition: background .12s; }
.repo-item:hover { background: #f5f5f7; }
.repo-item.active {
  background: #f5f5f7;
}
.repo-item.active .repo-row,
.repo-item.active .repo-row .name { color: #1d1d1f !important; }
.repo-item.active .repo-row svg { color: #1d1d1f !important; opacity: 0.7 !important; }
.repo-row { display: flex; align-items: center; gap: 8px; padding: 8px 11px; font-size: 13px; cursor: pointer; color: #1d1d1f; }
.repo-row .name { flex: 1; overflow: hidden; text-overflow: ellipsis; white-space: nowrap; font-weight: 500; }
.repo-row .status-dot { width: 7px; height: 7px; border-radius: 50%; flex-shrink: 0; }
.repo-row .status-dot.ready { background: #34c759; }
.repo-row .status-dot.cloning, .repo-row .status-dot.indexing { background: #0071e3; animation: pulse 1.5s infinite; }
.repo-row .status-dot.error { background: #ff3b30; }
.repo-row .status-dot.pending { background: #c7c7cc; }
@keyframes pulse { 0%, 100% { opacity: 1; } 50% { opacity: 0.4; } }
.repo-children { padding: 2px 0 6px 32px; }
.repo-children .child { display: flex; align-items: center; gap: 6px; padding: 6px 10px; border-radius: 8px; font-size: 12.5px; color: #1d1d1f; cursor: pointer; transition: background .1s; }
.repo-children .child:hover { background: #f5f5f7; }
.repo-children .child.active {
  color: #0071e3;
  background: #f5f5f7;
  font-weight: 500;
}
.repo-children .child .ic { width: 14px; height: 14px; opacity: .8; }
.sidebar-add { padding: 14px; border-top: 1px solid #d2d2d7; }

/* Top-level sidebar nav (above repo tree) */
.sidebar-nav { padding: 6px 10px 0; }
.sidebar-nav .nav-link {
  display: flex; align-items: center; gap: 8px;
  padding: 8px 11px; border-radius: 10px;
  font-size: 13px; font-weight: 500; color: #1d1d1f;
  cursor: pointer; transition: background .12s;
}
.sidebar-nav .nav-link:hover { background: #f5f5f7; }
.sidebar-nav .nav-link.active { background: #f5f5f7; color: #0071e3; }
.sidebar-nav .nav-link.active svg { color: #0071e3 !important; opacity: 1 !important; }
.task-badge {
  background: #ff3b30; color: #fff; padding: 1px 7px;
  border-radius: 980px; font-size: 10px; font-weight: 600;
  min-width: 18px; text-align: center;
}

/* 我的任务 view cards */
.tasks-list { display: flex; flex-direction: column; gap: 12px; max-width: 880px; }
.task-card {
  background: #fff; border: 1px solid #d2d2d7; border-radius: 12px;
  padding: 16px 18px; transition: box-shadow .15s, border-color .15s;
}
.task-card:hover { box-shadow: 0 4px 14px rgba(0,0,0,0.06); border-color: #c7c7cc; }
.task-card-top { display: flex; justify-content: space-between; align-items: center; margin-bottom: 6px; gap: 10px; }
.task-repo { font-weight: 600; font-size: 14px; color: #1d1d1f; overflow: hidden; text-overflow: ellipsis; white-space: nowrap; }
.task-status { padding: 2px 10px; border-radius: 980px; font-size: 11px; font-weight: 500; flex-shrink: 0; }
.status-queued { background: #f0f0f5; color: #424245; }
.status-running { background: #e6f0ff; color: #0040a3; }
.status-done { background: #e8f7ee; color: #1d6b3a; }
.status-error { background: #ffe9e7; color: #8a2a1f; }
.task-meta { font-size: 12px; color: #86868b; display: flex; flex-wrap: wrap; gap: 6px; margin-bottom: 4px; }
.task-progress { font-size: 12px; color: #424245; padding: 6px 10px; background: #f5f5f7; border-radius: 8px; margin-top: 6px; }
.task-error { font-size: 12px; color: #8a2a1f; padding: 6px 10px; background: #ffe9e7; border-radius: 8px; margin-top: 6px; }
.task-actions { margin-top: 10px; }
.btn-primary-pill {
  background: #0071e3; color: #fff; border: none;
  padding: 8px 18px; border-radius: 980px; font-size: 13px;
  font-weight: 500; cursor: pointer;
  display: inline-flex; align-items: center; justify-content: center;
  transition: background .15s;
}
.btn-primary-pill:hover { background: #0077ed; }
.btn-add-repo {
  width: 100%;
  display: inline-flex; align-items: center; justify-content: center; text-align: center; gap: 6px;
  padding: 10px 22px;
  border-radius: 980px;
  background: #0071e3;
  color: #fff;
  font-size: 14px;
  font-weight: 400;
  transition: background .15s;
}
.btn-add-repo:hover { background: #0077ed; }
.sidebar-foot { padding: 14px 18px; border-top: 1px solid #d2d2d7; display: flex; align-items: center; gap: 10px; }
.avatar {
  width: 34px; height: 34px;
  border-radius: 50%;
  background: #f5f5f7;
  color: #1d1d1f;
  border: 1px solid #d2d2d7;
  display: flex; align-items: center; justify-content: center;
  font-weight: 600; font-size: 13px;
  flex-shrink: 0;
}
.user-info { flex: 1; min-width: 0; }
.user-name { font-size: 13px; font-weight: 500; overflow: hidden; text-overflow: ellipsis; white-space: nowrap; color: #1d1d1f; }
.user-sub { font-size: 11px; color: #86868b; font-weight: 400; }
.btn-icon {
  padding: 6px;
  border-radius: 8px;
  color: #86868b;
  transition: all .15s;
  display: inline-flex; align-items: center; justify-content: center; text-align: center;
}
.btn-icon:hover { background: #f5f5f7; color: #1d1d1f; }

/* Main pane */
.main { display: flex; flex-direction: column; height: 100vh; overflow: hidden; }
.topbar {
  height: 52px;
  border-bottom: 1px solid #d2d2d7;
  background: #ffffff;
  display: flex; align-items: center; padding: 0 24px; gap: 14px;
  position: sticky; top: 0; z-index: 10;
}
.topbar .crumb { font-size: 13px; color: #86868b; display: flex; align-items: center; gap: 6px; font-weight: 400; }
.topbar .crumb b { color: #1d1d1f; font-weight: 600; }
.topbar .topbar-actions { margin-left: auto; display: flex; align-items: center; gap: 10px; }
.topbar a.platform-link { font-size: 13px; color: #1d1d1f; padding: 6px 12px; border-radius: 8px; transition: background .15s; font-weight: 400; }
.topbar a.platform-link:hover { background: #f5f5f7; color: #0071e3; }
.content { flex: 1; overflow-y: auto; padding: 30px 32px; background: #ffffff; }

/* Empty state */
.empty { padding: 80px 20px; text-align: center; color: #86868b; }
.empty-icon { font-size: 54px; margin-bottom: 14px; opacity: .7; }
.empty-title { font-size: 18px; font-weight: 600; color: #1d1d1f; margin-bottom: 6px; }
.empty-sub { font-size: 14px; color: #86868b; }

/* Repo overview cards */
.repo-grid { display: grid; grid-template-columns: repeat(auto-fill, minmax(320px, 1fr)); gap: 20px; }
.repo-card {
  background: #ffffff;
  border: 1px solid #d2d2d7;
  border-radius: 18px;
  padding: 22px;
  transition: transform .2s, box-shadow .2s;
  cursor: pointer;
  position: relative;
}
.repo-card:hover {
  transform: translateY(-2px);
  box-shadow: 0 4px 14px rgba(0,0,0,0.05);
}
.repo-card-head { display: flex; align-items: flex-start; gap: 12px; margin-bottom: 14px; }
.repo-avatar {
  width: 44px; height: 44px;
  border-radius: 14px;
  background: #0071e3;
  display: flex; align-items: center; justify-content: center;
  color: #fff; font-size: 18px; font-weight: 600;
  flex-shrink: 0;
}
.repo-meta { flex: 1; min-width: 0; }
.repo-card-name { font-size: 15px; font-weight: 600; overflow: hidden; text-overflow: ellipsis; white-space: nowrap; letter-spacing: -0.2px; color: #1d1d1f; }
.repo-card-org { font-size: 12px; color: #86868b; margin-top: 2px; overflow: hidden; text-overflow: ellipsis; white-space: nowrap; font-weight: 400; }
.repo-card-stats { display: flex; gap: 10px; font-size: 12px; color: #1d1d1f; margin-bottom: 12px; }
.stat-pill {
  display: inline-flex; align-items: center; gap: 5px;
  padding: 4px 10px;
  border-radius: 980px;
  background: #f5f5f7;
  color: #1d1d1f;
  font-weight: 400;
}
.repo-status-badge {
  display: inline-flex; align-items: center; gap: 5px;
  padding: 3px 11px;
  border-radius: 980px;
  font-size: 11.5px;
  font-weight: 600;
  letter-spacing: 0.2px;
}
.repo-status-badge.ready { background: rgba(52, 199, 89, 0.15); color: #248a3d; }
.repo-status-badge.cloning, .repo-status-badge.indexing { background: rgba(0, 113, 227, 0.12); color: #0071e3; }
.repo-status-badge.error { background: rgba(255, 59, 48, 0.12); color: #d70015; }
.repo-status-badge.pending { background: #f5f5f7; color: #86868b; }
.repo-card-foot { display: flex; align-items: center; justify-content: space-between; font-size: 12px; color: #86868b; margin-top: 12px; font-weight: 400; }
.repo-card-progress {
  font-size: 12px;
  margin-top: 8px;
  font-weight: 500;
  color: #0071e3;
}

/* Modal */
.modal-mask {
  position: fixed; inset: 0;
  background: rgba(0, 0, 0, 0.4);
  backdrop-filter: blur(6px);
  -webkit-backdrop-filter: blur(6px);
  display: flex; align-items: center; justify-content: center;
  z-index: 200;
  animation: fadeIn .2s ease;
}
@keyframes fadeIn { from { opacity: 0; } to { opacity: 1; } }
.modal-card {
  width: 540px; max-width: 90%;
  background: #ffffff;
  border-radius: 18px;
  padding: 30px;
  box-shadow: 0 20px 50px rgba(0,0,0,0.15);
  border: 1px solid #d2d2d7;
}
.modal-title {
  font-size: 22px; font-weight: 600; margin-bottom: 6px; letter-spacing: -0.3px;
  color: #1d1d1f;
}
.modal-sub { font-size: 14px; color: #86868b; margin-bottom: 22px; font-weight: 400; }
.modal-input {
  width: 100%; height: 44px;
  border: none;
  border-radius: 12px;
  padding: 0 14px;
  font-size: 14px;
  background: #f5f5f7;
  outline: none;
  transition: box-shadow .15s;
  color: #1d1d1f;
}
.modal-input:focus { box-shadow: 0 0 0 4px rgba(0,113,227,0.2); }
.modal-actions { display: flex; justify-content: flex-end; gap: 10px; margin-top: 20px; }
.btn-secondary {
  padding: 10px 22px;
  border-radius: 980px;
  border: 1px solid #0071e3;
  background: transparent;
  font-size: 14px;
  font-weight: 400;
  color: #0071e3;
  transition: background .15s;
  display: inline-flex; align-items: center; justify-content: center; text-align: center; gap: 6px;
}
.btn-secondary:hover { background: rgba(0,113,227,0.05); }
.btn-primary {
  padding: 10px 22px;
  border-radius: 980px;
  background: #0071e3;
  color: #fff;
  border: none;
  font-size: 14px;
  font-weight: 400;
  transition: background .15s;
  display: inline-flex; align-items: center; justify-content: center; text-align: center; gap: 6px;
}
.btn-primary:hover { background: #0077ed; }
.btn-primary:disabled { opacity: .4; cursor: not-allowed; }
.btn-danger {
  padding: 10px 22px;
  border-radius: 980px;
  border: 1px solid #ff3b30;
  color: #ff3b30;
  background: transparent;
  font-size: 14px;
  font-weight: 400;
  transition: background .15s;
  display: inline-flex; align-items: center; justify-content: center; text-align: center; gap: 6px;
}
.btn-danger:hover { background: rgba(255,59,48,0.06); }

/* Add modal tabs + dropzone */
.add-tabs {
  display: inline-flex;
  background: #f5f5f7;
  border-radius: 12px;
  padding: 4px;
  margin-bottom: 18px;
}
.add-tab {
  padding: 8px 18px;
  border-radius: 9px;
  font-size: 13px;
  font-weight: 500;
  color: #1d1d1f;
  transition: all .15s;
  display: inline-flex; align-items: center; justify-content: center; text-align: center;
}
.add-tab.active {
  background: #ffffff;
  color: #1d1d1f;
  box-shadow: 0 1px 3px rgba(0,0,0,0.05);
}
.drop-zone {
  border: 2px dashed #d2d2d7;
  border-radius: 14px;
  background: #fbfbfd;
  padding: 34px 20px;
  text-align: center;
  cursor: pointer;
  transition: all .2s;
  display: flex; flex-direction: column; align-items: center; gap: 8px;
  min-height: 170px; justify-content: center;
}
.drop-zone:hover { background: #f5f5f7; border-color: #0071e3; }
.drop-zone.dragging { background: #f5f5f7; border-color: #0071e3; }
.drop-zone.has-file {
  border-color: #0071e3;
  background: #ffffff;
  border-style: solid;
  cursor: default;
}
.drop-zone .drop-icon { color: #0071e3; }
.drop-zone .drop-text { font-size: 14px; color: #1d1d1f; font-weight: 500; }
.drop-zone .drop-hint { font-size: 12px; color: #86868b; }
.picked-file { display: flex; flex-direction: column; align-items: center; gap: 6px; }
.picked-name { font-size: 14px; font-weight: 600; color: #1d1d1f; word-break: break-all; max-width: 100%; }
.picked-size { font-size: 12px; color: #86868b; }
.picked-clear {
  margin-top: 6px;
  font-size: 12px;
  color: #ff3b30;
  background: transparent;
  padding: 5px 12px;
  border: 1px solid #ff3b30;
  border-radius: 980px;
  transition: background .15s;
}
.picked-clear:hover { background: rgba(255,59,48,0.06); }

/* Detail header */
.detail-head { padding-bottom: 22px; border-bottom: 1px solid #d2d2d7; margin-bottom: 26px; }
.detail-title-row { display: flex; align-items: center; gap: 14px; margin-bottom: 8px; }
.detail-title {
  font-size: 28px; font-weight: 600; letter-spacing: -0.5px;
  color: #1d1d1f;
}
.detail-sub { font-size: 14px; color: #86868b; font-weight: 400; }
.detail-stats-row { display: flex; align-items: center; gap: 14px; margin-top: 16px; flex-wrap: wrap; }
.stat-block {
  display: flex; align-items: center; gap: 8px;
  padding: 8px 14px;
  background: #ffffff;
  border: 1px solid #d2d2d7;
  border-radius: 12px;
  font-size: 13px;
}
.stat-block .num {
  font-size: 16px; font-weight: 600;
  color: #1d1d1f;
}
.stat-block .label { color: #86868b; font-weight: 400; }

/* Doc */
.doc-pane {
  background: #ffffff;
  border: 1px solid #d2d2d7;
  border-radius: 18px;
  overflow: hidden;
}
.doc-toolbar { padding: 16px 24px; border-bottom: 1px solid #d2d2d7; display: flex; align-items: center; justify-content: space-between; }
.doc-toolbar-title { font-size: 14px; font-weight: 600; color: #1d1d1f; }
.doc-body { padding: 30px 34px; min-height: 320px; font-size: 15px; line-height: 1.6; color: #1d1d1f; }
.doc-body h1, .doc-body h2 {
  font-size: 22px; font-weight: 600; margin: 26px 0 14px; padding-bottom: 10px;
  border-bottom: 1px solid #d2d2d7; letter-spacing: -0.3px;
  color: #1d1d1f;
}
.doc-body h2:first-child { margin-top: 0; }
.doc-body h3 { font-size: 17px; font-weight: 600; color: #1d1d1f; margin: 20px 0 8px; }
.doc-body h4 { font-size: 14px; font-weight: 600; color: #1d1d1f; margin: 14px 0 6px; }
.doc-body p { margin: 8px 0; line-height: 1.6; color: #1d1d1f; }
.doc-body ul, .doc-body ol { margin: 8px 0 12px 22px; color: #1d1d1f; }
.doc-body li { margin: 4px 0; line-height: 1.6; }
.doc-body code { background: #f5f5f7; padding: 2px 7px; border-radius: 5px; font-size: 13px; color: #1d1d1f; font-family: 'SF Mono', 'Monaco', monospace; font-weight: 400; }
.doc-body pre {
  background: #1d1d1f;
  color: #f5f5f7;
  padding: 18px 20px;
  border-radius: 12px;
  overflow-x: auto;
  font-family: 'SF Mono', 'Monaco', monospace;
  font-size: 13px;
  line-height: 1.6;
  margin: 14px 0;
}
.doc-body pre code { background: none; color: inherit; padding: 0; font-size: inherit; font-weight: normal; }
.doc-body table { width: 100%; border-collapse: collapse; margin: 14px 0; font-size: 13px; }
.doc-body th, .doc-body td { border: 1px solid #d2d2d7; padding: 9px 13px; text-align: left; }
.doc-body th { background: #f5f5f7; font-weight: 600; color: #1d1d1f; }
.doc-body blockquote {
  border-left: 3px solid #0071e3;
  background: #fbfbfd;
  padding: 12px 18px;
  border-radius: 0 10px 10px 0;
  color: #1d1d1f;
  margin: 14px 0;
}
.diagram {
  background: #1d1d1f;
  color: #f5f5f7;
  padding: 20px;
  border-radius: 12px;
  font-family: 'SF Mono', 'Monaco', monospace;
  font-size: 13px;
  line-height: 1.7;
  margin: 14px 0;
  white-space: pre;
  overflow-x: auto;
}
.diagram-label { font-size: 11px; color: #86868b; letter-spacing: 1px; margin-bottom: 6px; text-transform: uppercase; font-weight: 500; }

.ask-btn {
  background: #0071e3; color: white; border: none;
  padding: 6px 14px; border-radius: 980px; font-size: 13px;
  cursor: pointer; margin-right: 8px;
  display: inline-flex; align-items: center; justify-content: center;
}
.ask-btn:hover { background: #0077ed; }
.ask-item { padding: 8px 0; border-bottom: 1px solid #f5f5f7; }
.doc-body img {
  max-width: 100%; display: block; margin: 16px auto;
  border: 1px solid #d2d2d7; border-radius: 12px; padding: 12px;
  background: #fff;
}
.doc-body details {
  background: #f5f5f7; padding: 10px 14px; border-radius: 8px;
  margin: 8px 0 16px; font-size: 13px;
}
.doc-body details summary {
  cursor: pointer; color: #86868b; font-size: 12.5px; font-weight: 500;
  letter-spacing: 0.2px;
}
.doc-body details[open] summary { margin-bottom: 8px; }
.doc-body details pre { margin: 0; background: #1d1d1f; color: #f5f5f7; font-size: 12px; }
.doc-body h2 { margin-top: 24px; font-weight: 600; }
.doc-body h3 { margin-top: 16px; font-weight: 500; }

/* Doc template picker */
.doc-template-bar {
  display: flex; gap: 14px; align-items: flex-start; flex-wrap: wrap;
  padding: 14px 20px; border-bottom: 1px solid #d2d2d7;
  background: #fbfbfd;
}
.doc-template-bar label { font-size: 12px; color: #86868b; font-weight: 500; }
.doc-template-bar select {
  padding: 6px 10px; border: 1px solid #d2d2d7; border-radius: 8px;
  background: #fff; font-size: 13px; color: #1d1d1f; min-width: 160px;
}
.doc-template-desc { font-size: 12px; color: #86868b; align-self: center; }
.section-toggle-list {
  display: flex; flex-wrap: wrap; gap: 8px 14px; margin-top: 10px;
  padding: 12px; background: #fff; border: 1px solid #d2d2d7; border-radius: 10px;
  width: 100%;
}
.section-toggle {
  display: inline-flex; align-items: center; gap: 6px;
  font-size: 13px; color: #1d1d1f; cursor: pointer;
  padding: 4px 8px; border-radius: 6px;
}
.section-toggle:hover { background: #f5f5f7; }
.section-toggle input { cursor: pointer; }

/* User docs panel */
.user-docs-panel {
  background: #ffffff;
  border: 1px solid #d2d2d7;
  border-radius: 14px;
  padding: 18px 20px;
  margin-bottom: 18px;
}
.user-docs-head {
  display: flex; align-items: center; justify-content: space-between;
  margin-bottom: 10px;
}
.user-docs-title { font-size: 14px; font-weight: 600; color: #1d1d1f; }
.user-docs-sub { font-size: 12px; color: #86868b; margin-bottom: 12px; line-height: 1.6; }
.user-docs-drop {
  border: 1.5px dashed #d2d2d7; border-radius: 12px;
  padding: 22px; text-align: center; color: #86868b; font-size: 13px;
  cursor: pointer; transition: all .15s; background: #fbfbfd;
}
.user-docs-drop:hover, .user-docs-drop.dragover {
  border-color: #0071e3; color: #0071e3; background: #f0f7ff;
}
.user-docs-list { margin-top: 12px; display: flex; flex-direction: column; gap: 6px; }
.user-docs-item {
  display: flex; align-items: center; gap: 12px;
  padding: 9px 12px; background: #fbfbfd; border-radius: 8px;
  font-size: 13px;
}
.user-docs-item .fname { flex: 1; color: #1d1d1f; overflow: hidden; text-overflow: ellipsis; white-space: nowrap; }
.user-docs-item .size { color: #86868b; font-size: 12px; }
.user-docs-item .del-btn {
  background: none; border: none; color: #86868b; cursor: pointer;
  padding: 4px 8px; border-radius: 6px; font-size: 12px;
}
.user-docs-item .del-btn:hover { background: #ff3b30; color: #fff; }

/* Loading */
.loading-block { padding: 60px 20px; text-align: center; color: #86868b; }
.spinner-lg {
  display: inline-block;
  width: 36px; height: 36px;
  border: 3px solid #f5f5f7;
  border-top-color: #0071e3;
  border-radius: 50%;
  animation: appleSpin .8s linear infinite;
}

/* Chat */
.chat-pane {
  background: #ffffff;
  border: 1px solid #d2d2d7;
  border-radius: 18px;
  display: flex; flex-direction: column;
  height: calc(100vh - 220px); min-height: 500px;
  overflow: hidden;
}
.chat-head { padding: 16px 24px; border-bottom: 1px solid #d2d2d7; display: flex; align-items: center; gap: 12px; flex-shrink: 0; }
.chat-head-avatar {
  width: 38px; height: 38px;
  border-radius: 50%;
  background: #0071e3;
  display: flex; align-items: center; justify-content: center;
  color: #fff; font-weight: 600; font-size: 13px;
}
.chat-head-name { font-weight: 600; font-size: 14px; color: #1d1d1f; }
.chat-head-status { font-size: 12px; color: #248a3d; display: flex; align-items: center; gap: 6px; font-weight: 400; }
.chat-head-status::before { content: ""; width: 6px; height: 6px; border-radius: 50%; background: #34c759; }
.chat-stream {
  flex: 1; overflow-y: auto;
  padding: 26px 30px;
  display: flex; flex-direction: column;
  gap: 18px;
  background: #ffffff;
}
.bubble-row { display: flex; gap: 10px; animation: bubbleIn .25s ease; }
@keyframes bubbleIn { from { opacity: 0; transform: translateY(8px); } to { opacity: 1; transform: translateY(0); } }
.bubble-row.me { flex-direction: row-reverse; }
.bubble-avatar {
  width: 34px; height: 34px;
  border-radius: 50%;
  display: flex; align-items: center; justify-content: center;
  color: #fff; font-weight: 600; font-size: 12px;
  flex-shrink: 0;
}
.bubble-row.me .bubble-avatar { background: #1d1d1f; }
.bubble-row.bot .bubble-avatar { background: #0071e3; }
.bubble-body { max-width: 74%; }
.bubble-row.me .bubble-body { text-align: right; }
.bubble-meta { font-size: 11px; color: #86868b; margin-bottom: 4px; padding: 0 4px; font-weight: 400; }
.bubble {
  display: inline-block;
  padding: 12px 18px;
  border-radius: 20px;
  font-size: 14px;
  line-height: 1.6;
  white-space: pre-wrap;
  word-break: break-word;
  text-align: left;
}
.bubble-row.me .bubble {
  background: #0071e3;
  color: #fff;
  border-bottom-right-radius: 6px;
}
.bubble-row.bot .bubble {
  background: #f5f5f7;
  color: #1d1d1f;
  border-bottom-left-radius: 6px;
}
.thinking { display: flex; align-items: center; gap: 4px; padding: 14px 18px; }
.thinking-dot { width: 7px; height: 7px; border-radius: 50%; background: #86868b; animation: thinkBounce 1.2s infinite; }
.thinking-dot:nth-child(2) { animation-delay: .18s; }
.thinking-dot:nth-child(3) { animation-delay: .36s; }
@keyframes thinkBounce { 0%, 80%, 100% { transform: scale(.6); opacity: .4; } 40% { transform: scale(1); opacity: 1; } }
.chat-input-bar { border-top: 1px solid #d2d2d7; padding: 16px 20px; display: flex; align-items: flex-end; gap: 10px; background: #ffffff; }
.chat-textarea {
  flex: 1;
  border: none;
  border-radius: 14px;
  padding: 11px 14px;
  font-size: 14px;
  resize: none;
  outline: none;
  min-height: 44px;
  max-height: 120px;
  transition: box-shadow .15s;
  background: #f5f5f7;
  color: #1d1d1f;
}
.chat-textarea:focus { box-shadow: 0 0 0 4px rgba(0,113,227,0.2); }
.chat-send {
  width: 44px; height: 44px;
  border-radius: 50%;
  background: #0071e3;
  color: #fff;
  display: inline-flex; align-items: center; justify-content: center; text-align: center;
  transition: background .15s;
  flex-shrink: 0;
}
.chat-send:hover:not(:disabled) { background: #0077ed; }
.chat-send:disabled { opacity: .4; cursor: not-allowed; }

/* Hamburger */
.hamburger {
  display: none; width: 36px; height: 36px; border-radius: 10px;
  align-items: center; justify-content: center; text-align: center;
  color: #1d1d1f; background: transparent; transition: background .15s; flex-shrink: 0;
}
.hamburger:hover { background: #f5f5f7; }
.hamburger svg { width: 20px; height: 20px; }
.sidebar-backdrop { display: none; }

/* ============================================================
   RESPONSIVE DESIGN
   ============================================================ */

/* Tablet : 768px - 1024px */
@media (max-width: 1024px) {
  :root { --sidebar: 200px; }
  .content { padding: 24px 22px; }
  .topbar { padding: 0 18px; }
  .doc-body { padding: 24px 26px; }
  .repo-grid { grid-template-columns: repeat(auto-fill, minmax(280px, 1fr)); gap: 16px; }
  .detail-title { font-size: 24px; }
  .sidebar-head { padding: 16px 14px 12px; }
  .sidebar-section-title { padding: 14px 14px 6px; }
}

/* Mobile : < 768px */
@media (max-width: 768px) {
  html, body { font-size: 14px; }
  .shell { grid-template-columns: 1fr; }
  .sidebar {
    position: fixed; top: 0; left: 0;
    width: 280px; height: 100vh;
    z-index: 200;
    transform: translateX(-100%);
    transition: transform .28s cubic-bezier(.2,.8,.2,1);
    box-shadow: 0 10px 30px rgba(0,0,0,0.1);
    background: #ffffff;
  }
  .shell.sidebar-open .sidebar { transform: translateX(0); }
  .sidebar-backdrop {
    display: block;
    position: fixed; inset: 0;
    background: rgba(0, 0, 0, 0.4);
    backdrop-filter: blur(6px);
    -webkit-backdrop-filter: blur(6px);
    z-index: 150;
    opacity: 0; pointer-events: none;
    transition: opacity .25s ease;
  }
  .shell.sidebar-open .sidebar-backdrop { opacity: 1; pointer-events: auto; }
  .hamburger { display: inline-flex; }
  .topbar { padding: 0 14px; height: 50px; gap: 10px; }
  .topbar .crumb { font-size: 12.5px; flex: 1; min-width: 0; overflow: hidden; text-overflow: ellipsis; white-space: nowrap; }
  .topbar .crumb b { overflow: hidden; text-overflow: ellipsis; }
  .topbar .topbar-actions { gap: 4px; }
  .topbar a.platform-link { padding: 5px 8px; font-size: 11.5px; }
  .content { padding: 20px 16px; }
  .main { height: 100vh; }
  .detail-title { font-size: 22px; }
  .detail-title-row { flex-wrap: wrap; gap: 10px; }
  .detail-head { padding-bottom: 16px; margin-bottom: 18px; }
  .detail-stats-row { gap: 8px; }
  .stat-block { padding: 6px 10px; font-size: 12px; }
  .repo-grid { grid-template-columns: 1fr; gap: 14px; }
  .repo-card { padding: 18px; }
  .repo-card-head { gap: 10px; margin-bottom: 10px; }
  .repo-avatar { width: 38px; height: 38px; font-size: 15px; }
  .repo-card-name { font-size: 14.5px; }
  .doc-pane { border-radius: 14px; }
  .doc-toolbar { padding: 12px 16px; flex-wrap: wrap; gap: 8px; }
  .doc-toolbar-title { font-size: 13px; }
  .doc-body { padding: 20px 16px; font-size: 14px; }
  .doc-body h1, .doc-body h2 { font-size: 18px; }
  .doc-body pre { font-size: 12px; padding: 12px 14px; }
  .diagram { font-size: 11.5px; padding: 14px 14px; }
  .chat-pane { height: calc(100vh - 200px); min-height: 420px; border-radius: 14px; }
  .chat-head { padding: 12px 16px; }
  .chat-stream { padding: 16px 14px; gap: 14px; }
  .bubble-body { max-width: 85%; }
  .bubble { font-size: 13.5px; padding: 10px 14px; }
  .chat-input-bar { padding: 12px 14px; gap: 8px; }
  .chat-textarea { font-size: 14px; padding: 10px 12px; }
  .chat-send { width: 40px; height: 40px; }
  .btn-primary, .btn-secondary { padding: 11px 20px; }
  .modal-mask { padding: 0; align-items: stretch; }
  .modal-card {
    width: 100%; max-width: 100%;
    height: 100%; max-height: 100vh;
    border-radius: 0;
    padding: 24px 20px;
    display: flex; flex-direction: column;
  }
  .modal-actions { margin-top: auto; padding-top: 20px; }
  .modal-actions button { flex: 1; }
  .toast-wrap { top: 12px; right: 12px; left: 12px; }
  .toast { max-width: none; font-size: 13px; padding: 11px 14px; }
}

/* Small phone : < 480px */
@media (max-width: 480px) {
  .content { padding: 16px 12px; }
  .topbar { padding: 0 10px; height: 48px; }
  .topbar a.platform-link { display: none; }
  .topbar .topbar-actions { margin-left: auto; }
  .detail-title { font-size: 20px; letter-spacing: -0.3px; }
  .detail-head { padding-bottom: 12px; margin-bottom: 14px; }
  .detail-sub { font-size: 12.5px; }
  .repo-card { padding: 16px; border-radius: 14px; }
  .repo-card-org { font-size: 11.5px; }
  .repo-status-badge { font-size: 11px; padding: 2px 8px; }
  .doc-toolbar { padding: 10px 12px; }
  .doc-body { padding: 16px 12px; line-height: 1.6; }
  .chat-stream { padding: 14px 10px; }
  .bubble-body { max-width: 88%; }
  .bubble { font-size: 13px; padding: 9px 12px; border-radius: 16px; }
  .bubble-avatar { width: 30px; height: 30px; font-size: 11px; }
  .modal-card { padding: 20px 16px; }
  .modal-title { font-size: 19px; }
  .stat-block { font-size: 11.5px; padding: 5px 10px; }
  .stat-block .num { font-size: 14px; }
  .brand-sub { display: none; }
}

@media (max-width: 540px) {
  .modal-card { width: calc(100vw - 24px) !important; max-width: 100% !important; margin: 12px !important; }
  .apple-auth-card { width: calc(100vw - 32px) !important; max-width: 100% !important; padding: 28px 20px !important; }
  .toast { width: calc(100vw - 24px) !important; max-width: 100% !important; right: 12px !important; left: 12px !important; }
}
@media (max-width: 480px) {
  .modal-card { padding: 20px !important; }
  .apple-auth-card { padding: 24px 16px !important; }
  .apple-auth-title { font-size: 22px !important; }
  body { font-size: 14px !important; }
  .topbar { padding: 0 12px !important; }
}

.sidebar { height: 100vh; height: 100dvh; }
.repo-tree { flex: 1 1 auto; overflow-y: auto; -webkit-overflow-scrolling: touch; min-height: 0; }
.sidebar-add { flex-shrink: 0; }
.sidebar-foot { flex-shrink: 0; }
.sidebar-head { flex-shrink: 0; }
.sidebar-section-title { flex-shrink: 0; }
@media (max-width: 768px) {
  .sidebar { height: 100vh !important; height: 100dvh !important; max-height: 100vh; max-height: 100dvh; }
  .repo-tree { max-height: none !important; }
}


.groundedness-badge {
  display: inline-flex;
  align-items: center;
  padding: 4px 10px;
  border-radius: 12px;
  font-size: 11px;
  margin-top: 6px;
  background: #f5f5f7;
  color: #1d1d1f;
  border: 1px solid #d2d2d7;
  cursor: default;
}
.groundedness-badge.gnd-high { background: #e8f7ee; color: #1d6b3a; border-color: #b8e3c8; }
.groundedness-badge.gnd-mid { background: #fff4e0; color: #8a5a14; border-color: #f0d8a8; }
.groundedness-badge.gnd-low { background: #ffe9e7; color: #8a2a1f; border-color: #f0c0bc; }
.groundedness-badge .invalid-refs-toggle { color: inherit; opacity: 0.8; }
.groundedness-badge .invalid-refs-toggle:hover { opacity: 1; }
.invalid-refs-list {
  margin-top: 4px;
  font-size: 11px;
  color: #86868b;
  padding: 6px 10px;
  background: #fafafa;
  border-radius: 8px;
  border: 1px solid #e8e8ec;
}
.invalid-refs-list code {
  background: #fff;
  padding: 1px 5px;
  border-radius: 4px;
  font-family: 'SF Mono', Monaco, monospace;
  font-size: 10.5px;
  color: #1d1d1f;
}
</style>
</head>
<body>

<div id="app" v-cloak>

  <!-- Toasts -->
  <div class="toast-wrap">
    <div v-for="(t, i) in toasts" :key="i" :class="'toast ' + (t.type||'info')">{{ t.msg }}</div>
  </div>

  <!-- ===================== AUTH ===================== -->
  <div v-if="!token" class="apple-auth-page">
    <div class="apple-auth-card">
      <div class="apple-auth-logo">CodeDoc</div>
      <div class="apple-auth-subtitle">代码文档平台</div>
      <div class="apple-auth-form">
        <input v-model="authForm.username" class="apple-input" :placeholder="authMode==='login' ? '用户名' : '设置用户名'" @keydown.enter="doAuth" autocomplete="username">
        <input v-model="authForm.password" type="password" class="apple-input" :placeholder="authMode==='login' ? '密码' : '设置密码'" @keydown.enter="doAuth" autocomplete="current-password">
        <div v-if="authError" class="apple-auth-error">{{ authError }}</div>
        <button @click="doAuth" class="apple-btn" :disabled="authLoading">
          <span v-if="authLoading" class="apple-spinner"></span>
          {{ authMode === 'login' ? '登录' : '注册' }}
        </button>
        <div class="apple-auth-toggle">
          <span v-if="authMode==='login'" @click="authMode='register'">没有账号？创建账号</span>
          <span v-else @click="authMode='login'">已有账号？登录</span>
        </div>
      </div>
    </div>
  </div>

  <!-- ===================== MAIN ===================== -->
  <div v-else class="shell" :class="{'sidebar-open': sidebarOpen}">

    <!-- Backdrop for mobile drawer -->
    <div class="sidebar-backdrop" @click="sidebarOpen=false"></div>

    <!-- Sidebar -->
    <aside class="sidebar">
      <div class="sidebar-head">
        <div class="brand-mark">C</div>
        <div>
          <div class="brand-title">CodeDoc</div>
          <div class="brand-sub">代码文档平台</div>
        </div>
      </div>

      <!-- Top-level navigation (above repo tree) -->
      <div class="sidebar-nav">
        <a class="nav-link" :class="{active: currentTab === 'tasks'}" @click="goPage('tasks')">
          <svg width="14" height="14" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round" style="opacity:.7;flex-shrink:0">
            <path d="M9 11l3 3L22 4"/><path d="M21 12v7a2 2 0 0 1-2 2H5a2 2 0 0 1-2-2V5a2 2 0 0 1 2-2h11"/>
          </svg>
          <span style="flex:1">我的任务</span>
          <span v-if="runningTaskCount() > 0" class="task-badge">{{ runningTaskCount() }}</span>
        </a>
      </div>

      <div class="sidebar-section-title">仓库</div>

      <div class="repo-tree">
        <div v-for="r in repos" :key="r.name" class="repo-item" :class="{active: currentRepo && currentRepo.name===r.name && currentTab==='overview'}">
          <div class="repo-row" @click="selectRepo(r,'overview')">
            <svg width="14" height="14" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round" style="opacity:.5;flex-shrink:0">
              <path d="M21 16V8a2 2 0 0 0-1-1.73l-7-4a2 2 0 0 0-2 0l-7 4A2 2 0 0 0 3 8v8a2 2 0 0 0 1 1.73l7 4a2 2 0 0 0 2 0l7-4A2 2 0 0 0 21 16z"/>
            </svg>
            <span class="name" :title="r.name">{{ shortRepoName(r.name) }}</span>
            <span class="status-dot" :class="r.status"></span>
          </div>
          <div v-if="r.status==='ready' && currentRepo && currentRepo.name===r.name" class="repo-children">
            <div class="child" :class="{active: currentTab==='doc'}" @click="selectRepo(r,'doc')">
              <svg class="ic" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><path d="M14 2H6a2 2 0 0 0-2 2v16a2 2 0 0 0 2 2h12a2 2 0 0 0 2-2V8z"/><polyline points="14 2 14 8 20 8"/></svg>
              设计文档
            </div>
            <div class="child" :class="{active: currentTab==='qa'}" @click="selectRepo(r,'qa')">
              <svg class="ic" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><path d="M21 15a2 2 0 0 1-2 2H7l-4 4V5a2 2 0 0 1 2-2h14a2 2 0 0 1 2 2z"/></svg>
              代码问答
            </div>
          </div>
        </div>
        <div v-if="repos.length===0 && !reposLoading" style="padding:20px 12px;font-size:12.5px;color:var(--text-3);text-align:center;line-height:1.6">
          还没有仓库<br>点击下方按钮添加
        </div>
      </div>

      <div class="sidebar-add">
        <button class="btn-add-repo" @click="showAddModal=true; sidebarOpen=false">
          <svg width="14" height="14" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2.5" stroke-linecap="round" stroke-linejoin="round"><line x1="12" y1="5" x2="12" y2="19"/><line x1="5" y1="12" x2="19" y2="12"/></svg>
          添加仓库
        </button>
      </div>

      <div class="sidebar-foot">
        <div class="avatar">{{ (username||'?').charAt(0).toUpperCase() }}</div>
        <div class="user-info">
          <div class="user-name">{{ username }}</div>
          <div class="user-sub">已登录</div>
        </div>
        <button class="btn-icon" @click="logout" title="退出">
          <svg width="16" height="16" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><path d="M9 21H5a2 2 0 0 1-2-2V5a2 2 0 0 1 2-2h4"/><polyline points="16 17 21 12 16 7"/><line x1="21" y1="12" x2="9" y2="12"/></svg>
        </button>
      </div>
    </aside>

    <!-- Main -->
    <main class="main">
      <header class="topbar">
        <button class="hamburger" @click="sidebarOpen=!sidebarOpen" :aria-label="sidebarOpen ? '关闭菜单' : '打开菜单'">
          <svg viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round">
            <line v-if="!sidebarOpen" x1="3" y1="6" x2="21" y2="6"/>
            <line v-if="!sidebarOpen" x1="3" y1="12" x2="21" y2="12"/>
            <line v-if="!sidebarOpen" x1="3" y1="18" x2="21" y2="18"/>
            <line v-if="sidebarOpen" x1="18" y1="6" x2="6" y2="18"/>
            <line v-if="sidebarOpen" x1="6" y1="6" x2="18" y2="18"/>
          </svg>
        </button>
        <div class="crumb">
          <span>CodeDoc</span>
          <template v-if="currentRepo"><span>/</span><b>{{ currentRepo.name }}</b></template>
          <template v-if="currentRepo && currentTab==='doc'"><span>/</span><span>设计文档</span></template>
          <template v-if="currentRepo && currentTab==='qa'"><span>/</span><span>代码问答</span></template>
        </div>
        <div class="topbar-actions">
          <a href="/platform/" class="platform-link">平台首页</a>
          <a href="/devbot/" class="platform-link">DevBot 评审</a>
        </div>
      </header>

      <div class="content">

        <!-- 我的任务 view -->
        <div v-if="currentTab === 'tasks'">
          <div class="detail-head">
            <div class="detail-title-row">
              <h1 class="detail-title">我的任务</h1>
              <span style="font-size:12px;color:var(--text-3);font-weight:500">{{ docgenTasks.length }} 条记录</span>
            </div>
            <div class="detail-sub">所有的文档生成任务记录（最近 50 条）。任务在后台运行，可关闭页面，完成后回来查看。</div>
          </div>
          <div v-if="!docgenTasks.length" class="empty" style="padding:60px 20px">
            <div class="empty-icon">📋</div>
            <div class="empty-title">还没有任务</div>
            <div class="empty-sub">前往任意仓库的"设计文档"页面点击"生成文档"即可创建任务</div>
          </div>
          <div v-else class="tasks-list">
            <div v-for="t in docgenTasks" :key="t.task_id" class="task-card">
              <div class="task-card-top">
                <div class="task-repo">{{ t.repo }}</div>
                <div class="task-status" :class="'status-' + t.status">
                  {{ TASK_STATUS_LABELS[t.status] || t.status }}
                </div>
              </div>
              <div class="task-meta">
                <span>{{ taskTemplateLabel(t) }}</span>
                <span>·</span>
                <span>{{ fmtTaskTime(t.submitted_at) }}</span>
                <span v-if="t.status === 'queued' && t.position">· 排队第 {{ t.position }} 位</span>
                <span v-if="t.stage && t.status === 'running'">· {{ DOC_STAGE_LABELS[t.stage] || t.stage }}</span>
              </div>
              <div v-if="t.status === 'running' && t.progress" class="task-progress">
                {{ t.progress }}
              </div>
              <div v-if="t.status === 'error'" class="task-error">
                {{ t.error || '未知错误' }}
              </div>
              <div v-if="t.status === 'done'" class="task-actions">
                <button class="btn-primary-pill" @click="viewTaskDocument(t)">查看文档</button>
              </div>
            </div>
          </div>
        </div>

        <!-- Overview (no repo selected OR overview tab) -->
        <div v-if="currentTab !== 'tasks' && (!currentRepo || currentTab==='overview')">
          <div class="detail-head" v-if="!currentRepo">
            <div class="detail-title-row">
              <h1 class="detail-title">我的仓库</h1>
            </div>
            <div class="detail-sub">管理你导入的 GitHub 仓库，每个仓库都会自动构建代码文档图谱。</div>
          </div>
          <div class="detail-head" v-else>
            <div class="detail-title-row">
              <h1 class="detail-title">{{ currentRepo.name }}</h1>
              <span class="repo-status-badge" :class="currentRepo.status">{{ statusText(currentRepo.status) }}</span>
            </div>
            <div class="detail-sub">{{ currentRepo.url }}</div>
            <div class="detail-stats-row" v-if="currentRepo.status==='ready'">
              <div class="stat-block"><span class="num">{{ currentRepo.nodes }}</span><span class="label">节点</span></div>
              <div class="stat-block"><span class="num">{{ currentRepo.edges }}</span><span class="label">边</span></div>
            </div>
          </div>

          <div v-if="reposLoading && repos.length===0" class="loading-block">
            <div class="spinner-lg"></div>
          </div>

          <div v-else-if="!currentRepo && repos.length===0" class="empty">
            <div class="empty-icon">📦</div>
            <div class="empty-title">还没有仓库</div>
            <div class="empty-sub">点击左侧"添加仓库"开始（支持 GitHub 链接或上传 ZIP）</div>
          </div>

          <div v-else-if="!currentRepo" class="repo-grid">
            <div v-for="r in repos" :key="r.name" class="repo-card" @click="selectRepo(r, r.status==='ready'?'doc':'overview')">
              <div class="repo-card-head">
                <div class="repo-avatar">{{ (r.name.split('/')[0]||'?').charAt(0).toUpperCase() }}</div>
                <div class="repo-meta">
                  <div class="repo-card-name">{{ r.name.split('/')[1] || r.name }}</div>
                  <div class="repo-card-org">{{ r.name.split('/')[0] }}</div>
                </div>
                <span class="repo-status-badge" :class="r.status">{{ statusText(r.status) }}</span>
              </div>
              <div class="repo-card-stats" v-if="r.status==='ready'">
                <div class="stat-pill">
                  <svg width="11" height="11" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2.5" stroke-linecap="round" stroke-linejoin="round"><circle cx="12" cy="12" r="3"/><circle cx="12" cy="2" r="2"/><circle cx="12" cy="22" r="2"/><circle cx="2" cy="12" r="2"/><circle cx="22" cy="12" r="2"/></svg>
                  {{ r.nodes }} 节点
                </div>
                <div class="stat-pill">
                  <svg width="11" height="11" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2.5" stroke-linecap="round" stroke-linejoin="round"><line x1="7" y1="17" x2="17" y2="7"/><polyline points="7 7 17 7 17 17"/></svg>
                  {{ r.edges }} 边
                </div>
              </div>
              <div v-if="r.progress && r.status!=='ready' && r.status!=='error'" class="repo-card-progress">
                {{ r.progress }}
              </div>
              <div class="repo-card-foot">
                <span>{{ formatDate(r.created_at) }}</span>
                <button class="btn-icon" @click.stop="deleteRepo(r)" title="删除" style="color:var(--danger)">
                  <svg width="14" height="14" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><polyline points="3 6 5 6 21 6"/><path d="M19 6l-2 14a2 2 0 0 1-2 2H9a2 2 0 0 1-2-2L5 6"/><path d="M10 11v6M14 11v6"/></svg>
                </button>
              </div>
            </div>
          </div>

          <!-- Repo overview detail (when selected and status=ready) -->
          <div v-else-if="currentRepo && currentRepo.status==='ready'">
            <div style="display:grid;grid-template-columns:repeat(auto-fit,minmax(220px,1fr));gap:14px;margin-bottom:24px">
              <div class="repo-card" @click="selectRepo(currentRepo,'doc')" style="cursor:pointer">
                <div style="font-size:13px;color:var(--text-3);margin-bottom:6px;letter-spacing:0.4px">📄 设计文档</div>
                <div style="font-size:15px;font-weight:600;margin-bottom:4px">AI 生成架构文档</div>
                <div style="font-size:12.5px;color:var(--text-3);line-height:1.6">基于代码文档图谱，自动生成模块、调用关系和接口说明的设计文档。</div>
              </div>
              <div class="repo-card" @click="selectRepo(currentRepo,'qa')" style="cursor:pointer">
                <div style="font-size:13px;color:var(--text-3);margin-bottom:6px;letter-spacing:0.4px">💬 代码问答</div>
                <div style="font-size:15px;font-weight:600;margin-bottom:4px">与代码对话</div>
                <div style="font-size:12.5px;color:var(--text-3);line-height:1.6">用自然语言提问，AI 基于知识图谱检索相关代码并回答你的问题。</div>
              </div>
            </div>
          </div>
        </div>

        <!-- Doc tab -->
        <div v-if="currentRepo && currentTab==='doc'">
          <div class="detail-head">
            <div class="detail-title-row">
              <h1 class="detail-title">设计文档</h1>
              <span style="font-size:12px;color:var(--text-3);font-weight:500;background:#fff;padding:4px 10px;border-radius:8px;border:1px solid var(--border)">{{ currentRepo.name }}</span>
            </div>
            <div class="detail-sub">基于代码文档图谱（{{ currentRepo.nodes }} 节点 / {{ currentRepo.edges }} 边）生成的架构设计文档，使用 PlantUML 渲染</div>
          </div>

          <!-- User-uploaded reference docs panel -->
          <div class="user-docs-panel">
            <div class="user-docs-head">
              <div class="user-docs-title">📎 参考资料</div>
              <span style="font-size:12px;color:var(--text-3)">{{ userDocs.length }} / 不限</span>
            </div>
            <div class="user-docs-sub">
              上传项目相关文档（需求/设计/规范等），代码问答时 AI 会综合参考。
              支持 PDF / DOCX / MD / TXT / HTML，单文件最大 20MB。
            </div>
            <div class="user-docs-drop"
                 :class="{dragover: docDragOver}"
                 @click="$refs.userDocInput.click()"
                 @dragover.prevent="docDragOver=true"
                 @dragleave.prevent="docDragOver=false"
                 @drop.prevent="onDropUserDoc">
              <input type="file" ref="userDocInput" hidden accept=".pdf,.docx,.md,.txt,.html,.htm" @change="onPickUserDoc">
              <span v-if="userDocUploading">上传解析中...</span>
              <span v-else>拖拽文件到这里，或<span style="color:#0071e3;font-weight:500">点击选择文件</span></span>
            </div>
            <div v-if="userDocs.length" class="user-docs-list">
              <div v-for="d in userDocs" :key="d.id" class="user-docs-item">
                <span style="font-size:14px">📄</span>
                <span class="fname" :title="d.filename">{{ d.filename }}</span>
                <span class="size">{{ humanSize(d.size) }} · 文本 {{ d.text_length }} 字</span>
                <button class="del-btn" @click="deleteUserDoc(d)">删除</button>
              </div>
            </div>
          </div>

          <div class="doc-pane">
            <div class="doc-template-bar">
              <div style="display:flex; gap:8px; align-items:center;">
                <label>文档模板</label>
                <select v-model="docTemplate" @change="onTemplateChange">
                  <option v-for="(t,k) in docTemplates" :key="k" :value="k">{{ t.name }}</option>
                  <option value="custom">自定义</option>
                </select>
              </div>
              <span v-if="docTemplate!=='custom' && docTemplates[docTemplate]" class="doc-template-desc">
                {{ docTemplates[docTemplate].description }}
              </span>
              <div v-if="docTemplate==='custom'" class="section-toggle-list">
                <label v-for="(label, sec) in sectionLabels" :key="sec" class="section-toggle">
                  <input type="checkbox" :value="sec" v-model="customSections">
                  <span>{{ label }}</span>
                </label>
              </div>
            </div>
            <div class="doc-toolbar">
              <span class="doc-toolbar-title">📐 架构设计文档</span>
              <button @click="doGenDoc" class="btn-primary" :disabled="docLoading">
                <span v-if="docLoading" class="apple-spinner" style="width:13px;height:13px;border-width:2px"></span>
                <svg v-else width="13" height="13" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2.5" stroke-linecap="round" stroke-linejoin="round"><polyline points="23 4 23 10 17 10"/><path d="M20.49 15a9 9 0 1 1-2.12-9.36L23 10"/></svg>
                {{ docContent ? '重新生成' : '生成文档' }}
              </button>
            </div>
            <div class="doc-body">
              <div v-if="docLoading" class="loading-block">
                <div class="spinner-lg"></div>
                <div style="margin-top:14px;font-size:14px;color:var(--text-2);font-weight:500">{{ docStage || '正在分析代码结构，生成专业文档...' }}</div>
                <div v-if="docProgress" style="margin-top:6px;font-size:12.5px;color:var(--text-3)">{{ docProgress }}</div>
              </div>
              <div v-else-if="docContent" v-html="docContent"></div>
              <div v-else class="empty" style="padding:60px 20px">
                <div class="empty-icon">📐</div>
                <div class="empty-title">还没有生成文档</div>
                <div class="empty-sub">选择模板后点击"生成文档"，AI 将基于代码文档图谱与 PlantUML 渲染图表</div>
              </div>
            </div>
          </div>
        </div>

        <!-- QA tab -->
        <div v-if="currentRepo && currentTab==='qa'">
          <div class="detail-head">
            <div class="detail-title-row">
              <h1 class="detail-title">代码问答</h1>
              <span style="font-size:12px;color:var(--text-3);font-weight:500;background:#fff;padding:4px 10px;border-radius:8px;border:1px solid var(--border)">{{ currentRepo.name }}</span>
            </div>
            <div class="detail-sub">用自然语言提问，AI 基于知识图谱回答关于这个仓库代码的任何问题</div>
          </div>

          <div class="chat-pane">
            <div class="chat-head">
              <div class="chat-head-avatar">AI</div>
              <div style="flex:1">
                <div class="chat-head-name">代码文档助手</div>
                <div class="chat-head-status">在线 · 基于 DeepSeek-V3</div>
              </div>
            </div>
            <div class="chat-stream" ref="chatBox">
              <div v-for="(m, i) in chatMessages" :key="i" :class="'bubble-row ' + m.role">
                <div class="bubble-avatar">{{ m.role==='me' ? (username||'?').charAt(0).toUpperCase() : 'AI' }}</div>
                <div class="bubble-body">
                  <div class="bubble-meta">{{ m.role==='me' ? username : '代码助手' }} · {{ m.time }}</div>
                  <div class="bubble">{{ m.content }}</div>
                  <div v-if="m.role==='bot' && m.groundedness !== undefined && m.total_refs > 0" class="groundedness-badge" :class="groundednessClass(m.groundedness)" :title="groundednessTitle(m)">
                    <svg width="11" height="11" viewBox="0 0 24 24" fill="currentColor" style="vertical-align:-1px;margin-right:3px"><path d="M12 2L4 5v6c0 5 3.5 9.5 8 11 4.5-1.5 8-6 8-11V5l-8-3z"/></svg>
                    <span>引用可信度 {{ Math.round(m.groundedness * 100) }}%</span>
                    <span style="opacity:0.75;margin-left:4px">· {{ m.valid_refs_count }}/{{ m.total_refs }} 已核对</span>
                    <span v-if="m.invalid_refs && m.invalid_refs.length" class="invalid-refs-toggle" @click="toggleInvalidRefs(i)" style="margin-left:6px;cursor:pointer;text-decoration:underline">{{ expandedInvalid[i] ? '收起' : '查看未匹配' }}</span>
                  </div>
                  <div v-if="m.role==='bot' && expandedInvalid[i] && m.invalid_refs && m.invalid_refs.length" class="invalid-refs-list">
                    未在知识图谱中找到：<code v-for="r in m.invalid_refs" :key="r" style="margin-right:6px">{{ r }}</code>
                  </div>
                </div>
              </div>
              <div v-if="chatLoading" class="bubble-row bot">
                <div class="bubble-avatar">AI</div>
                <div class="bubble-body">
                  <div class="bubble-meta">代码助手 · 思考中</div>
                  <div class="bubble thinking">
                    <span class="thinking-dot"></span>
                    <span class="thinking-dot"></span>
                    <span class="thinking-dot"></span>
                  </div>
                </div>
              </div>
            </div>
            <div class="chat-input-bar">
              <textarea v-model="chatInput" class="chat-textarea" placeholder="输入关于代码的问题，按 Enter 发送..." rows="1" @keydown.enter.exact.prevent="doAsk"></textarea>
              <button class="chat-send" @click="doAsk" :disabled="chatLoading || !chatInput.trim()" title="发送">
                <svg width="18" height="18" viewBox="0 0 24 24" fill="currentColor"><path d="M2 21l21-9L2 3v7l15 2-15 2v7z"/></svg>
              </button>
            </div>
          </div>
        </div>

      </div>
    </main>

    <!-- Add Modal -->
    <div v-if="showAddModal" class="modal-mask" @click.self="closeAddModal">
      <div class="modal-card">
        <div class="modal-title">添加仓库</div>
        <div class="modal-sub">从 GitHub 导入或上传本地 ZIP，系统将自动构建代码文档图谱</div>

        <div class="add-tabs">
          <button class="add-tab" :class="{active: addTab==='github'}" @click="addTab='github'">GitHub 链接</button>
          <button class="add-tab" :class="{active: addTab==='upload'}" @click="addTab='upload'">本地上传</button>
        </div>

        <!-- GitHub tab -->
        <div v-if="addTab==='github'">
          <input v-model="addUrl" class="modal-input" placeholder="https://github.com/owner/repo" @keydown.enter="doAddRepo">
          <div class="modal-actions">
            <button @click="closeAddModal" class="btn-secondary">取消</button>
            <button @click="doAddRepo" class="btn-primary" :disabled="addLoading">
              <span v-if="addLoading" class="apple-spinner" style="width:13px;height:13px;border-width:2px"></span>
              开始导入
            </button>
          </div>
        </div>

        <!-- Upload tab -->
        <div v-else-if="addTab==='upload'">
          <input v-model="uploadName" class="modal-input" placeholder="仓库名称（字母 / 数字 / - / _）" style="margin-bottom:12px">
          <input ref="fileInput" type="file" accept=".zip" style="display:none" @change="onFilePicked">
          <div class="drop-zone"
               :class="{'has-file': !!uploadFile, dragging: dragOver}"
               @click="$refs.fileInput && $refs.fileInput.click()"
               @dragover.prevent="dragOver=true"
               @dragleave.prevent="dragOver=false"
               @drop.prevent="onFileDropped">
            <svg v-if="!uploadFile" class="drop-icon" width="40" height="40" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="1.5" stroke-linecap="round" stroke-linejoin="round">
              <path d="M21 15v4a2 2 0 0 1-2 2H5a2 2 0 0 1-2-2v-4"/>
              <polyline points="17 8 12 3 7 8"/>
              <line x1="12" y1="3" x2="12" y2="15"/>
            </svg>
            <div v-if="!uploadFile" class="drop-text">拖拽 ZIP 文件到此处或点击选择</div>
            <div v-if="!uploadFile" class="drop-hint">（最大 50MB）</div>
            <div v-else class="picked-file">
              <div class="picked-name">📦 {{ uploadFile.name }}</div>
              <div class="picked-size">{{ humanSize(uploadFile.size) }}</div>
              <button class="picked-clear" @click.stop="uploadFile=null">移除</button>
            </div>
          </div>
          <button @click="doUploadRepo" class="btn-primary" style="width:100%;margin-top:14px;justify-content:center" :disabled="uploadLoading || !uploadFile || !uploadName.trim()">
            <span v-if="uploadLoading" class="apple-spinner" style="width:13px;height:13px;border-width:2px"></span>
            {{ uploadLoading ? '上传中…' : '上传并索引' }}
          </button>
          <div class="modal-actions" style="margin-top:8px">
            <button @click="closeAddModal" class="btn-secondary">取消</button>
          </div>
        </div>
      </div>
    </div>

  </div>

</div>

<script>
const BASE = '/codedoc';
const { createApp, ref, reactive, onMounted, nextTick, watch } = Vue;

createApp({
  setup() {
    const token = ref(localStorage.getItem('cw_token') || '');
    const username = ref(localStorage.getItem('cw_username') || '');
    const sidebarOpen = ref(false);
    const authMode = ref('login');
    const authForm = reactive({ username: '', password: '' });
    const authLoading = ref(false);
    const authError = ref('');

    const toasts = ref([]);
    function toast(msg, type='info') {
      const t = { msg, type };
      toasts.value.push(t);
      setTimeout(() => { const i = toasts.value.indexOf(t); if (i>=0) toasts.value.splice(i,1); }, 3500);
    }

    async function api(path, opts={}) {
      const headers = { 'Content-Type': 'application/json', ...(opts.headers||{}) };
      if (token.value) headers['Authorization'] = 'Bearer ' + token.value;
      const res = await fetch(BASE + path, { ...opts, headers });
      if (res.status === 401) {
        token.value = ''; localStorage.removeItem('cw_token');
        toast('会话已过期，请重新登录', 'error');
        throw new Error('unauthorized');
      }
      let data;
      try { data = await res.json(); } catch { data = {}; }
      if (!res.ok) { const msg = data.detail || data.error || '请求失败'; toast(msg, 'error'); throw new Error(msg); }
      return data;
    }

    async function doAuth() {
      authError.value = '';
      if (!authForm.username || !authForm.password) { authError.value = '请填写用户名和密码'; return; }
      authLoading.value = true;
      try {
        const endpoint = authMode.value === 'login' ? '/api/v1/auth/login' : '/api/v1/auth/register';
        const data = await api(endpoint, { method: 'POST', body: JSON.stringify(authForm) });
        token.value = data.token;
        username.value = data.username;
        localStorage.setItem('cw_token', data.token);
        localStorage.setItem('cw_username', data.username);
        toast(authMode.value === 'login' ? '欢迎回来 · ' + data.username : '注册成功 · 欢迎使用', 'success');
        await loadRepos();
      } catch(e) { authError.value = e.message; }
      authLoading.value = false;
    }
    function logout() {
      token.value = '';
      username.value = '';
      localStorage.removeItem('cw_token');
      localStorage.removeItem('cw_username');
      repos.value = [];
      currentRepo.value = null;
    }

    // Repos
    const repos = ref([]);
    const reposLoading = ref(false);
    const showAddModal = ref(false);
    const addTab = ref('github');
    const addUrl = ref('');
    const addLoading = ref(false);
    // Upload state
    const uploadName = ref('');
    const uploadFile = ref(null);
    const uploadLoading = ref(false);
    const dragOver = ref(false);
    const fileInput = ref(null);
    let pollTimer = null;

    function closeAddModal() {
      showAddModal.value = false;
      // Reset transient state when closed (keep tab choice for UX)
      addUrl.value = '';
      uploadFile.value = null;
      uploadName.value = '';
      dragOver.value = false;
    }
    function onFilePicked(e) {
      const f = e.target.files && e.target.files[0];
      if (f) acceptFile(f);
      e.target.value = '';
    }
    function onFileDropped(e) {
      dragOver.value = false;
      const f = e.dataTransfer.files && e.dataTransfer.files[0];
      if (f) acceptFile(f);
    }
    function acceptFile(f) {
      if (!/\.zip$/i.test(f.name)) { toast('请选择 .zip 文件', 'error'); return; }
      if (f.size > 50 * 1024 * 1024) { toast('文件大小不能超过 50MB', 'error'); return; }
      uploadFile.value = f;
      if (!uploadName.value.trim()) {
        const base = f.name.replace(/\.zip$/i, '');
        uploadName.value = base;
      }
    }
    function humanSize(n) {
      if (n < 1024) return n + ' B';
      if (n < 1024 * 1024) return (n / 1024).toFixed(1) + ' KB';
      return (n / 1024 / 1024).toFixed(2) + ' MB';
    }

    async function doUploadRepo() {
      if (!uploadFile.value) { toast('请先选择 ZIP 文件', 'error'); return; }
      const rawName = uploadName.value.trim();
      if (!rawName) { toast('请输入仓库名称', 'error'); return; }
      uploadLoading.value = true;
      try {
        const fd = new FormData();
        fd.append('file', uploadFile.value);
        fd.append('name', rawName);
        const headers = {};
        if (token.value) headers['Authorization'] = 'Bearer ' + token.value;
        const res = await fetch(BASE + '/api/v1/repos/upload', { method: 'POST', body: fd, headers });
        let data; try { data = await res.json(); } catch { data = {}; }
        if (!res.ok) { throw new Error(data.detail || data.error || '上传失败'); }
        toast('已开始索引 ' + (data.name || rawName), 'success');
        closeAddModal();
        await loadRepos();
      } catch(e) {
        toast(e.message || '上传失败', 'error');
      }
      uploadLoading.value = false;
    }

    async function loadRepos() {
      reposLoading.value = true;
      try {
        const data = await api('/api/v1/repos');
        repos.value = data.repos || [];
        // refresh currentRepo from list
        if (currentRepo.value) {
          const fresh = repos.value.find(r => r.name === currentRepo.value.name);
          if (fresh) currentRepo.value = fresh;
        }
        checkPolling();
      } catch(e) {}
      reposLoading.value = false;
    }

    function checkPolling() {
      const needsPoll = repos.value.some(r => r.status === 'cloning' || r.status === 'indexing');
      if (needsPoll && !pollTimer) {
        pollTimer = setInterval(async () => {
          try {
            const data = await api('/api/v1/repos');
            repos.value = data.repos || [];
            if (currentRepo.value) {
              const fresh = repos.value.find(r => r.name === currentRepo.value.name);
              if (fresh) currentRepo.value = fresh;
            }
            const still = repos.value.some(r => r.status === 'cloning' || r.status === 'indexing');
            if (!still) { clearInterval(pollTimer); pollTimer = null; }
          } catch(e) { clearInterval(pollTimer); pollTimer = null; }
        }, 3000);
      } else if (!needsPoll && pollTimer) {
        clearInterval(pollTimer);
        pollTimer = null;
      }
    }

    async function doAddRepo() {
      if (!addUrl.value.trim()) { toast('请输入 GitHub 地址', 'error'); return; }
      addLoading.value = true;
      try {
        const data = await api('/api/v1/repos', { method: 'POST', body: JSON.stringify({ url: addUrl.value.trim() }) });
        toast('已开始导入 ' + data.name, 'success');
        showAddModal.value = false;
        addUrl.value = '';
        await loadRepos();
      } catch(e) {}
      addLoading.value = false;
    }

    async function deleteRepo(r) {
      if (!confirm('确定删除仓库 ' + r.name + '？')) return;
      try {
        await api('/api/v1/repos/' + r.name, { method: 'DELETE' });
        toast('已删除 ' + r.name, 'success');
        if (currentRepo.value && currentRepo.value.name === r.name) {
          currentRepo.value = null;
          currentTab.value = 'overview';
        }
        await loadRepos();
      } catch(e) {}
    }

    function statusText(s) {
      return ({ pending:'等待中', cloning:'克隆中', indexing:'索引中', ready:'已就绪', error:'失败' })[s] || s;
    }
    function shortRepoName(name){
      const parts = (name||'').split('/');
      return parts.length===2 ? parts[1] : name;
    }
    function formatDate(s) {
      if (!s) return '';
      try {
        const d = new Date(s.replace(' ', 'T') + 'Z');
        const now = new Date();
        const diff = now - d;
        if (diff < 60000) return '刚刚';
        if (diff < 3600000) return Math.floor(diff/60000) + ' 分钟前';
        if (diff < 86400000) return Math.floor(diff/3600000) + ' 小时前';
        return d.toLocaleDateString('zh-CN', { month: 'short', day: 'numeric' });
      } catch { return s; }
    }
    function fmtTime(d=new Date()){
      const hh = String(d.getHours()).padStart(2,'0');
      const mm = String(d.getMinutes()).padStart(2,'0');
      return hh+':'+mm;
    }

    // Selection
    const currentRepo = ref(null);
    const currentTab = ref('overview');

    function selectRepo(r, tab) {
      currentRepo.value = r;
      currentTab.value = tab || 'overview';
      // Auto-close mobile drawer after selection
      if (window.innerWidth <= 768) sidebarOpen.value = false;
      if (tab === 'qa') {
        if (chatMessages.value.length === 0) {
          chatMessages.value.push({
            role: 'bot',
            content: '你好！我是代码文档助手，可以基于这个仓库的知识图谱回答你的问题。\n\n试着问我：\n  · 这个项目的核心模块有哪些？\n  · 路由是怎么注册的？\n  · X 函数的调用关系是什么？',
            time: fmtTime()
          });
        }
        nextTick(() => { if (chatBox.value) chatBox.value.scrollTop = chatBox.value.scrollHeight; });
      }
      if (tab === 'doc') {
        loadUserDocs();
        if (!Object.keys(docTemplates.value).length) loadDocTemplates();
        // Re-attach to a running docgen for this repo (if any), or auto-load
        // the most recent completed doc — so user sees progress / cached doc
        // even after navigating away and back.
        var running = docgenTasks.value.find(function(t){
          return t.repo === r.name && (t.status === 'queued' || t.status === 'running');
        });
        if (running) {
          docTaskId.value = running.task_id;
          docLoading.value = true;
          docContent.value = '';
          docStage.value = DOC_STAGE_LABELS[running.stage] || running.stage || '处理中';
          docProgress.value = running.progress || '';
          pollDocgenTask(running.task_id);
        } else {
          var lastDone = docgenTasks.value.find(function(t){
            return t.repo === r.name && t.status === 'done';
          });
          if (lastDone && (!docContent.value || docTaskId.value !== lastDone.task_id)) {
            docTaskId.value = lastDone.task_id;
            renderDocFromTask(lastDone.task_id);
          }
        }
      }
    }

    // Navigate to a non-repo top-level page (e.g. 'tasks').
    function goPage(page) {
      currentTab.value = page;
      currentRepo.value = null;
      if (window.innerWidth <= 768) sidebarOpen.value = false;
      if (page === 'tasks') loadDocgenTasks();
    }

    // Doc
    const docContent = ref('');
    const docLoading = ref(false);
    const docStage = ref('');      // human-readable stage label, e.g. "生成项目概览"
    const docProgress = ref('');   // secondary detail line, e.g. "分析模块 2/5: foo.bar"
    const docTaskId = ref('');     // active async docgen task id
    const docTemplates = ref({});
    const sectionLabels = ref({});
    const docTemplate = ref('default');
    const customSections = ref([]);

    async function loadDocTemplates(){
      if (!token.value) return;
      try {
        const data = await api('/api/v1/docgen/templates');
        docTemplates.value = data.templates || {};
        sectionLabels.value = data.sections || {};
        if (docTemplates.value[docTemplate.value] && Array.isArray(docTemplates.value[docTemplate.value].sections)) {
          customSections.value = [...docTemplates.value[docTemplate.value].sections];
        }
      } catch(e) { /* silent */ }
    }

    function onTemplateChange(){
      var t = docTemplates.value[docTemplate.value];
      if (t && Array.isArray(t.sections)) {
        customSections.value = [...t.sections];
      } else if (docTemplate.value === 'custom' && customSections.value.length === 0) {
        customSections.value = Object.keys(sectionLabels.value);
      }
    }

    // ---- 我的任务 (per-user docgen task list) ----------------------------
    const docgenTasks = ref([]);
    let _taskPollInterval = null;

    const TASK_STATUS_LABELS = {
      queued: '排队中', running: '生成中', done: '已完成', error: '失败',
    };

    function runningTaskCount() {
      return docgenTasks.value.filter(function(t){
        return t.status === 'queued' || t.status === 'running';
      }).length;
    }

    function fmtTaskTime(ts) {
      if (!ts) return '';
      try {
        var d = new Date(ts * 1000);
        var now = new Date();
        var diff = now - d;
        if (diff < 60000) return '刚刚';
        if (diff < 3600000) return Math.floor(diff/60000) + ' 分钟前';
        if (diff < 86400000) return Math.floor(diff/3600000) + ' 小时前';
        return d.toLocaleString('zh-CN', { month: 'short', day: 'numeric', hour: '2-digit', minute: '2-digit' });
      } catch(e) { return ''; }
    }

    function taskTemplateLabel(t) {
      var tmpl = docTemplates.value[t.template];
      return (tmpl && tmpl.name) || t.template || '';
    }

    async function loadDocgenTasks() {
      if (!token.value) return;
      try {
        var r = await api('/api/v1/docgen/tasks');
        docgenTasks.value = r.tasks || [];
        // Re-attach pollers for any tasks still in flight
        docgenTasks.value.forEach(function(t){
          if (t.status === 'queued' || t.status === 'running') {
            pollDocgenTask(t.task_id);
          }
        });
      } catch(e) { /* ignore */ }
    }

    function startTaskAutoReload() {
      if (_taskPollInterval) clearInterval(_taskPollInterval);
      _taskPollInterval = setInterval(function(){
        if (!token.value) return;
        var hasPending = docgenTasks.value.some(function(t){
          return t.status === 'queued' || t.status === 'running';
        });
        if (hasPending || currentTab.value === 'tasks') {
          loadDocgenTasks();
        }
      }, 5000);
    }

    async function viewTaskDocument(task) {
      var repo = repos.value.find(function(r){ return r.name === task.repo; });
      docTaskId.value = task.task_id;
      docContent.value = '';
      docLoading.value = true;
      docStage.value = '加载文档...';
      docProgress.value = '';
      if (repo) {
        currentRepo.value = repo;
      }
      currentTab.value = 'doc';
      if (window.innerWidth <= 768) sidebarOpen.value = false;
      await renderDocFromTask(task.task_id);
    }

    // User-uploaded reference docs
    const userDocs = ref([]);
    const userDocUploading = ref(false);
    const docDragOver = ref(false);

    async function loadUserDocs(){
      if (!currentRepo.value) { userDocs.value = []; return; }
      try {
        var data = await api('/api/v1/repos/' + currentRepo.value.name + '/docs');
        userDocs.value = data.docs || [];
      } catch(e) { userDocs.value = []; }
    }

    async function uploadUserDocFile(file){
      if (!currentRepo.value || !file) return;
      var fd = new FormData();
      fd.append('file', file);
      userDocUploading.value = true;
      try {
        var r = await fetch((window.__API_BASE__||'') + '/api/v1/repos/' + currentRepo.value.name + '/docs/upload', {
          method: 'POST',
          headers: { 'Authorization': 'Bearer ' + token.value },
          body: fd,
        });
        var data = await r.json();
        if (!r.ok) { toast(data.detail || '上传失败', 'error'); }
        else { toast('已上传 ' + data.filename, 'success'); await loadUserDocs(); }
      } catch(e) { toast('上传失败: ' + e.message, 'error'); }
      userDocUploading.value = false;
    }

    function onPickUserDoc(ev){
      var f = ev.target.files && ev.target.files[0];
      if (f) uploadUserDocFile(f);
      ev.target.value = '';
    }

    function onDropUserDoc(ev){
      docDragOver.value = false;
      var f = ev.dataTransfer && ev.dataTransfer.files && ev.dataTransfer.files[0];
      if (f) uploadUserDocFile(f);
    }

    async function deleteUserDoc(d){
      if (!confirm('删除文档 ' + d.filename + '？')) return;
      try {
        await api('/api/v1/repos/' + currentRepo.value.name + '/docs/' + d.id, { method: 'DELETE' });
        toast('已删除', 'success');
        await loadUserDocs();
      } catch(e) { toast('删除失败: ' + e.message, 'error'); }
    }

    function renderDoc(raw) {
      if (!raw) return '';
      let html = raw;
      // If looks like HTML already, just keep; otherwise process markdown lite
      if (!/<h[1-6]|<p|<ul|<table/i.test(html)) {
        // basic markdown to html
        html = html
          .replace(/```([\s\S]*?)```/g, (m, code) => '<pre><code>' + code.replace(/</g,'&lt;') + '</code></pre>')
          .replace(/^### (.*)$/gm, '<h3>$1</h3>')
          .replace(/^## (.*)$/gm, '<h2>$1</h2>')
          .replace(/^# (.*)$/gm, '<h2>$1</h2>')
          .replace(/^\* (.*)$/gm, '<li>$1</li>')
          .replace(/^- (.*)$/gm, '<li>$1</li>')
          .replace(/(<li>.*<\/li>\n?)+/g, m => '<ul>' + m + '</ul>')
          .replace(/`([^`]+)`/g, '<code>$1</code>')
          .replace(/\n\n/g, '</p><p>')
          .replace(/^(?!<)(.*)$/gm, '<p>$1</p>')
          .replace(/<p><\/p>/g, '')
          .replace(/<p>(<h[1-6])/g,'$1').replace(/(<\/h[1-6]>)<\/p>/g,'$1')
          .replace(/<p>(<ul)/g,'$1').replace(/(<\/ul>)<\/p>/g,'$1')
          .replace(/<p>(<pre)/g,'$1').replace(/(<\/pre>)<\/p>/g,'$1');
      }
      return html;
    }

    // Human-readable labels for backend stage codes
    const DOC_STAGE_LABELS = {
      queued: '排队中',
      loading: '加载仓库索引',
      overview: '生成项目概览',
      arch_diagram: '绘制系统架构图',
      class_uml: '绘制核心类 UML',
      call_flow: '绘制调用关系图',
      modules: '分析关键模块',
      routes: '整理接口文档',
      recommendations: '生成推荐问答',
      finalizing: '整理文档',
      done: '完成',
      error: '出错',
    };

    function renderDocMarkdown(md) {
      // Replace ASK markers with clickable buttons
      md = md.replace(/- ASK::(.+)/g, function(_, q){
        var safe = q.replace(/"/g, '&quot;').trim();
        return '<div class="ask-item"><button class="ask-btn" data-q="' + safe + '">问这个 →</button> ' + safe + '</div>';
      });
      if (window.marked && marked.parse) return marked.parse(md);
      return renderDoc(md);
    }

    // Fire-and-forget: submit task and immediately return control. A non-blocking
    // background poller updates docgenTasks state and (if user stays on this doc
    // tab for this task) updates docContent when done.
    async function doGenDoc() {
      if (!currentRepo.value) return;
      try {
        var payload = { repo: currentRepo.value.name, template: docTemplate.value };
        if (docTemplate.value === 'custom') {
          payload.sections = customSections.value;
        }
        const submit = await api('/api/v1/docgen', { method: 'POST', body: JSON.stringify(payload) });
        if (submit.error) {
          toast('生成失败: ' + submit.error, 'error');
          return;
        }
        var tName = (docTemplates.value[docTemplate.value] && docTemplates.value[docTemplate.value].name) || '完整文档';
        toast('任务已提交（' + tName + '），可继续操作其他功能', 'success');

        // Optimistically attach to this task on the doc tab so the user sees
        // progress immediately — but the polling is NON-BLOCKING.
        docTaskId.value = submit.task_id;
        docLoading.value = true;
        docContent.value = '';
        docStage.value = '排队中（第 ' + (submit.position || 1) + ' 位）';
        docProgress.value = '';

        await loadDocgenTasks();          // surface the new task in 我的任务
        pollDocgenTask(submit.task_id);   // fire-and-forget; does not await
      } catch(e) {
        toast('提交失败: ' + e.message, 'error');
      }
    }

    // Background poller — never blocks the UI. Updates docgenTasks state and,
    // if the user is currently viewing this task on the doc tab, updates the
    // doc content + stage labels too. Exits when status reaches done/error,
    // when the user logs out, or when polling errors out.
    const docgenPollers = ref({});  // task_id -> true while a poller is active
    async function pollDocgenTask(taskId) {
      if (docgenPollers.value[taskId]) return;  // already polling
      docgenPollers.value[taskId] = true;
      try {
        while (true) {
          await new Promise(function(r){ setTimeout(r, 3000); });
          if (!token.value) return;
          var s;
          try {
            s = await api('/api/v1/docgen/' + taskId + '/status');
          } catch(e) {
            return;
          }
          // Update the matching row in docgenTasks (if loaded)
          var idx = docgenTasks.value.findIndex(function(t){ return t.task_id === taskId; });
          if (idx >= 0) {
            docgenTasks.value[idx].status = s.status;
            docgenTasks.value[idx].stage = s.stage;
            docgenTasks.value[idx].progress = s.progress;
            docgenTasks.value[idx].position = s.position || 0;
            docgenTasks.value[idx].error = s.error || '';
          }
          // If the user is currently looking at this task's doc tab, update the live progress display
          if (docTaskId.value === taskId) {
            if (s.status === 'queued') {
              docStage.value = '排队中（第 ' + s.position + ' 位 / 共 ' + s.queue_total + ' 个任务）';
              docProgress.value = '等待前面的任务完成...';
            } else if (s.status === 'running') {
              docStage.value = DOC_STAGE_LABELS[s.stage] || s.stage || '处理中';
              docProgress.value = s.progress || '';
            }
          }
          if (s.status === 'done') {
            toast(((docgenTasks.value[idx] && docgenTasks.value[idx].repo) || '') + ' 文档已生成', 'success');
            if (docTaskId.value === taskId) {
              await renderDocFromTask(taskId);
            }
            return;
          }
          if (s.status === 'error') {
            toast('生成失败: ' + (s.error || '未知错误'), 'error');
            if (docTaskId.value === taskId) {
              docContent.value = '<p style="color:#ff3b30">生成失败: ' + (s.error || '未知错误') + '</p>';
              docStage.value = '';
              docProgress.value = '';
              docLoading.value = false;
            }
            return;
          }
        }
      } finally {
        delete docgenPollers.value[taskId];
      }
    }

    // Load the cached document for a task and render it into the doc pane.
    async function renderDocFromTask(taskId) {
      try {
        var r = await api('/api/v1/docgen/' + taskId + '/document');
        if (r.status !== 'done' || !r.document) {
          // still in progress — show stage if attached
          if (docTaskId.value === taskId) {
            docStage.value = DOC_STAGE_LABELS[r.stage] || r.stage || '处理中';
            docProgress.value = r.progress || '';
          }
          return;
        }
        docContent.value = renderDocMarkdown(r.document);
        docLoading.value = false;
        docStage.value = '';
        docProgress.value = '';
        await nextTick();
        document.querySelectorAll('.doc-body .ask-btn').forEach(function(btn){
          btn.onclick = function(){
            var q = btn.dataset.q || '';
            currentTab.value = 'qa';
            chatInput.value = q;
            nextTick().then(function(){ doAsk(); });
          };
        });
      } catch(e) {
        toast('加载文档失败: ' + e.message, 'error');
      }
    }

    // Chat
    const chatMessages = ref([]);
    const expandedInvalid = ref({});
    const chatInput = ref('');
    const chatLoading = ref(false);
    const chatBox = ref(null);

    async function doAsk() {
      const q = chatInput.value.trim();
      if (!q || !currentRepo.value || chatLoading.value) return;
      chatMessages.value.push({ role: 'me', content: q, time: fmtTime() });
      chatInput.value = '';
      chatLoading.value = true;
      await nextTick();
      if (chatBox.value) chatBox.value.scrollTop = chatBox.value.scrollHeight;
      try {
        const data = await api('/api/v1/ask', { method: 'POST', body: JSON.stringify({ question: q, repo: currentRepo.value.name }) });
        chatMessages.value.push({
          role: 'bot',
          content: data.answer || '暂无回答',
          time: fmtTime(),
          groundedness: data.groundedness,
          total_refs: data.total_refs,
          valid_refs_count: data.valid_refs_count,
          invalid_refs: data.invalid_refs || [],
        });
      } catch(e) {
        chatMessages.value.push({ role: 'bot', content: '请求失败: ' + e.message, time: fmtTime() });
      }
      chatLoading.value = false;
      await nextTick();
      if (chatBox.value) chatBox.value.scrollTop = chatBox.value.scrollHeight;
    }

    // Reset chat when switching repo. Doc state is intentionally NOT reset
    // here — selectRepo() takes care of re-attaching to a running task or
    // auto-loading the most recent completed doc for the new repo.
    watch(currentRepo, (val, old) => {
      if (!val || !old || val.name !== old.name) {
        chatMessages.value = [];
        userDocs.value = [];
        // Clear the doc pane only if there's no in-flight task for the new repo
        // and no cached doc to show; selectRepo() will repopulate when known.
        var hasRunning = val && docgenTasks.value.some(function(t){
          return t.repo === val.name && (t.status === 'queued' || t.status === 'running');
        });
        var hasDone = val && docgenTasks.value.some(function(t){
          return t.repo === val.name && t.status === 'done';
        });
        if (!hasRunning && !hasDone) {
          docContent.value = '';
          docStage.value = '';
          docProgress.value = '';
          docTaskId.value = '';
          docLoading.value = false;
        }
        if (val) loadUserDocs();
      }
    });

    onMounted(() => {
      if (token.value) {
        loadRepos();
        loadDocTemplates();
        loadDocgenTasks();
        startTaskAutoReload();
      }
    });

    // Whenever the user logs in/out, refresh tasks + auto-poll state.
    watch(token, function(v){
      if (v) {
        loadDocgenTasks();
        startTaskAutoReload();
      } else if (_taskPollInterval) {
        clearInterval(_taskPollInterval);
        _taskPollInterval = null;
        docgenTasks.value = [];
      }
    });
    function groundednessClass(g) {
      if (g >= 0.85) return 'gnd-high';
      if (g >= 0.6) return 'gnd-mid';
      return 'gnd-low';
    }
    function groundednessTitle(m) {
      const v = m.valid_refs_count || 0;
      const t = m.total_refs || 0;
      const missing = (m.invalid_refs || []).length;
      return `共 ${t} 个代码引用，其中 ${v} 个已在知识图谱中验证` + (missing ? `，${missing} 个未匹配` : '');
    }
    function toggleInvalidRefs(i) {
      expandedInvalid.value = { ...expandedInvalid.value, [i]: !expandedInvalid.value[i] };
    }



    return {
      token, username, sidebarOpen, authMode, authForm, authLoading, authError, doAuth, logout,
      toasts,
      repos, reposLoading, showAddModal, addTab, addUrl, addLoading,
      uploadName, uploadFile, uploadLoading, dragOver, fileInput,
      closeAddModal, onFilePicked, onFileDropped, doUploadRepo, humanSize,
      doAddRepo, deleteRepo, statusText, shortRepoName, formatDate,
      currentRepo, currentTab, selectRepo, goPage,
      docContent, docLoading, docStage, docProgress, doGenDoc,
      docTemplates, sectionLabels, docTemplate, customSections, onTemplateChange,
      userDocs, userDocUploading, docDragOver,
      onPickUserDoc, onDropUserDoc, deleteUserDoc,
      chatMessages, expandedInvalid, groundednessClass, groundednessTitle, toggleInvalidRefs, chatInput, chatLoading, chatBox, doAsk,
      docgenTasks, runningTaskCount, viewTaskDocument,
      TASK_STATUS_LABELS, DOC_STAGE_LABELS, fmtTaskTime, taskTemplateLabel,
    };
  }
}).mount('#app');
</script>
</body>
</html>"""

@app.get("/", response_class=HTMLResponse)
def index():
    return FRONTEND_HTML


# PLATFORM_PORTAL_ROUTE 综合门户:codedoc + devbot 一个入口
_PLATFORM_HTML = """<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>研发智能平台</title>
<style>
*{margin:0;padding:0;box-sizing:border-box;font-family:-apple-system,BlinkMacSystemFont,"Segoe UI","Microsoft YaHei",sans-serif}
body{min-height:100vh;display:flex;flex-direction:column;align-items:center;justify-content:center;background:linear-gradient(135deg,#1e293b 0%,#0f172a 100%);color:#e2e8f0;padding:40px}
.hero{text-align:center;margin-bottom:46px}
.hero h1{font-size:38px;font-weight:700;letter-spacing:-.5px;background:linear-gradient(90deg,#60a5fa,#a78bfa);-webkit-background-clip:text;background-clip:text;-webkit-text-fill-color:transparent}
.hero p{margin-top:12px;font-size:15px;color:#94a3b8}
.cards{display:flex;gap:26px;flex-wrap:wrap;justify-content:center;max-width:860px}
.card{display:block;width:378px;background:rgba(30,41,59,.7);border:1px solid rgba(148,163,184,.18);border-radius:18px;padding:30px;text-decoration:none;color:inherit;transition:.22s;backdrop-filter:blur(8px)}
.card:hover{transform:translateY(-5px);border-color:rgba(96,165,250,.55);box-shadow:0 16px 40px rgba(0,0,0,.42)}
.ic{width:52px;height:52px;border-radius:13px;display:flex;align-items:center;justify-content:center;font-size:25px;margin-bottom:18px;font-weight:700;color:#dbeafe}
.ic.cd{background:rgba(96,165,250,.16)}
.ic.db{background:rgba(167,139,250,.16)}
.card h2{font-size:22px;font-weight:650;margin-bottom:9px}
.card p{font-size:13.5px;line-height:1.65;color:#94a3b8;min-height:70px}
.tags{margin-top:14px;display:flex;gap:8px;flex-wrap:wrap}
.tag{font-size:11px;padding:4px 10px;border-radius:20px;background:rgba(148,163,184,.12);color:#cbd5e1}
.go{margin-top:18px;font-size:14px;font-weight:600;color:#60a5fa}
.foot{margin-top:42px;font-size:12px;color:#64748b}
</style>
</head>
<body>
<div class="hero">
  <h1>研发智能平台</h1>
  <p>代码知识图谱问答 · PR 智能评审 — 一套底座（PostgreSQL + 图谱 + Agent 编排）</p>
</div>
<div class="cards">
  <a class="card" href="/">
    <div class="ic cd">CD</div>
    <h2>CodeDoc</h2>
    <p>团队级代码知识库。AST 多语言图谱 + 意图分流混合检索 + 多仓多 Agent 问答，逐符号抗幻觉核验，自动出设计文档。</p>
    <div class="tags"><span class="tag">代码问答</span><span class="tag">文档生成</span><span class="tag">影响分析</span></div>
    <div class="go">进入 CodeDoc →</div>
  </a>
  <a class="card" id="devbot-link" href="#">
    <div class="ic db">DB</div>
    <h2>DevBot</h2>
    <p>研发能效 Agent 平台。LangGraph 多 Critic 并行 PR 评审 + 沙箱实跑验证，代码/单测生成，经 MCP 接 CodeDoc 图谱取变更影响。</p>
    <div class="tags"><span class="tag">PR 评审</span><span class="tag">代码生成</span><span class="tag">单测生成</span></div>
    <div class="go">进入 DevBot →</div>
  </a>
</div>
<div class="foot">codedoc:8501 · devbot:8502 · 同一 PostgreSQL 底座</div>
<script>document.getElementById("devbot-link").href="http://"+location.hostname+":8502/";</script>
</body>
</html>"""

_DEVBOT_REDIRECT = """<!DOCTYPE html><html lang="zh-CN"><head><meta charset="UTF-8">
<script>location.replace("http://"+location.hostname+":8502/");</script>
</head><body style="font-family:sans-serif;padding:40px;color:#555">正在跳转到 DevBot…若未跳转 <a id="l" href="#">点此</a>
<script>document.getElementById("l").href="http://"+location.hostname+":8502/";</script></body></html>"""


@app.get("/platform/", response_class=HTMLResponse)
@app.get("/platform", response_class=HTMLResponse)
def platform_home():
    return _PLATFORM_HTML


@app.get("/devbot/", response_class=HTMLResponse)
@app.get("/devbot", response_class=HTMLResponse)
def devbot_redirect():
    return _DEVBOT_REDIRECT




def _start_watchdog():
    """watchdog 增量同步:监听 REPOS_DIR,文件变 -> 重建该仓内存图谱+落盘(纯 CPU,不重嵌入)。"""
    import os, threading, pathlib
    if os.environ.get("CODEDOC_WATCHDOG", "1") != "1":
        return

    def _loop():
        try:
            from watchfiles import watch
            from codedoc.config import load_config
            from codedoc.parser.runner import parse_repo
            from codedoc.graph.memory_backend import MemoryGraphStore
            from codedoc.graph import graph_persist
            root = str(REPOS_DIR)
            print("[watchdog] watching", root)
            for changes in watch(root, debounce=2000, step=500):
                repos = set()
                for _chg, path in changes:
                    if not path.endswith((".py", ".java", ".js", ".ts")):
                        continue
                    try:
                        rel = pathlib.Path(path).relative_to(root)
                        if rel.parts:
                            repos.add(rel.parts[0])
                    except Exception:
                        pass
                for rd in repos:
                    d = pathlib.Path(root) / rd
                    if not d.is_dir():
                        continue
                    try:
                        cfg = load_config(str(d))
                        nodes, edges = parse_repo(cfg)
                        st = MemoryGraphStore(cfg)
                        st.upsert_nodes(nodes); st.upsert_edges(edges)
                        for _nm, _info in list(indexed_repos.items()):
                            if _info.get("path") == str(d):
                                _info["store"] = st
                                _info["nodes"] = len(nodes); _info["edges"] = len(edges)
                                _info.pop("_comm_summaries", None)
                        graph_persist.save_graph(rd, nodes, edges)
                        print("[watchdog] re-synced graph repo=%s (%d nodes)" % (rd, len(nodes)))
                    except Exception as _we:
                        print("[watchdog] sync failed for %s: %s" % (rd, _we))
        except Exception as _e:
            print("[watchdog] disabled:", _e)

    threading.Thread(target=_loop, daemon=True).start()


_start_watchdog()
_start_summary_worker()
_start_docgen_worker()
_start_repo_worker()
